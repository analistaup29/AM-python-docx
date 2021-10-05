#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Sep 21 15:54:48 2021
"""

# Importar librerías ----------------------------------------------------------
import docx
import pandas as pd
import nums_from_string
import os
from datetime import datetime
from pyprojroot import here
from janitor import clean_names # pip install pyjanitor
from pathlib import Path

# Fechas de corte -------------------------------------------------------------
# Importamos los nombres de los archivos en la carpeta input
lista_archivos = os.listdir(Path(here() / "input"))

## A) Base disponibilidad
# Nos quedamos con el nombre de archivo para la base de disponibilidad
fecha_corte_disponibilidad = str([s for s in lista_archivos if "Disponibilidad_Presupuestal" in s])
# Extraemos la fecha del nombre de archivo
fecha_corte_disponibilidad = nums_from_string.get_numeric_string_tokens(fecha_corte_disponibilidad)
# Convertimos a formato string
fecha_corte_disponibilidad = ''.join(fecha_corte_disponibilidad) 
# Convertimos a formato numérico
fecha_corte_disponibilidad_date = datetime.strptime(fecha_corte_disponibilidad, '%Y%m%d').date()

## B) Siaf de mascarillas
fecha_corte_mascarillas = "2021-09-20"

# C) Compromisos de desempeño
fecha_corte_compromisos = "2021-09-21"

# Creamos tabla con fechas de corte
tabla_fechas_corte = (
    ("Intervenciones pedagógicas", fecha_corte_disponibilidad_date),
    ("Mascarillas y protectores faciales", fecha_corte_mascarillas),
    ("Compromisos de desempeño", fecha_corte_compromisos)
)

# Transformación de Datasets --------------------------------------------------

# Base de datos región
## Cargamos nombres de regiones
nombre_regiones = pd.read_excel(here() / "input/nombre_regiones.xlsx")

# A) Base de disponibilidad
## Cargamos base de disponibilidad
data_intervenciones = pd.read_excel(here() / f"input/Disponibilidad_Presupuestal_{fecha_corte_disponibilidad}interv.xlsx")
data_intervenciones = clean_names(data_intervenciones) # Normalizamos nombres

# Mantenemos variables de interés (PIM, DEVENGADO, COMPROMETIDO CERTIFICADO) y 
# colapsamos a nivel de Region, Intervencion Pedagogica y Cas-No-Cas
data_intervenciones = data_intervenciones[["region", "cas_no_cas","intervencion_pedagogica", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_certificado_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"]]. \
   groupby(by = ["region", "cas_no_cas", "intervencion_pedagogica"] , as_index=False).sum()

# Eliminamos filas de "No hay Intervenciones pedagogicas"
data_intervenciones = data_intervenciones[data_intervenciones['intervencion_pedagogica'] != "No hay Intervenciones Pedagógicas"]

# (PENDIENTE) FOR LOOP de número de intervenciones pedagógicas, Mascarillas y CDD
numero_intervenciones = "8" ## PENDIENTE

# B) Siaf de mascarillas
## Cargamos la base insumo de mascarillas
data_mascarillas = pd.read_excel(here() / "input/Incorporación_DU_SIAF_20210921.xlsx", sheet_name='Sheet1')
data_mascarillas = clean_names(data_mascarillas) # Normalizamos nombres

# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
tabla_mascarillas = data_mascarillas[["region","nom_ue","certificado","comprometido_anual","devengado","transferencia"]]. \
   groupby(by = ["region", "nom_ue"], as_index=False).sum()

tabla_mascarillas["region"] = tabla_mascarillas["region"].str.split(". ", n=1).apply(lambda l: "".join(l[1]))

# C) Compromisos de desempeño

## Cargamos data de compromisos de desempeño
data_cdd = pd.read_excel(here() / "input/regiones_BD_CDD.xlsx")
data_cdd = clean_names(data_cdd) # Normalizamos nombres
data_cdd["pliego"] = data_cdd["pliego"].str.split(". ", n=1, expand = True)
data_cdd['pliego'] = data_cdd['pliego'].astype('int64') # Convertimos ubigeo a integer

# Corregimos genericas
data_cdd["generica"] = data_cdd["generica"].replace("3. BIENES Y SERVICIOS", "2.3. BIENES Y SERVICIOS")
data_cdd["generica"] = data_cdd["generica"].replace("6. ADQUISICION DE ACTIVOS NO FINANCIEROS", "2.6. ADQUISICION DE ACTIVOS NO FINANCIEROS")	

# Hacemos merge con base de datos región
data_cdd = data_cdd.merge(right = nombre_regiones, how="left", on = "pliego")

# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
tabla_cdd = data_cdd[["region", "unidad_ejecutora", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region", "programa_presupuestal", "generica"], as_index=False).sum()


# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS", "TACNA", "AREQUIPA"]

# Generar docx ----------------------------------------------------------------

# For loop para cada región
for region in lista_regiones:
    ##########################################################################
    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_intervenciones['region'] == region #Seleccionar region
    tabla_intervenciones = data_intervenciones[region_seleccionada]    
    pim_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    ejecucion_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    # Generamos la tabla "tabla1_region" - mantiene la región i de la lista de regiones
    tabla_intervenciones_formato = data_intervenciones[region_seleccionada]
    # Formato para la tabla
    formato_tabla_intervenciones = {
        "cas_no_cas": "{}",
        "intervencion_pedagogica" : "{}",
        f"pim_reporte_siaf_{fecha_corte_disponibilidad}": "{:,.0f}",
        f"presupuesto_certificado_reporte_siaf_{fecha_corte_disponibilidad}" : "{:,.0f}",
        f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}" : "{:,.0f}",
        f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}": "{:,.0f}",
        }
    tabla_intervenciones_formato = tabla_intervenciones_formato.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})        
    ##########################################################################
    # Generamos la tabla "tabla1_mascarilla" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = tabla_mascarillas['region'] == region
    tabla2_region = tabla_mascarillas[region_seleccionada]
    # Generamos los indicadores de PIM y ejecución de intervenciones
    transferencia_mascarilla = str('{:,.1f}'.format(tabla2_region["transferencia"].sum()/1000000))
    devengado_mascarillas=str('{:.1%}'.format(tabla2_region["devengado"].sum()/tabla2_region["transferencia"].sum()))
    ##########################################################################
    # Generamos la tabla "tabla1_cdd" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = tabla_cdd['region'] == region
    tabla3_region = tabla_cdd[region_seleccionada]
    # Generamos CDD transferido
    cdd_transferido = str('{:,.0f}'.format(tabla3_region["monto"].sum()))
    cdd_acciones_centrales = "88,888"
    ##########################################################################
    # Incluimos el código del Documento
    document = docx.Document(here() / "input/formato.docx") # Creación del documento en base al template
    title=document.add_heading('AYUDA MEMORIA', 0) #Título del documento
    run = title.add_run()
    run.add_break()
    title.add_run('REGIÓN ')
    title.add_run(region)
    run = title.add_run()
    run.add_break()
    title.add_run(datetime.today().strftime('%d-%m-%y'))
    ##########################################################################
    # Incluimos sección 1 de intervenciones pedagógicas
    document.add_heading("1. Intervenciones pedagógicas", level=1) # 1) Intervenciones pedagógicas
    interv_parrafo1 = document.add_paragraph(
    "Las Unidades Ejecutoras de Educación de la región " , style="List Bullet")
    interv_parrafo1.add_run(region)
    interv_parrafo1.add_run(" vienen implementando ")
    interv_parrafo1.add_run(numero_intervenciones)
    interv_parrafo1.add_run(
    " intervenciones y acciones pedagógicas en el Año 2021, en el marco de la \
    Norma Técnica “Disposiciones para la implementación de las intervenciones \
    y acciones pedagógicas del Ministerio de Educación en los Gobiernos Regionales \
    y Lima Metropolitana en el Año Fiscal 2021”, aprobada mediante \
    RM N° 043-2021-MINEDU y modificada RM N° 159-2021-MINEDU." )
    interv_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    # Incluimos tabla 1 intervenciones
    tabla1_interv = document.add_table(tabla_intervenciones_formato.shape[0]+1, tabla_intervenciones_formato.shape[1])
    tabla1_interv.style = "Colorful List Accent 1"
    ## Header de la tabla
    row = tabla1_interv.rows[0].cells
    row[0].text = "CAS - NO CAS"
    row[1].text = "Intervencion Pedagogica"
    row[2].text = "PIM"
    row[3].text = "Certificado"
    row[4].text = "Comprometido"
    row[5].text = "Devengado"
    ## Contenido de la tabla
    for i in range(tabla_intervenciones_formato.shape[0]):
        for j in range(tabla_intervenciones_formato.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato.values[i,j])
    interv_parrafo2 = document.add_paragraph(
    "A través de los Decretos Supremos N°s 092, 169, 189, 209 y 210-2021-EF, \
    se realizaron todas las transferencias de partidas programadas para el Año \
    Fiscal 2021 para el financiamiento de las intervenciones y acciones pedagógicas \
    hasta el 31 de diciembre.", style="List Bullet")
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    interv_parrafo3 = document.add_paragraph(
    "Es importante considerar que la ejecución en los Contratos Administrativos \
    de Servicios (CAS) se ha visto afectada por la vigencia de la Ley N° 31131. \
    Por otro lado, la ejecución en bienes y servicios (excluyendo CAS) es menor \
    a lo esperado debido al bajo número de IIEE que brindan servicios presenciales \
    y semipresenciales a nivel nacional.", style="List Bullet")
    interv_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    interv_parrafo4 = document.add_paragraph(
    "Las Unidades Ejecutoras de Educación de la región " , style="List Bullet")
    interv_parrafo4.add_run(region)
    interv_parrafo4.add_run(" cuentan con ")
    interv_parrafo4.add_run(pim_intervenciones_region)
    interv_parrafo4.add_run(
    " millones en su Presupuesto Institucional Modificado (PIM) para el \
    financiamiento de intervenciones y acciones pedagógicas, de los cuales se han ejecutado S/ ")
    interv_parrafo4.add_run(ejecucion_intervenciones_region)
    interv_parrafo4.add_run(".")
    interv_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    ###########################################################################
    # Incluimos sección 2 Mascarillas y protectores faciales
    document.add_heading("2. Mascarillas y protectores faciales", level=1)
    mascarillas_parrafo1 = document.add_paragraph(
    "Mediante el Decreto de Urgencia N° 021-2021 y la Resolución de Secretaría \
General N° 047-2021-MINEDU, se transfirieron S/ ", style="List Bullet")
    mascarillas_parrafo1.add_run(transferencia_mascarilla)
    mascarillas_parrafo1.add_run(" millones de soles para \
la adquisición y distribución de mascarillas faciales textiles de uso \
comunitario para estudiantes y personal que labora en instituciones \
educativas públicas, así como protectores faciales para el mencionado \
personal.")
    mascarillas_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    mascarillas_parrafo2 = document.add_paragraph(
    "La adquisición de mascarillas y protectores faciales es condición necesaria \
para el retorno seguro a los servicios educativos presenciales y \
semipresenciales, según lo dispuesto por las “Disposiciones para la \
prestación del servicio en las instituciones y programas educativos \
públicos y privados de la Educación Básica de los ámbitos urbanos y \
rurales, en el marco de la emergencia sanitaria de la COVID-19”, \
aprobado mediante Resolución Ministerial N° 121-2021- MINEDU y modificado \
con Resoluciones Ministeriales N° 199-2021-MINEDU y N° 273-2021- MINEDU.", style="List Bullet")
    mascarillas_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    mascarillas_parrafo3 = document.add_paragraph(
    "Con fecha de corte al ", style="List Bullet")
    mascarillas_parrafo3.add_run(fecha_corte_mascarillas)
    mascarillas_parrafo3.add_run(
    ", la ejecución a nivel regional de los recursos de mascarillas faciales \
textiles protectores faciales fue del ")
    mascarillas_parrafo3.add_run(devengado_mascarillas)
    mascarillas_parrafo3.add_run(" (devengado) según se presenta a continuación:")
    mascarillas_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    # Incluir tabla 2 mascarillas
    tabla2_interv = document.add_table(tabla2_region.shape[0]+1, tabla2_region.shape[1])
    tabla2_interv.style = "Colorful List Accent 1"
    for j in range(tabla2_region.shape[-1]):
        tabla2_interv.cell(0,j).text = tabla2_region.columns[j]
    for i in range(tabla2_region.shape[0]):
        for j in range(tabla2_region.shape[-1]):
            tabla2_interv.cell(i+1,j).text = str(tabla2_region.values[i,j])
    ##########################################################################
    # Incluimos sección 3 Compromisos de desempeño
    document.add_heading("3. Compromisos de desempeño", level=1)
    cdd_parrafo1 = document.add_paragraph(
        "En el marco de la Norma Técnica para la implementación del mecanismo \
denominado Compromisos de Desempeño 2021, aprobada por Resolución Ministerial \
N° 042-2021-MINEDU y modificada por la Resolución Ministerial N° 160-2021-MINEDU, \
se han realizado transferencias de partidas a favor de las Unidades Ejecutoras de \
Educación del Gobierno Regional de ", style="List Bullet")
    cdd_parrafo1.add_run(region)
    cdd_parrafo1.add_run("por la suma de")
    cdd_parrafo1.add_run(cdd_transferido)
    cdd_parrafo1.add_run(" De dichos recursos, ")
    cdd_parrafo1.add_run(cdd_acciones_centrales)
    cdd_parrafo1.add_run(" corresponden a las acciones centrales, según el \
siguiente detalle")
    cdd_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    tabla1_cdd = document.add_table(tabla3_region.shape[0]+1, tabla3_region.shape[1])
    tabla1_cdd.style = "Colorful List Accent 1"
    for j in range(tabla3_region.shape[-1]):
        tabla1_cdd.cell(0,j).text = tabla3_region.columns[j]
    for i in range(tabla3_region.shape[0]):
        for j in range(tabla3_region.shape[-1]):
            tabla1_cdd.cell(i+1,j).text = str(tabla3_region.values[i,j])
    ##########################################################################
    # Incluimos anexo de fechas de actualización
    document.add_heading("4. Anexos", level=1)
    anexo_parrafo1 = document.add_paragraph("Las fechas de actualización para las \
secciones del documento se presentan en la tabla siguiente")
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Sección"
    hdr_cells[1].text = "Fecha de actualización"
    for id, name in tabla_fechas_corte:
        row = table.add_row().cells
        row[0].text = str(id)
        row[1].text = str(name)
    ##########################################################################
    # Guardamos documento
    document.save(here() / f'output/{region}_AM.docx')
