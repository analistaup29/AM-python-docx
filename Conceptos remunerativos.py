# -*- coding: utf-8 -*-
"""
Created on Mon Oct  4 15:42:30 2021

@author: Hugo Fernandez y Carlos Ramirez
"""

# Librerias

import docx
import pandas as pd
from datetime import datetime
from pyprojroot import here # pip install pypro
from janitor import clean_names # pip install pyjanitor
import os
import numpy as np

pd.options.display.float_format = "{:,.2f}".format

# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS", "TACNA", "AREQUIPA"]

# Import las bases
os.chdir('D:\Minedu\AM-python-docx') # Directorio del proyecto 
    # Excel CONCEPTOS CONSOLIDADOS
        #Encargaturas
df_consolidado_enc = pd.read_excel('input\Bases plazas\CONCEPTOS CONSOLIDADOS.xlsx',sheet_name = 'ENC-CONSOLIDADO-VF') 
df_consolidado_enc.fillna(0, inplace =  True)
df_consolidado_enc['COSTO'] = df_consolidado_enc['COSTO-TRAMO I'] + df_consolidado_enc['COSTO-TRAMO2']
df_consolidado_enc['PROGRAMADO POR MINEDU'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']+ df_consolidado_enc['NM ENCARGATURAS']
df_consolidado_enc.rename(columns={'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                   'TRANSFERENCIA DS 217':'TRANSFERENCIAPOR DS 217-2021',
                                   'UNIDADEJECUTORA':'UNIDAD EJECUTORA'},inplace=True)
df_consolidado_enc['APMeINCREMENTOS'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']

tabla_encargaturas = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIAPOR DS 217-2021',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]

tabla_encargaturas_resumen = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIAPOR DS 217-2021',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]

        #Asignaciones temporales
df_consolidado_at = pd.read_excel('input\Bases plazas\CONCEPTOS CONSOLIDADOS.xlsx',sheet_name = 'AT-CONSOLIDADO-VF')   
df_consolidado_at.fillna(0, inplace =  True)      
df_consolidado_at.rename(columns={'REGIÓN':'REGION',
                                  'COSTO-TRAMO I':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                  'TRANSFERENCIA DS 187':'TRANSFERENCIA POR DS 187-2021'
    },inplace=True)
df_at=df_consolidado_at[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                         'PROGRAMADO POR EL PLIEGO REGIONAL','TRANSFERENCIA POR DS 187-2021']]

        #Beneficios sociales
df_consolidado_bf = pd.read_excel('input\Bases plazas\CONCEPTOS CONSOLIDADOS.xlsx',sheet_name = 'BS-CONSOLIDADO-VF')           
df_consolidado_bf.fillna(0,inplace = True)
df_consolidado_bf['COSTO BENEFICIARIOS 2020 Y 2021'] = df_consolidado_bf['LISTAS-2021'] + df_consolidado_bf['TRAMO I-BS 2020'] + df_consolidado_bf['TRAMO II-BS 2021']
df_consolidado_bf.rename(columns={'REGIÓN':'REGION',
                                  'APM':'PROGRAMADO POR MINEDU (BENEFICIARIOS 2021)',
                                  'TRANSFERENCIA DS 072-BS 2020':'TRANSFERENCIA POR DS 072-2021 (BENEFICIARIOS 2020)',
                                  'TRANSFERENCIA DS 256-BS 2021':'TRANSFERENCIA POR DS 256-2021 (BENEFICIARIOS 2021)'
    },inplace=True)
df_bs = df_consolidado_bf[['REGION','UNIDAD EJECUTORA','COSTO BENEFICIARIOS 2020 Y 2021',
                           'PROGRAMADO POR MINEDU (BENEFICIARIOS 2021)',
                           'TRANSFERENCIA POR DS 072-2021 (BENEFICIARIOS 2020)',
                           'TRANSFERENCIA POR DS 256-2021 (BENEFICIARIOS 2021)']]


# For loop para cada región
for region in lista_regiones:
    
    ########################
    # Tablas Encargaturas #
    ########################

    region_seleccionada = df_consolidado_enc['REGION'] == region
    tabla1 = tabla_encargaturas[region_seleccionada]
    costo_enc = str('{:,.0f}'.format(tabla1["COSTO"].sum())) 
    apmeincre = str('{:,.0f}'.format(tabla1["APMeINCREMENTOS"].sum()))
    prog_gore = str('{:,.0f}'.format(tabla1["PROGRAMADO POR EL PLIEGO REGIONAL"].sum()))
    nm_enca = str('{:,.0f}'.format(tabla1['NM ENCARGATURAS'].sum()))
    ds_217 = str('{:,.0f}'.format(tabla1['TRANSFERENCIAPOR DS 217-2021'].sum()))
    
    tabla_encargaturas_resumen = tabla1.groupby(['UNIDAD EJECUTORA'], as_index=False).sum()
    tabla_encargaturas_resumen = tabla_encargaturas_resumen[['UNIDAD EJECUTORA',
                                                             'COSTO', 
                                                             'PROGRAMADO POR MINEDU', 
                                                             'PROGRAMADO POR EL PLIEGO REGIONAL',
                                                             'TRANSFERENCIAPOR DS 217-2021']] 
    tabla_encargaturas_resumen = tabla_encargaturas_resumen.round(2)

    ###################################
    #  Tablas Asignaciones Temporales #
    ###################################
    
    region_seleccionada = df_at['REGION'] == region
    tabla2 = df_at[region_seleccionada]
    costo_at = str('{:,.0f}'.format(tabla2["COSTO"].sum()))
    apm_at = str('{:,.0f}'.format(tabla2["PROGRAMADO POR MINEDU"].sum())) 
    diploma_gore_at = str('{:,.0f}'.format(tabla2["PROGRAMADO POR MINEDU"].sum())) 
    ds_187_at = str('{:,.0f}'.format(tabla2["TRANSFERENCIA POR DS 187-2021"].sum())) 
    tabla_at_resumen = tabla2.groupby(['UNIDAD EJECUTORA'], as_index=False).sum() 
    tabla_at_resumen = tabla_at_resumen.round(2)

    ###################################
    #    Tablas Beneficios Sociales   #
    ###################################
    
    region_seleccionada = df_bs['REGION'] == region
    tabla3 = df_bs[region_seleccionada]
    costo_bs = str('{:,.0f}'.format(tabla3["COSTO BENEFICIARIOS 2020 Y 2021"].sum()))
    apm_bs = str('{:,.0f}'.format(tabla3["PROGRAMADO POR MINEDU (BENEFICIARIOS 2021)"].sum()))    
    ds_72_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS 072-2021 (BENEFICIARIOS 2020)"].sum())) 
    ds_256_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS 256-2021 (BENEFICIARIOS 2021)"].sum()))
    tabla_bs_resumen = tabla3.groupby(['UNIDAD EJECUTORA'], as_index=False).sum()
    tabla_bs_resumen = tabla_bs_resumen.round(2)
    
#-----------------------------------------------------------------------------#

    # Incluimos el código del Documento
    document = docx.Document(here() / "input/Formato.docx") # Creación del documento en base al template
    title=document.add_heading('AM TEMAS PRESUPUESTALES - REGION ', 0) #Título del documento
    title.add_run(region)
    run = title.add_run()
    run.add_break()
    run = title.add_run()
    run.add_break()
    title.add_run(datetime.today().strftime('%d-%m-%y'))
    
#-----------------------------------------------------------------------------#

    document.add_heading("Sobre el financiamiento de conceptos remunerativos", level=1)
    run.underline = True
    
    ########################
    ##   1) Encargaturas  ##
    ########################
    document.add_heading("1.Pago de Encargaturas", level=1)
    
    encarg_parrafo1 = document.add_paragraph("Para la región ")
    encarg_parrafo1.add_run(region)
    encarg_parrafo1.add_run(', por concepto de encargaturas, se ha calculado para el ')
    encarg_parrafo1.add_run(datetime.today().strftime('%Y'))
    encarg_parrafo1.add_run(', un costo de S/.')
    encarg_parrafo1.add_run(f'{costo_enc}') #Insertar valor de base de datos
    encarg_parrafo1.add_run(
    ' que incluye la Jornada de Trabajo Adicional de 10 horas \
la carga social vinculada y la asignación por cargo de \
los profesores que asumen cargos de mayor responsabilidad \
mediante encargaturas')
    encarg_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo2 = document.add_paragraph('Para financiar estos conceptos en el 2020,')
    encarg_parrafo2.add_run('el MINEDU gestionó una programación directa de recursos\
    en el PIA 2021 de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo2.add_run(region)  
    encarg_parrafo2.add_run(' por el monto de S/.')
    encarg_parrafo2.add_run(f'{apmeincre}') #Insertar valor de base de datos
    encarg_parrafo2.add_run('  en la finalidad 0267929 Pago de la asignación por jornada \
de trabajo adicional y asignación por cargo de mayor responsabilidad, \
la cuál es usada para financiar las encargaturas. Asimismo, el Pliego Regional \
ya contaba con una programación de ')
    encarg_parrafo2.add_run(f'{prog_gore}') #Insertar valor de base de datos
    encarg_parrafo2.add_run(' en la misma finalidad y mediante ')
    encarg_parrafo2.add_run(' Oficio Múltiple N° 00082-2021-MINEDU/SPE-OPEP-UPP, ')
    encarg_parrafo2.add_run('se le solicitó a las Unidades Ejecutoras del Pliego Regional \
realizar modificaciones presupuestarias por el monto de S/.')
    encarg_parrafo2.add_run(f'{nm_enca}') #Insertar valor de base de datos
    encarg_parrafo2.add_run(' para habilitar la finalidad 0267929')
    encarg_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo3 = document.add_paragraph('Con ')
    encarg_parrafo3.add_run('Decreto Supremo 217-2021 \
publicado el 27 de agosto de 2021 en el marco de lo autorizado en el literal b) \
del numeral 40.1 de la Ley de Presupuesto 2021, se ha realizado una transferencia \
de partidas por el monto de S/.')  # Este párrafo tendrá que variar año tras año
    encarg_parrafo3.add_run(f'{ds_217}') #Insertar valor de base de datos    
    encarg_parrafo3.add_run(' a favor de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo3.add_run(region)
    encarg_parrafo3.add_run(' para financiar el costo diferencial.')
    encarg_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo4 = document.add_paragraph('En el siguiente cuadro se muestran el costo \
y los montos programados/transferidos a la Región ')
    encarg_parrafo4.add_run(region)
    encarg_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla_enc = document.add_table(tabla_encargaturas_resumen.shape[0]+1, tabla_encargaturas_resumen.shape[1])
    tabla_enc.style = "Colorful List Accent 1"
    for j in range(tabla_encargaturas_resumen.shape[-1]):
        tabla_enc.cell(0,j).text = tabla_encargaturas_resumen.columns[j]
    for i in range(tabla_encargaturas_resumen.shape[0]):
        for j in range(tabla_encargaturas_resumen.shape[-1]):
            tabla_enc.cell(i+1,j).text = str(tabla_encargaturas_resumen.values[i,j])

#-----------------------------------------------------------------------------#

    ###################################
    ##   2) Asginaciones Temporales  ##
    ###################################

    document.add_heading("2.Pago de Asignaciones Temporales", level=1)
    
    ##Párrafos
    at_parrafo1 = document.add_paragraph("Para la Región ")    
    at_parrafo1.add_run(region)
    at_parrafo1 = document.add_paragraph(', por concepto de Asignaciones Temporales \
por prestar servicios en condiciones especiales, se ha calculado para el 2021 un costo de \
S/ ')    
    at_parrafo1.add_run(f'{costo_at} ') #Insertar valor de base de datos    
    at_parrafo1 = document.add_paragraph('que incluye el pago por prestar servicios \
en zonas rurales, de frontera, VRAEM, Instituciones Educativas Unidocentes, Multigrado \
Bilingüe y acreditar dominio de lengua originaria, de los profesores y auxiliares de \
educación nombrados y contratados.')        
    at_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    at_parrafo2 = document.add_paragraph('Para financiar estos conceptos, en el ')
    at_parrafo2.add_run('2020, ') #Calcular valor de año anterior
    at_parrafo2.add_run('el MINEDU gestionó una programación directa de recursos \
en el PIA ')
    at_parrafo2.add_run(datetime.today().strftime('%Y')) #Año actual
    at_parrafo2.add_run(' de las Unidades Ejecutoras de Educación de la Región ')    
    at_parrafo2.add_run(region)
    at_parrafo2.add_run('por el monto de S/. ')
    at_parrafo2.add_run(f'{apm_at}')
    at_parrafo2.add_run(' en la finalidad 0267928. Pago de las asignaciones \
por tipo y ubicacion de Institucion Educativa la \
cuál es usada para financiar las asignaciones temporales. \
Asimismo, el Pliego Regional ya contaba con una programación de S/ ')
    at_parrafo2.add_run(f'{diploma_gore_at}') 
    at_parrafo2.add_run(' en la misma finalidad.')
    
    at_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
  
    at_parrafo3 = document.add_paragraph('Con Decreto Supremo 187-2021 \
publicado el 22 de julio de 2021 en el marco de lo autorizado en los literales \
a), c), d) y e) del numeral 40.1 de la Ley de Presupuesto 2021, ') # Este párrafo tendrá que variar año tras año
    at_parrafo3.add_run('se ha realizado una transferencia \
de partidas por el monto de S/. ')
    at_parrafo3.add_run(f'{ds_187_at}') #Insertar valor de base de datos    
    at_parrafo3.add_run(' a favor de las Unidades Ejecutoras de Educación \
de la Región ') 
    at_parrafo3.add_run(region)
    at_parrafo3.add_run(' para financiar el costo diferencial de las asignaciones \
temporales a favor los profesores y auxiliares de educación nombrados y contratados.')    
    at_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    at_parrafo4 = document.add_paragraph('Actualmente está en gestión en el MINEDU \
la segunda transferencia de recursos por concepto de asignaciones temporales, \
el cual debería realizarse antes del ')
    at_parrafo4.add_run('26 de noviembre del 2021') #Esta fecha se actualizará año a año 
    at_parrafo4.add_run(' de acuerdo al plazo legal establecido en la Ley de Presupuesto') 
    at_parrafo4.add_run(datetime.today().strftime('%Y')) #Año actual
    at_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    at_parrafo5 = document.add_paragraph('En el siguiente cuadro se muestran el costo y los \
montos programados/transferidos a la Región ')
    at_parrafo5.add_run(region)
    at_parrafo5.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY    

    tabla_at = document.add_table(tabla_at_resumen.shape[0]+1, tabla_at_resumen.shape[1])
    tabla_at.style = "Colorful List Accent 1"
    for j in range(tabla_at_resumen.shape[-1]):
        tabla_at.cell(0,j).text = tabla_at_resumen.columns[j]
    for i in range(tabla_at_resumen.shape[0]):
        for j in range(tabla_at_resumen.shape[-1]):
            tabla_at.cell(i+1,j).text = str(tabla_at_resumen.values[i,j])

#-----------------------------------------------------------------------------#

    ###################################
    ##     3) Beneficios sociales    ##
    ###################################
    
    document.add_heading("3.Pago de Beneficios Sociales", level=1)

    ##Párrafos    
    bs_parrafo1 = document.add_paragraph('Para la Región ')
    bs_parrafo1.add_run(region)
    bs_parrafo1.add_run(', de acuerdo con la nueva estrategia para el pago oportuno \
de Beneficios Sociales implementada por el MINEDU, se han aprobado pagos por concepto de \
Asignación por Tiempo de Servicios (ATS), Compensación por Tiempo de Servicios (CTS) y, \
Subsidio por Luto y Sepelio (SLS) hasta por  un costo de S/ ')
    bs_parrafo1.add_run(f'{costo_bs}') #Insertar valor de base de datos    
    bs_parrafo1.add_run(' a la fecha , a favor de los profesores y auxiliares de educación \
nombrados y contratados.')    
    bs_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    bs_parrafo2 = document.add_paragraph('Para financiar estos conceptos, el ')
    bs_parrafo2.add_run('2020 ') #Calcular valor de año anterior
    bs_parrafo2.add_run('el MINEDU gestionó una programación directa de recursos en el PIA ')
    bs_parrafo2.add_run(datetime.today().strftime('%Y')) #Año actual
    bs_parrafo2.add_run(' de las Unidades Ejecutoras de Educación de la Región ')
    bs_parrafo2.add_run(region)
    bs_parrafo2.add_run(' por el monto de S/. ')
    bs_parrafo2.add_run(f'{apm_bs}.') #Insertar valor de base de datos    
    bs_parrafo2.add_run('  Lo cual fue comunicado a través del ')
    bs_parrafo2.add_run('Oficio Múltiple N° 00011-2021-MINEDU/SPE-OPEP-UPP') #Esto cambiará cada año
    bs_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    bs_parrafo3 = document.add_paragraph('Con ')
    bs_parrafo3.add_run('Decreto Supremo 072-2021 publicado el 21 de abril de 2021')
    bs_parrafo3.add_run('en el marco de lo autorizado en los literales a), d) y e) \
del numeral 40.1 de la Ley de Presupuesto ')
    bs_parrafo3.add_run(datetime.today().strftime('%Y')) #Año actual
    bs_parrafo3.add_run(', se ha realizado una transferencia de partidas por el monto de S/ ')   
    bs_parrafo3.add_run(f'{ds_72_bs}.') #Insertar valor de base de datos   
    bs_parrafo3.add_run(', a favor de las Unidades Ejecutoras de Educación de la Región ')   
    bs_parrafo3.add_run(region)
    bs_parrafo3.add_run(' para financiar el pago de los beneficios sociales a favor los profesores \
y auxiliares de educación nombrados y contratados que fueron reconocidos hasta el ')   
    bs_parrafo3.add_run('2020') #Calcular valor de año anterior
    bs_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    bs_parrafo4 = document.add_paragraph('Asimismo, mediante ')
    bs_parrafo4.add_run('Decreto Supremo 256-2021 publicado el 24 de setiembre de 2021, ')  #Esto cambiará cada año
    bs_parrafo4.add_run(', se realizó la segunda transferencia de recursos por concepto de \
beneficios sociales a favor de docentes y auxiliares nombrados y contratados, cuyos beneficios fueron \
reconocidos durante el año ')
    bs_parrafo4.add_run(datetime.today().strftime('%Y')) #Año actual
    bs_parrafo4.add_run(' transfiriéndose S/. ')
    bs_parrafo4.add_run(f'{ds_256_bs}.') #Insertar valor de base de datos   
    bs_parrafo4.add_run('  a las Unidades Ejecutoras de Educación de la Región ')
    bs_parrafo4.add_run(region)
    bs_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
   
    bs_parrafo5 = document.add_paragraph('En el siguiente cuadro se muestran el costo y los montos \
programados/transferidos a la Región ')
    bs_parrafo5.add_run(region)
    bs_parrafo5.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla_bs = document.add_table(tabla_bs_resumen.shape[0]+1, tabla_bs_resumen.shape[1])
    tabla_bs.style = "Colorful List Accent 1"
    for j in range(tabla_bs_resumen.shape[-1]):
        tabla_bs.cell(0,j).text = tabla_bs_resumen.columns[j]
    for i in range(tabla_bs_resumen.shape[0]):
        for j in range(tabla_bs_resumen.shape[-1]):
            tabla_bs.cell(i+1,j).text = str(tabla_bs_resumen.values[i,j])

    bs_parrafo6 = document.add_paragraph(' De la misma forma, durante el presente año, para la Región ')  
    bs_parrafo6.add_run(region)
    bs_parrafo6.add_run(' se ha realizado las siguientes transferencias')
    bs_parrafo6.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    bs_parrafo7 = document.add_paragraph(' Por otro lado, para el año ')
    bs_parrafo7.add_run('2022 ') #Calcular año posterior
    bs_parrafo7.add_run('el MINEDU está gestionando la programación parcial de recursos en los \
presupuestos de las Unidades Ejecutoras para atender encargaturas, asignaciones temporales, \
beneficios sociales, entre otros y, el financiamiento restante, se realizará de manera oportuna el ')
    bs_parrafo7.add_run('2022, ') #Calcular año posterior
    bs_parrafo7.add_run('preferentemente antes que termine el primer semestre de dicho año fiscal.')    
    bs_parrafo7.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY    
            
    document.save(here() / f'output/{region}_AM_GG1.docx')