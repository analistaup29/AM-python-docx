#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Sep 21 15:54:48 2021
"""
import docx
import pandas as pd
import numpy as np
import nums_from_string
import os
import getpass
import glob
import platform
from datetime import datetime
#from pyprojroot import here
from janitor import clean_names # pip install pyjanitor
from pathlib import Path
from docx.shared import Pt

###############################################################################
# Rutas de los archivos #
###############################################################################

if getpass.getuser() == "analistaup18": # PC Analista UP 18 Minedu
    github = Path("C:/Users/ANALISTAUP18/Documents/GitHub/AM-python-docx")
    proyecto = Path("B:/OneDrive - Ministerio de Educación/unidad_B/2021/4. Herramientas de Seguimiento/13.AM_automatizada")
elif  getpass.getuser() == "bran": # PC Brandon
    github = Path("/Users/bran/Documents/GitHub/AM-python-docx")
    proyecto = Path("/Users/bran/Documents/GitHub/AM-python-docx")

###############################################################################
# Fechas de corte #
###############################################################################

# Importamos los nombres de los archivos en la carpeta input
lista_archivos = os.listdir(Path(proyecto, "input", "intervenciones_pedagogicas"))

# Fecha de hoy
fecha_actual = datetime.today().strftime('%d-%m-%y')

## A) Base disponibilidad
# Importamos los nombres de los archivos en la carpeta intervenciones pedagogicas
lista_archivos_int = glob.glob(os.path.join(proyecto,"input/intervenciones_pedagogicas/*"))
#Mantenemos el corte de disponibilidad más reciente
fecha_corte_disponibilidad = max(lista_archivos_int, key=os.path.getctime)
# Nos quedamos con el nombre de archivo para la base de disponibilidad
fecha_corte_disponibilidad = os.path.split(fecha_corte_disponibilidad)
fecha_corte_disponibilidad = fecha_corte_disponibilidad[1]
# Extraemos la fecha del nombre de archivo
fecha_corte_disponibilidad = nums_from_string.get_numeric_string_tokens(fecha_corte_disponibilidad)
# Convertimos a formato string
fecha_corte_disponibilidad = ''.join(fecha_corte_disponibilidad) 
# Convertimos a formato numérico
fecha_corte_disponibilidad_date = datetime.strptime(fecha_corte_disponibilidad, '%Y%m%d').date()
mes_disponibilidad = fecha_corte_disponibilidad_date.month
# Damos estilo
fecha_corte_disponibilidad_date = fecha_corte_disponibilidad_date.strftime("%d %b %Y")

## B) Siaf de mascarillas
fecha_corte_mascarillas = "03 Oct 2021"

# C) Compromisos de desempeño
fecha_corte_compromisos = "21 Sep 2021"

# Creamos tabla con fechas de corte
tabla_fechas_corte = (
    ("Intervenciones pedagógicas", fecha_corte_disponibilidad_date),
    ("Mascarillas y protectores faciales", fecha_corte_mascarillas),
    ("Compromisos de desempeño", fecha_corte_compromisos)
)


###############################################################################
# Creación de carpeta donde se guardan los outputs #
###############################################################################

# Creación de carpeta
dir = os.path.join(proyecto, f"output/AM_{fecha_actual}")
if not os.path.exists(dir):
    os.mkdir(dir)
    print("Se creó una nueva carpeta")
else:
    print("Ya existe la carpeta")
        
# Path de nueva carpeta
nueva_carpeta = Path(proyecto/ f"output/AM_{fecha_actual}")

###############################################################################
# Transformación de Datasets #
###############################################################################

##########################
# Base de disponibilidad #
##########################

# Base de datos región
## Cargamos nombres de regiones
nombre_regiones = pd.read_excel(proyecto / "input/otros/nombre_regiones.xlsx")

# A) Base de disponibilidad
## Cargamos base de disponibilidad
data_intervenciones = pd.read_excel(proyecto / f"input/intervenciones_pedagogicas/Disponibilidad_Presupuestal_{fecha_corte_disponibilidad}interv.xlsx")
data_intervenciones = clean_names(data_intervenciones) # Normalizamos nombres

# Eliminamos filas de "No hay Intervenciones pedagogicas"
data_intervenciones = data_intervenciones[data_intervenciones['intervencion_pedagogica'] != "No hay Intervenciones Pedagógicas"]
# Eliminamos COAR
data_intervenciones = data_intervenciones[data_intervenciones['intervencion_pedagogica'] != "COAR"]
# Eliminamos  Vacaciones Truncas 
data_intervenciones = data_intervenciones[data_intervenciones['especifica_de_gasto'] != "3.2.8.1.5. VACACIONES TRUNCAS DE C.A.S."]

# Mantenemos variables de interés (PIM, DEVENGADO, COMPROMETIDO CERTIFICADO) y 
# colapsamos a nivel de Region, Intervencion Pedagogica y Cas-No-Cas
data_intervenciones_total = data_intervenciones[["region", "cas_no_cas","intervencion_pedagogica", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}", "costo_enero", "costo_febrero", "costo_marzo", "costo_abril", "costo_mayo", "costo_junio", "costo_julio", "costo_agosto", "costo_septiembre", "costo_octubre", "costo_noviembre", "costo_diciembre"]]. \
   groupby(by = ["region", "intervencion_pedagogica"] , as_index=False).sum()

data_intervenciones = data_intervenciones[["region", "cas_no_cas","intervencion_pedagogica", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}", "costo_enero", "costo_febrero", "costo_marzo", "costo_abril", "costo_mayo", "costo_junio", "costo_julio", "costo_agosto", "costo_septiembre", "costo_octubre", "costo_noviembre", "costo_diciembre"]]. \
   groupby(by = ["region","cas_no_cas", "intervencion_pedagogica"] , as_index=False).sum()

# Eliminamos filas con 0 PIM
condicion_elim = (data_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"] != 0) 
data_intervenciones = data_intervenciones[condicion_elim]

# Calculamos porcentajes
## Avance PIM
data_intervenciones["avance_pim"] = data_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"]/data_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"]

## Avance costo actual
data_intervenciones["costo_actual"] = data_intervenciones["costo_enero"] + data_intervenciones["costo_febrero"] + data_intervenciones["costo_marzo"] + data_intervenciones["costo_abril"] + data_intervenciones["costo_mayo"] + data_intervenciones["costo_junio"] + data_intervenciones["costo_julio"] + data_intervenciones["costo_agosto"] + data_intervenciones["costo_septiembre"] + data_intervenciones["costo_octubre"]
data_intervenciones["avance_costo"] = data_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"]/data_intervenciones["costo_actual"]

# Mantenemos y ordenamos columnas
data_intervenciones = data_intervenciones[['region', "intervencion_pedagogica", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}", "avance_pim", "costo_actual", "avance_costo"]]

## Reemplazamos inf por NaN
data_intervenciones.replace([np.inf, -np.inf], np.nan, inplace=True)

## Reemplazamos NaN por 0
data_intervenciones['avance_pim'] = data_intervenciones['avance_pim'].fillna("0").astype(float)
data_intervenciones['avance_costo'] = data_intervenciones['avance_costo'].fillna("0").astype(float)

# Tabla intervenciones CAS
#data_intervenciones_cas = data_intervenciones[data_intervenciones['cas_no_cas'] != "NO CAS"]
# Tabla intervenciones NO CAS
#data_intervenciones_nocas = data_intervenciones[data_intervenciones['cas_no_cas'] != "CAS"]

#######################
# Siaf de mascarillas #
#######################

## Cargamos la base insumo de mascarillas
data_mascarillas = pd.read_excel(proyecto / "input/mascarillas/Incorporación_DU_SIAF_20211003.xlsx", sheet_name='Sheet1')
data_mascarillas = clean_names(data_mascarillas) # Normalizamos nombres

# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
data_mascarillas = data_mascarillas[["region","nom_ue","transferencia","pim","certificado","comprometido_anual","devengado"]]. \
   groupby(by = ["region", "nom_ue"], as_index=False).sum()

data_mascarillas["region"] = data_mascarillas["region"].str.split(". ", n=1).apply(lambda l: "".join(l[1]))

data_mascarillas["UNIDAD EJECUTORA"]=data_mascarillas["nom_ue"]
data_mascarillas["RECURSOS TRANSF. (*)"]=data_mascarillas["transferencia"]
data_mascarillas["PIM"]=data_mascarillas["pim"]
data_mascarillas["CERT. (%)"]=data_mascarillas["certificado"]/data_mascarillas["pim"]
data_mascarillas["COMPRO. (%)"]=data_mascarillas["comprometido_anual"]/data_mascarillas["pim"]
data_mascarillas["DEVENGADO (%)"]=data_mascarillas["devengado"]/data_mascarillas["pim"]

############################
# Compromisos de desempeño #
############################

## Cargamos data de compromisos de desempeño
data_cdd = pd.read_excel(proyecto / "input/compromisos_desempeno/regiones_BD_CDD.xlsx")
data_cdd = clean_names(data_cdd) # Normalizamos nombres
data_cdd["pliego"] = data_cdd["pliego"].str.split(". ", n=1, expand = True)
data_cdd['pliego'] = data_cdd['pliego'].astype('int64') # Convertimos ubigeo a integer

# Corregimos genericas
data_cdd["generica"] = data_cdd["generica"].replace("3. BIENES Y SERVICIOS", "2.3. BIENES Y SERVICIOS")
data_cdd["generica"] = data_cdd["generica"].replace("6. ADQUISICION DE ACTIVOS NO FINANCIEROS", "2.6. ADQUISICION DE ACTIVOS NO FINANCIEROS")	

# Hacemos merge con base de datos región
data_cdd = data_cdd.merge(right = nombre_regiones, how="left", on = "pliego")

# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
data_cdd = data_cdd[["region", "unidad_ejecutora", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region", "programa_presupuestal", "generica"], as_index=False).sum()

######################################################
# Sobre el financiamiento de conceptos remunerativos #
######################################################

# C) Base de Encargaturas
df_consolidado_enc = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'ENC-CONSOLIDADO-VF') 
df_consolidado_enc.fillna(0, inplace =  True)
df_consolidado_enc['COSTO'] = df_consolidado_enc['COSTO-TRAMO I']
df_consolidado_enc['PROGRAMADO POR MINEDU'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']+ df_consolidado_enc['NM ENCARGATURAS']
df_consolidado_enc.rename(columns={'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                   'TRANSFERENCIA DS 217':'TRANSFERENCIA POR DS N° 217-2021-EF',
                                   'UNIDADEJECUTORA':'UNIDAD EJECUTORA'},inplace=True)
df_consolidado_enc['APMeINCREMENTOS'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']

tabla_encargaturas = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIA POR DS N° 217-2021-EF',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]

tabla_encargaturas_resumen = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIA POR DS N° 217-2021-EF',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]



# D) Base de Asignaciones Temporales
df_consolidado_at = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'AT-CONSOLIDADO-VF')   
df_consolidado_at.fillna(0, inplace =  True)      
df_consolidado_at.rename(columns={'REGIÓN':'REGION',
                                  'COSTO-TRAMO I':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                  'TRANSFERENCIA DS 187':'TRANSFERENCIA POR DS N° 187-2021-EF'
    },inplace=True)
df_at=df_consolidado_at[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU','TRANSFERENCIA POR DS N° 187-2021-EF']]

# E) Base de Beneficios Sociales
df_consolidado_bf = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'BS-CONSOLIDADO-VF')           
df_consolidado_bf.fillna(0,inplace = True)
df_consolidado_bf['COSTO BENEFICIARIOS 2020 Y 2021'] = df_consolidado_bf['LISTAS-2021'] + df_consolidado_bf['TRAMO I-BS 2020'] + df_consolidado_bf['TRAMO II-BS 2021']
df_consolidado_bf.rename(columns={'REGIÓN':'REGION',
                                  'COSTO BENEFICIARIOS 2020 Y 2021':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'TRANSFERENCIA DS 072-BS 2020':'TRANSFERENCIA POR DS N° 072-2021-EF',
                                  'TRANSFERENCIA DS 256-BS 2021':'TRANSFERENCIA POR DS N° 256-2021-EF'
    },inplace=True)
df_bs = df_consolidado_bf[['REGION','UNIDAD EJECUTORA','COSTO',
                           'PROGRAMADO POR MINEDU',
                           'TRANSFERENCIA POR DS N° 072-2021-EF',
                           'TRANSFERENCIA POR DS N° 256-2021-EF']]

df_transferencia = pd.read_excel(proyecto / 'input/normas_transferencias/TRANSFERENCIAS 2021.xlsx',sheet_name = 'TRANSFERENCIAS')           
df_transferencia.fillna(0,inplace = True)
df_transferencia = clean_names(df_transferencia)
df_transferencia = df_transferencia[['region', 'norma_de_transferencia', 'concepto', 'monto_transferido']].\
groupby(by = ["region", 'norma_de_transferencia', 'concepto'] , as_index=False).sum()


normas = ["DECRETO DE URGENCIA N 065-2021", "DECRETO SUPREMO N 044-2021-EF", "DECRETO SUPREMO N 078-2021-EF"]


df_transferencia = df_transferencia.loc[df_transferencia['norma_de_transferencia'].isin(normas)]
df_transferencia["concepto"].replace({"CONTRATACIÓN MINDEF": "Contratación de plazas docentes en instituciones educativas de educación básica del Ministerio de Defensa"}, inplace=True)


######################################################
# Sobre el proceso de racionalización #
######################################################

data_creacion = pd.read_excel(proyecto / "input/creaciones/plazas_creacion_racio_2021.xlsx", sheet_name="BD",  skiprows = 2)
data_creacion = clean_names(data_creacion) # Normalizamos nombres
data_creacion = data_creacion[['d_region', 'd_dreugel', 'nivel', 'creacion_total']].\
groupby(by = ["d_region", 'd_dreugel', 'nivel'] , as_index=False).sum()
data_creacion['d_region'] = data_creacion['d_region'].str.split(r'DRE ').str[-1]
data_creacion.loc[data_creacion['d_region']=="DE DIOS", 'd_region'] = "MADRE DE DIOS"


data_creacion = data_creacion.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})
data_creacion.loc[data_creacion['nivel']=="inicial", 'inicial'] = data_creacion['creacion_total']
data_creacion.loc[data_creacion['nivel']=="primaria", 'primaria'] = data_creacion['creacion_total']
data_creacion.loc[data_creacion['nivel']=="secundaria", 'secundaria'] = data_creacion['creacion_total']

data_creacion = data_creacion[['region', 'ugel','inicial', 'primaria', 'secundaria', 'creacion_total']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()

data_creacion.fillna(0, inplace =  True)

## Creación plazas docentes -PEM 2021

data_creacion_pem = pd.read_excel(proyecto / "input/creaciones/plazas_creacion_pem_2021.xlsx")
data_creacion_pem = clean_names(data_creacion_pem) # Normalizamos nombres
data_creacion_pem = data_creacion_pem[['d_region', 'd_dreugel', 'modalidad', 'req_doc', 'req_bolsa', 'req_director', 'req_subdir']].\
groupby(by = ["d_region", 'd_dreugel', 'modalidad'] , as_index=False).sum()
data_creacion_pem['d_region'] = data_creacion_pem['d_region'].str.split(r'DRE ').str[-1]
data_creacion_pem.loc[data_creacion_pem['d_region']=="DE DIOS", 'd_region'] = "MADRE DE DIOS"

data_creacion_pem = data_creacion_pem.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})
data_creacion_pem.fillna(0, inplace =  True)

filtro_ebr = data_creacion_pem['modalidad']=="EBR"
creacion_ebr_pem = data_creacion_pem[filtro_ebr]
filtro_ebe = data_creacion_pem['modalidad']=="EBE"
creacion_ebe_pem = data_creacion_pem[filtro_ebe]

#----------------------------------------------------------------------#
## Brecha de plazas docentes
data_brecha = pd.read_excel(proyecto / "input/brecha/brecha_ugel_2020.xlsx", sheet_name="Data")
# Normalizamos nombres


# Mantenemos variables de interés
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()
data_brecha['doc_e_n_cub_req'] = data_brecha['doc_e_n'] - data_brecha['nom_exd_mov1']
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']]
data_brecha.loc[data_brecha['brecha_net']<=0, 'req_neto'] = -1*data_brecha['brecha_net']
data_brecha.loc[data_brecha['brecha_net']>0, 'exc_neto'] = data_brecha['brecha_net']

#Cantidad de UGEL con requerimiento neto
data_brecha.loc[data_brecha['brecha_net']<0, 'cant_ugel_req'] = 1
data_brecha.loc[data_brecha['brecha_net']>0, 'cant_ugel_exc'] = 1

data_brecha_regional = data_brecha[['region', 'brecha_net', 'cant_ugel_req', 'cant_ugel_exc']].groupby(by = ["region"] , as_index=False).sum()
data_brecha_regional.loc[data_brecha_regional['brecha_net']<=0, 'brecha_net'] = -1*data_brecha_regional['brecha_net']
data_brecha_regional.loc[data_brecha_regional['brecha_net']>0, 'brecha_net'] = data_brecha_regional['brecha_net']
data_brecha_regional.fillna(0, inplace =  True)


data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'req_neto', 'exc_neto', 'brecha_net']]
data_brecha.fillna(0, inplace =  True)


#----------------------------------------------------------------------#
## Bloqueo de plazas
data_bloqueo = pd.read_excel(proyecto / "input/bloqueo/plazas_bloqueo_2020.xlsx")
data_bloqueo = clean_names(data_bloqueo) # Normalizamos nombres
data_bloqueo.fillna(0, inplace =  True)

# Mantenemos variables de interés
data_bloqueo['cant_bloqueos'] = 1
data_bloqueo = data_bloqueo[["descreg", 'cant_bloqueos']].groupby(by = ["descreg"] , as_index=False).sum()
data_bloqueo = data_bloqueo.rename(columns={'descreg':'region'})

#----------------------------------------------------------------------#
## Deuda social
data_deuda_social = pd.read_excel(proyecto / "input/deudas_sociales/deudas_sociales.xlsx")
data_deuda_social = clean_names(data_deuda_social) # Normalizamos nombres

###############################################################################
# Creación del documento en docx 
###############################################################################

# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS", "ANCASH", "APURIMAC", "AREQUIPA", "AYACUCHO", "CAJAMARCA", "CUSCO", "HUANCAVELICA", "HUANUCO", "ICA", "JUNIN", "LA LIBERTAD", "LAMBAYEQUE", "LORETO", "MADRE DE DIOS", "MOQUEGUA", "PASCO", "PIURA", "PUNO", "SAN MARTIN", "TACNA", "TUMBES", "UCAYALI", "LIMA PROVINCIAS", "CALLAO"]

# For loop para cada región
for region in lista_regiones:

    ###########################################################################
    # Construcción de tablas e indicadores #
    ###########################################################################
    
    ############################################
    # Tablas e indicadores base disponibilidad #
    ############################################

    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_intervenciones['region'] == region #Seleccionar region
    tabla_intervenciones = data_intervenciones[region_seleccionada]    
    pim_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    ejecucion_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    porcentaje_ejecucion = str('{:,.1%}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    porcentaje_costomesactual = str('{:,.1%}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones["costo_actual"].sum()))
    
    # TOTAL
    
    tabla_intervenciones_formato = data_intervenciones[region_seleccionada]
    
    # Generamos porcentaje de avance
    porcentaje_ejecucion_a = tabla_intervenciones_formato[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()
    porcentaje_costomesactual_a = tabla_intervenciones_formato[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato["costo_actual"].sum()
    
    # Generamos fila total
    total_int = tabla_intervenciones_formato.groupby(by = ["region"], as_index=False).sum()

    # Realizamos append del total en la tabla
    tabla_intervenciones_formato = tabla_intervenciones_formato.append(total_int, ignore_index=True)

    # Reemplazamos % de avance pim y avance costo por los valores correctos en fila total
    tabla_intervenciones_formato.iloc[-1, tabla_intervenciones_formato.columns.get_loc('avance_pim')] = porcentaje_ejecucion_a
    tabla_intervenciones_formato.iloc[-1, tabla_intervenciones_formato.columns.get_loc('avance_costo')] = porcentaje_costomesactual_a

    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_intervenciones_formato['intervencion_pedagogica'] = tabla_intervenciones_formato['intervencion_pedagogica'].fillna("Total")
    tabla_intervenciones_formato['avance_pim'] = tabla_intervenciones_formato['avance_pim'].fillna("0").astype(float)
    tabla_intervenciones_formato['avance_costo'] = tabla_intervenciones_formato['avance_costo'].fillna("0").astype(float)
    

    # CAS
    #tabla_intervenciones_formato_cas = data_intervenciones_cas[region_seleccionada]
    
    # Generamos porcentaje de avance CAS
    #porcentaje_ejecucion_cas = tabla_intervenciones_formato_cas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_cas[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()
    #porcentaje_costomesactual_cas = tabla_intervenciones_formato_cas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_cas["costo_actual"].sum()
    
    # Generamos fila total
    #total_int_cas = tabla_intervenciones_formato_cas.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_cas = tabla_intervenciones_formato_cas.append(total_int_cas, ignore_index=True)
    
    # Reemplazamos % de avance pim y avance costo por los valores correctos en fila total
    #tabla_intervenciones_formato_cas.iloc[-1, tabla_intervenciones_formato_cas.columns.get_loc('avance_pim')] = porcentaje_ejecucion_cas
    #tabla_intervenciones_formato_cas.iloc[-1, tabla_intervenciones_formato_cas.columns.get_loc('avance_costo')] = porcentaje_costomesactual_cas
    
    #Incluimos palabra "total" y "-" en vez de NaN
    #tabla_intervenciones_formato_cas['intervencion_pedagogica'] = tabla_intervenciones_formato_cas['intervencion_pedagogica'].fillna("Total")
    #tabla_intervenciones_formato_cas['avance_pim'] = tabla_intervenciones_formato_cas['avance_pim'].fillna("0").astype(float)
    #tabla_intervenciones_formato_cas['avance_costo'] = tabla_intervenciones_formato_cas['avance_costo'].fillna("0").astype(float)
    
    # NO CAS
    #tabla_intervenciones_formato_nocas = data_intervenciones_nocas[region_seleccionada]
    
    # Generamos porcentaje de avance NO CAS
    #porcentaje_ejecucion_nocas = tabla_intervenciones_formato_nocas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_nocas[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()
    #porcentaje_costomesactual_nocas = tabla_intervenciones_formato_nocas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_nocas["costo_actual"].sum()
    
    # Generamos fila total
    #total_int_nocas = tabla_intervenciones_formato_nocas.groupby(by = ["region"], as_index=False).sum()
    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_nocas = tabla_intervenciones_formato_nocas.append(total_int_nocas, ignore_index=True)
    
    # Reemplazamos % de avance pim y avance costo por los valores correctos en fila total
    #tabla_intervenciones_formato_nocas.iloc[-1, tabla_intervenciones_formato_nocas.columns.get_loc('avance_pim')] = porcentaje_ejecucion_nocas
    #tabla_intervenciones_formato_nocas.iloc[-1, tabla_intervenciones_formato_nocas.columns.get_loc('avance_costo')] = porcentaje_costomesactual_nocas
    
    #Incluimos palabra "total" y "-" en vez de NaN
    #tabla_intervenciones_formato_nocas['intervencion_pedagogica'] = tabla_intervenciones_formato_nocas['intervencion_pedagogica'].fillna("Total")
    #tabla_intervenciones_formato_nocas['avance_pim'] = tabla_intervenciones_formato_nocas['avance_pim'].fillna("0").astype(float)
    #tabla_intervenciones_formato_nocas['avance_costo'] = tabla_intervenciones_formato_nocas['avance_costo'].fillna("0").astype(float)
    
    # Formato para la tabla
    formato_tabla_intervenciones = {
        "intervencion_pedagogica" : "{}",
        f"pim_reporte_siaf_{fecha_corte_disponibilidad}": "{:,.0f}",
        f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}" : "{:,.0f}",
        f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}" : "{:,.0f}",
    #    "avance_pim": "{:,.1%}",
    #    "costo_actual": "{:,.0f}",
        "avance_costo": "{:,.1%}",
        }
    #tabla_intervenciones_formato_cas = tabla_intervenciones_formato_cas.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})            
    #tabla_intervenciones_formato_nocas = tabla_intervenciones_formato_nocas.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})            

    tabla_intervenciones_formato = tabla_intervenciones_formato.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})            
    
    ############################################
    # Tablas e indicadores mascarillas #
    ############################################

    # Generamos la tabla "tabla1_mascarilla" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = data_mascarillas['region'] == region
    tabla_mascarillas = data_mascarillas[region_seleccionada]
    # Generamos los indicadores de PIM y ejecución de intervenciones
    transferencia_mascarilla = str('{:,.1f}'.format(tabla_mascarillas["transferencia"].sum()/1000000))
    devengado_mascarillas=str('{:.1%}'.format(tabla_mascarillas["devengado"].sum()/tabla_mascarillas["transferencia"].sum()))
    # Generamos la tabla "tabla_mascarillas_formato" - mantiene la región i de la lista de regiones
    tabla_mascarillas_formato = data_mascarillas[region_seleccionada]
    # Formato para la tabla
    formato_tabla_mascarillas = {
        "UNIDAD EJECUTORA": "{}",
        "RECURSOS TRANSF. (*)": "{:,.0f}",
        "PIM" : "{:,.0f}",
        "CERT. (%)" : "{:.1%}",
        "COMPRO. (%)": "{:.1%}",
        "DEVENGADO (%)": "{:.1%}",  
        }
    tabla_mascarillas_formato = tabla_mascarillas_formato.transform({k: v.format for k, v in formato_tabla_mascarillas.items()})  
    
    #################################################
    # Tablas e indicadores compromisos de desempeño #
    #################################################

    # Generamos la tabla "tabla_cdd" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = data_cdd['region'] == region
    tabla_cdd = data_cdd[region_seleccionada]
    
    # Generamos fila total
    total_cdd = tabla_cdd[["region", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_cdd = tabla_cdd.append(total_cdd, ignore_index=True)
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_cdd['programa_presupuestal'] = tabla_cdd['programa_presupuestal'].fillna("Total")
    tabla_cdd['generica'] = tabla_cdd['generica'].fillna("-")

     # Generamos la tabla con formato
    formato_tabla_cdd = {
        "programa_presupuestal": "{}",
        "generica": "{}",
        "monto": "{:,.0f}",
        "ds_085_2021_ef": "{:,.0f}",
        "ds_218_2021_ef": "{:,.0f}",
        "ds_220_2021_ef": "{:,.0f}",
        }
    tabla_cdd_formato = tabla_cdd
    tabla_cdd_formato = tabla_cdd_formato.transform({k: v.format for k, v in formato_tabla_cdd.items()})            
    
    # Generamos CDD transferido 
    cdd_transferido = str('{:,.0f}'.format(tabla_cdd["monto"].sum()))
   
    ########################
    # Tablas Encargaturas #
    ########################

    region_seleccionada = df_consolidado_enc['REGION'] == region
    tabla1 = tabla_encargaturas[region_seleccionada]
    costo_enc = str('{:,.0f}'.format(tabla1["COSTO"].sum())) 
    apmeincre = str('{:,.0f}'.format(tabla1["APMeINCREMENTOS"].sum()))
    prog_gore = str('{:,.0f}'.format(tabla1["PROGRAMADO POR EL PLIEGO REGIONAL"].sum()))
    nm_enca = str('{:,.0f}'.format(tabla1['NM ENCARGATURAS'].sum()))
    ds_217 = str('{:,.0f}'.format(tabla1['TRANSFERENCIA POR DS N° 217-2021-EF'].sum()))
    
    tabla_encargaturas_resumen = tabla1.groupby(['UNIDAD EJECUTORA'], as_index=False).sum()
    tabla_encargaturas_resumen = tabla_encargaturas_resumen[['UNIDAD EJECUTORA',
                                                             'COSTO', 
                                                             'PROGRAMADO POR MINEDU', 
                                                             'PROGRAMADO POR EL PLIEGO REGIONAL',
                                                             'TRANSFERENCIA POR DS N° 217-2021-EF']] 
    
    # Generamos fila total
    total_enc = tabla_encargaturas_resumen[["UNIDAD EJECUTORA", "COSTO", \
"PROGRAMADO POR MINEDU", "PROGRAMADO POR EL PLIEGO REGIONAL", "TRANSFERENCIA POR DS N° 217-2021-EF"]].sum()
    
    # Realizamos append del total en la tabla
    tabla_encargaturas_resumen = tabla_encargaturas_resumen.append(total_enc, ignore_index=True)
    
    #Incluimos palabra "total"
    tabla_encargaturas_resumen.iloc[-1, tabla_encargaturas_resumen.columns.get_loc('UNIDAD EJECUTORA')] = "Total"

    tabla_encargaturas_resumen = tabla_encargaturas_resumen.round(2)

    ###################################
    #  Tablas Asignaciones Temporales #
    ###################################
    
    region_seleccionada = df_at['REGION'] == region
    tabla2 = df_at[region_seleccionada]
    costo_at = str('{:,.0f}'.format(tabla2["COSTO"].sum()))
    apm_at = str('{:,.0f}'.format(tabla2["PROGRAMADO POR MINEDU"].sum())) 
    ds_187_at = str('{:,.0f}'.format(tabla2["TRANSFERENCIA POR DS N° 187-2021-EF"].sum())) 
    tabla_at_resumen = tabla2.groupby(['UNIDAD EJECUTORA'], as_index=False).sum() 
    
    # Generamos fila total
    total_at = tabla_at_resumen[["UNIDAD EJECUTORA", "COSTO", \
"PROGRAMADO POR MINEDU", "TRANSFERENCIA POR DS N° 187-2021-EF"]].sum()
    
    # Realizamos append del total en la tabla
    tabla_at_resumen = tabla_at_resumen.append(total_at, ignore_index=True)
    
    #Incluimos palabra "total"
    tabla_at_resumen.iloc[-1, tabla_at_resumen.columns.get_loc('UNIDAD EJECUTORA')] = "Total"

    tabla_at_resumen = tabla_at_resumen.round(2)

    ###################################
    #    Tablas Beneficios Sociales   #
    ###################################
    
    region_seleccionada = df_bs['REGION'] == region
    region_seleccionada2 = df_consolidado_bf['REGION'] == region
    tabla3 = df_bs[region_seleccionada]
    tabla3_2 = df_consolidado_bf[region_seleccionada]
    costo_bs = str('{:,.0f}'.format(tabla3["COSTO"].sum()))
    apm_bs = str('{:,.0f}'.format(tabla3["PROGRAMADO POR MINEDU"].sum()))    
    ds_72_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS N° 072-2021-EF"].sum())) 
    ds_256_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS N° 256-2021-EF"].sum()))
    tabla_bs_resumen = tabla3.groupby(['UNIDAD EJECUTORA'], as_index=False).sum()

    # Generamos fila total
    total_bs = tabla_bs_resumen[["UNIDAD EJECUTORA", "COSTO", \
"PROGRAMADO POR MINEDU", "TRANSFERENCIA POR DS N° 072-2021-EF", "TRANSFERENCIA POR DS N° 256-2021-EF"]].sum()

    # Realizamos append del total en la tabla
    tabla_bs_resumen = tabla_bs_resumen.append(total_bs, ignore_index=True)
    
    #Incluimos palabra "total"
    tabla_bs_resumen.iloc[-1, tabla_bs_resumen.columns.get_loc('UNIDAD EJECUTORA')] = "Total"

    tabla_bs_resumen = tabla_bs_resumen.round(2)
    
    lista_bf = str('{:,.0f}'.format(tabla3_2["LISTAS-2021"].sum()))
    
    ###########################################################################
    # Inclusión del texto del documento #
    ###########################################################################
    document = docx.Document(proyecto / "input/otros/Formato.docx") # Creación del documento en base al template
    title=document.add_heading('AYUDA MEMORIA', 0) #Título del documento
    run = title.add_run()
    run.add_break()
    title.add_run('REGIÓN ')
    title.add_run(region)
    run = title.add_run()
    
    #####################################################
    # Incluimos sección 1 de intervenciones pedagógicas #
    #####################################################
    
    document.add_heading("1. Intervenciones pedagógicas", level=1) # 1) Intervenciones pedagógicas    
    interv_fecha_actualizacion = document.add_paragraph()
    interv_fecha_negrita1 = interv_fecha_actualizacion.add_run("Actualizado al ")
    interv_fecha_negrita2 = interv_fecha_actualizacion.add_run(fecha_corte_disponibilidad_date)
    interv_fecha_negrita1.bold = True    
    interv_fecha_negrita2.bold = True
    
    interv_parrafo1 = document.add_paragraph(
    "Las Unidades Ejecutoras de Educación de la región " , style="List Bullet")
    interv_parrafo1.add_run(region)
    interv_parrafo1.add_run(" vienen implementando ")
    interv_parrafo1.add_run(
    " intervenciones y acciones pedagógicas en el Año 2021, en el marco de la \
    Norma Técnica “Disposiciones para la implementación de las intervenciones \
    y acciones pedagógicas del Ministerio de Educación en los Gobiernos Regionales \
    y Lima Metropolitana en el Año Fiscal 2021”, aprobada mediante \
    RM N° 043-2021-MINEDU y modificada RM N° 159-2021-MINEDU." )

    interv_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    interv_parrafo2 = document.add_paragraph(
    "Las Unidades Ejecutoras de Educación de la región " , style="List Bullet")
    interv_parrafo2.add_run(region)
    interv_parrafo2.add_run(" cuentan con ")
    interv_parrafo2.add_run(pim_intervenciones_region)
    interv_parrafo2.add_run(
    " millones en su Presupuesto Institucional Modificado (PIM) para el \
    financiamiento de intervenciones y acciones pedagógicas, de los cuales se han ejecutado S/ ")
    interv_parrafo2.add_run(ejecucion_intervenciones_region)
    interv_parrafo2.add_run(" lo cual corresponde a ")
    interv_parrafo2.add_run(porcentaje_ejecucion)
    interv_parrafo2.add_run(".")

    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    # Incluimos tabla 1 intervenciones CAS    
    int_titulo = document.add_paragraph()
    int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    int_titulo_negrita = int_titulo.add_run("Intervenciones Pedagógicas (soles)")
    int_titulo_negrita.bold = True    
    tabla1_interv = document.add_table(tabla_intervenciones_formato.shape[0]+1, tabla_intervenciones_formato.shape[1])
    tabla1_interv.autofit = False
    tabla1_interv.allow_autofit = True
    tabla1_interv.style = "formato_tabla_minedu"

    
    # Incluimos tabla 1 intervenciones CAS
    #cas_titulo = document.add_paragraph()
    #cas_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    #cas_titulo_negrita = cas_titulo.add_run("Intervenciones Pedagógicas - Componente CAS (soles)")
    #cas_titulo_negrita.bold = True    
    #tabla1_interv = document.add_table(tabla_intervenciones_formato_cas.shape[0]+1, tabla_intervenciones_formato_cas.shape[1])
    #tabla1_interv.autofit = False
    #tabla1_interv.allow_autofit = True
    #tabla1_interv.style = "formato_tabla_minedu"
    #nocas_titulo = document.add_paragraph()
    #nocas_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    #nocas_tit2 = nocas_titulo.add_run("Intervenciones Pedagógicas - Componente NO CAS (soles)")
    #nocas_tit2.bold = True    
    #tabla2_interv = document.add_table(tabla_intervenciones_formato_nocas.shape[0]+1, tabla_intervenciones_formato_nocas.shape[1])
    #tabla2_interv.allow_autofit
    #tabla2_interv.style = "formato_tabla_minedu"
    
    ## Header de la tabla
    row = tabla1_interv.rows[0].cells
    row[0].text = "INTERVENCIONES"
    row[1].text = "PIM"
    row[2].text = "COMP."
    row[3].text = "DEV."
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
    for i in range(tabla_intervenciones_formato.shape[0]):
        for j in range(tabla_intervenciones_formato.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato.values[i,j])
        
    ## Header de la tabla
    #row = tabla2_interv.rows[0].cells
    #row[0].text = "INTERVENCIONES"
    #row[1].text = "PIM"
    #row[2].text = "COMP."
    #row[3].text = "DEV."
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    #row[6].text = "% DEV COSTO AL MES"
    #for i in range(tabla_intervenciones_formato_nocas.shape[0]):
    #    for j in range(tabla_intervenciones_formato_nocas.shape[-1]):
    #        tabla2_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_nocas.values[i,j])
    
    parrafo_espacio_int = document.add_paragraph('')   
    interv_parrafo3 = document.add_paragraph(
    "A través de los Decretos Supremos N°s 092, 169, 189, 209 y 210-2021-EF, \
    se realizaron todas las transferencias de partidas programadas para el Año \
    Fiscal 2021 para el financiamiento de las intervenciones y acciones pedagógicas \
    hasta el 31 de diciembre.", style="List Bullet")
    interv_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    interv_parrafo4 = document.add_paragraph(
    "Es importante considerar que la ejecución en los Contratos Administrativos de Servicios \
    (CAS) se ha visto afectada por la vigencia de la Ley N° 31131. Actualmente , \
    el Decreto de Urgencia N° 083-2021, deja sin efecto la vigencia de la Ley \
    N° 31131, por lo cual se autoriza a las entidades de la Administración \
    Pública a contratar.", style="List Bullet")
    interv_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    ###########################################################
    # Incluimos sección 2 Mascarillas y protectores faciales #
    ###########################################################
    document.add_heading("2. Mascarillas y protectores faciales", level=1)
    mascarillas_fecha_actualizacion = document.add_paragraph()
    mascarillas_fecha_negrita1 = mascarillas_fecha_actualizacion.add_run("Actualizado al ")
    mascarillas_fecha_negrita2 = mascarillas_fecha_actualizacion.add_run(fecha_corte_mascarillas)
    mascarillas_fecha_negrita1.bold = True    
    mascarillas_fecha_negrita2.bold = True
    
    mascarillas_parrafo1 = document.add_paragraph(
    "Mediante el Decreto de Urgencia N° 021-2021 y la Resolución de Secretaría \
General N° 047-2021-MINEDU, se transfirieron S/ ", style="List Bullet")
    mascarillas_parrafo1.add_run(transferencia_mascarilla)
    mascarillas_parrafo1.add_run(" millones de soles para \
la adquisición y distribución de mascarillas faciales textiles de uso \
comunitario para estudiantes y personal que labora en instituciones \
educativas públicas, así como protectores faciales para el mencionado \
personal.")
    mascarillas_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    mascarillas_parrafo2 = document.add_paragraph(
    "La adquisición de mascarillas y protectores faciales es condición necesaria \
para el retorno seguro a los servicios educativos presenciales y \
semipresenciales, según lo dispuesto por las “Disposiciones para la \
prestación del servicio en las instituciones y programas educativos \
públicos y privados de la Educación Básica de los ámbitos urbanos y \
rurales, en el marco de la emergencia sanitaria de la COVID-19”, \
aprobado mediante Resolución Ministerial N° 121-2021- MINEDU y modificado \
con Resoluciones Ministeriales N° 199-2021-MINEDU y N° 273-2021- MINEDU.", style="List Bullet")
    mascarillas_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    mascarillas_parrafo3 = document.add_paragraph(
    "Con fecha de corte al ", style="List Bullet")
    mascarillas_parrafo3.add_run(fecha_corte_mascarillas)
    mascarillas_parrafo3.add_run(
    ", la ejecución a nivel regional de los recursos de mascarillas faciales \
textiles protectores faciales fue del ")
    mascarillas_parrafo3.add_run(devengado_mascarillas)
    mascarillas_parrafo3.add_run(" (devengado) según se presenta a continuación:")
    mascarillas_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    # Incluir tabla 1 mascarillas
    tabla1_mascarillas = document.add_table(tabla_mascarillas_formato.shape[0]+1, tabla_mascarillas_formato.shape[1])
    tabla1_mascarillas.style = "formato_tabla_minedu"
    #Formato de tabla pendiente, al parecer funciona sin problema en libre office
    #from docx.shared import Cm, Inches
    #tabla1_mascarillas.allow_autofit = False
    #for cell in tabla1_mascarillas.columns[0].cells:
    #    cell.width = Inches(0.5)
    ## Header de la tabla revisar
    row_mascarilla = tabla1_mascarillas.rows[0].cells
    row_mascarilla[0].text = "Unidad Ejecutora"
    row_mascarilla[1].text = "Transferencia"
    row_mascarilla[2].text = "PIM"
    row_mascarilla[3].text = "% Certificado"
    row_mascarilla[4].text = "% Comprometido"
    row_mascarilla[5].text = "% Devengado"
    ## Contenido de la tabla
    for j in range(tabla_mascarillas_formato.shape[-1]):
        tabla1_mascarillas.cell(0,j).text = tabla_mascarillas_formato.columns[j]
    for i in range(tabla_mascarillas_formato.shape[0]):
        for j in range(tabla_mascarillas_formato.shape[-1]):
            tabla1_mascarillas.cell(i+1,j).text = str(tabla_mascarillas_formato.values[i,j])

    mascarillas_parrafo4 = document.add_paragraph().add_run("Notas: Recursos transferidos mediante el Decreto de Urgencia N° 021-2021. Fuente: SIAF MPP al 03 de octubre de 2021.")
    mascarillas_parrafo4.font.size = Pt(7)
    
    ################################################
    # Incluimos sección 3 Compromisos de desempeño #
    ################################################

    document.add_heading("3. Compromisos de desempeño", level=1)
    cdd_fecha_actualizacion = document.add_paragraph()
    cdd_fecha_negrita1 = cdd_fecha_actualizacion.add_run("Actualizado al ")
    cdd_fecha_negrita2 = cdd_fecha_actualizacion.add_run(fecha_corte_compromisos)
    cdd_fecha_negrita1.bold = True    
    cdd_fecha_negrita2.bold = True
    
    cdd_parrafo1 = document.add_paragraph(
        "En el marco de la Norma Técnica para la implementación del mecanismo \
denominado Compromisos de Desempeño 2021, aprobada por Resolución Ministerial \
N° 042-2021-MINEDU y modificada por la Resolución Ministerial N° 160-2021-MINEDU, \
se han realizado transferencias de partidas a favor de las Unidades Ejecutoras de \
Educación del Gobierno Regional de ", style="List Bullet")
    cdd_parrafo1.add_run(region)
    cdd_parrafo1.add_run(" por la suma de S/. ")
    cdd_parrafo1.add_run(cdd_transferido)
    cdd_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    # Incluimos tabla de CDD
    cdd_titulo = document.add_paragraph()
    cdd_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    cdd_tit = cdd_titulo.add_run("Transferencias de compromisos de desempeño")
    cdd_tit.bold = True        
    tabla1_cdd = document.add_table(tabla_cdd_formato.shape[0]+1, tabla_cdd_formato.shape[1])
    tabla1_cdd.autofit = True
    tabla1_cdd.allow_autofit = True
    tabla1_cdd.style = "formato_tabla_minedu"
    # Header de la tabla
    row = tabla1_cdd.rows[0].cells
    row[0].text = "Programa presupuestal"
    row[1].text = "Genérica"
    row[2].text = "Total"
    row[3].text = "DS 085-2021-EF"
    row[4].text = "DS 218-2021-EF"
    row[5].text = "DS 220-2021-EF"
    ## Contenido de la tabla
    for i in range(tabla_cdd_formato.shape[0]):
        for j in range(tabla_cdd_formato.shape[-1]):
            tabla1_cdd.cell(i+1,j).text = str(tabla_cdd_formato.values[i,j])
    
    #################################################################
    # Incluimos sección 4 financiamiento de conceptos remunerativos #
    #################################################################
            
    # Incluimos sección 1 conceptos remunerativos
    document.add_heading("Sobre el financiamiento de conceptos remunerativos", level=1)
    run.underline = True
    document.add_heading("1.Pago de Encargaturas", level=1)
    
    ##Párrafos
    encarg_parrafo1 = document.add_paragraph(
    "Para la región ")
    encarg_parrafo1.add_run(f'{region}, por concepto de encargaturas, ') 
    encarg_parrafo1.add_run('se ha calculado para el ')
    encarg_parrafo1.add_run(datetime.today().strftime('%Y'))
    encarg_parrafo1.add_run(' un costo de S/.')
    encarg_parrafo1.add_run(f'{costo_enc}') #Insertar valor de base de datos
    encarg_parrafo1.add_run(
    ' que incluye la Jornada de Trabajo Adicional de 10 horas, \
la carga social vinculada y la asignación por cargo de \
los profesores que asumen cargos de mayor responsabilidad \
mediante encargaturas')
    nota1 = encarg_parrafo1.add_run('[1]')
    nota1.font.superscript = True

    encarg_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo2 = document.add_paragraph('Para financiar estos conceptos, el')
    #Insertar valor del año pasado
    encarg_parrafo2.add_run(' MINEDU gestionó una programación directa de recursos\
    en el PIA 2021 de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo2.add_run(region)  
    encarg_parrafo2.add_run(' por el monto de S/.')
    encarg_parrafo2.add_run(f'{apmeincre}') #Insertar valor de base de datos
    encarg_parrafo2.add_run('  en la finalidad 0267929 Pago de la asignación por jornada \
de trabajo adicional y asignación por cargo de mayor responsabilidad, \
la cuál es usada para financiar las encargaturas. Asimismo, el Pliego Regional \
ya contaba con una programación de ')
    encarg_parrafo2.add_run(f'{prog_gore}') #Insertar valor de base de datos
    nota2 = encarg_parrafo2.add_run('[2]')
    nota2.font.superscript = True
    
    encarg_parrafo2.add_run(' en la misma finalidad y mediante ')
    encarg_parrafo2.add_run(' Oficio Múltiple N° 00082-2021-MINEDU/SPE-OPEP-UPP, ')
    encarg_parrafo2.add_run('se le solicitó a las Unidades Ejecutoras del Pliego Regional \
realizar modificaciones presupuestarias por el monto de S/.')
    encarg_parrafo2.add_run(f'{nm_enca}') #Insertar valor de base de datos
    encarg_parrafo2.add_run(' para habilitar la finalidad 0267929')
    nota3 = encarg_parrafo2.add_run('[3]')
    nota3.font.superscript = True
    
    encarg_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo3 = document.add_paragraph('Con ')
    encarg_parrafo3.add_run('Decreto Supremo N° 217-2021-EF \
publicado el 27 de agosto de 2021 en el marco de lo autorizado en el literal b) \
del numeral 40.1 de la Ley de Presupuesto 2021, se ha realizado una transferencia \
de partidas por el monto de S/.')  # Este párrafo tendrá que variar año tras año
    encarg_parrafo3.add_run(f'{ds_217}') #Insertar valor de base de datos    
    encarg_parrafo3.add_run(' a favor de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo3.add_run(region)
    encarg_parrafo3.add_run(' para financiar el costo diferencial.')
    encarg_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo_add = document.add_paragraph('Actualmente está en gestión \
en el MINEDU la segunda transferencia de recursos por concepto de encargaturas, \
el cual debería aprobarse como máximo el 26 de noviembre del 2021 \
de acuerdo al plazo legal establecido en la Ley de Presupuesto 2021. ')

    encarg_parrafo4 = document.add_paragraph('En el siguiente cuadro se muestran el costo \
y los montos programados/transferidos a la Región ')
    encarg_parrafo4.add_run(region)
    encarg_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla_enc = document.add_table(tabla_encargaturas_resumen.shape[0]+1, tabla_encargaturas_resumen.shape[1])
    tabla_enc.style = "formato_tabla_minedu"
    for j in range(tabla_encargaturas_resumen.shape[-1]):
        tabla_enc.cell(0,j).text = tabla_encargaturas_resumen.columns[j]
    for i in range(tabla_encargaturas_resumen.shape[0]):
        for j in range(tabla_encargaturas_resumen.shape[-1]):
            tabla_enc.cell(i+1,j).text = str(tabla_encargaturas_resumen.values[i,j])
    
    nota_tabla = document.add_paragraph().add_run('Nota: Para el cálculo del monto \
transferido, se ha tomado en cuenta los recursos programados por el Minedu y \
las UE del Pliego Regional a nivel de estructura funcional programática (EFP). \
En casos en los que la necesidad (costo) por EFP haya sido inferior a los \
programado, se generarían saldos, ocasionando que exista diferencias entre el \
costo y la suma de los montos transferidos y programados.')    

    nota_tabla.font.size = Pt(7)


#----------------------------------------------------------------------------------------------#
    document.add_heading("2.Pago de Asignaciones Temporales", level=1)
    
    ##Párrafos
    encarg_parrafo5 = document.add_paragraph('Para la Región ')    
    encarg_parrafo5.add_run(region)
    encarg_parrafo5.add_run(', por concepto de Asignaciones Temporales \
por prestar servicios en condiciones especiales, se ha calculado para el 2021 un costo de \
S/ ')    
    encarg_parrafo5.add_run(f'{costo_at} ') #Insertar valor de base de datos    
    encarg_parrafo5.add_run('que incluye el pago por prestar servicios \
en zonas rurales, de frontera, VRAEM, Instituciones Educativas Unidocentes, Multigrado \
Bilingüe y acreditar dominio de lengua originaria, de los profesores y auxiliares de \
educación nombrados y contratados.')        
    encarg_parrafo5.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    
    encarg_parrafo6 = document.add_paragraph('Para financiar estos conceptos, el ')
    encarg_parrafo6.add_run('2020 ') #Calcular valor de año anterior
    encarg_parrafo6.add_run('el MINEDU gestionó una programación directa de recursos \
en el PIA ')
    encarg_parrafo6.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo6.add_run(' de las Unidades Ejecutoras de Educación de la Región ')    
    encarg_parrafo6.add_run('por el monto de S/.')
    encarg_parrafo6.add_run(f'{apm_at}') #Insertar valor de base de datos    

    encarg_parrafo6.add_run(' en la finalidad 0267928. Pago de las asignaciones \
por tipo y ubicacion de Institucion Educativa la cuál es usada para financiar \
las asignaciones temporales.')    

    encarg_parrafo6.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
  
    encarg_parrafo7 = document.add_paragraph('Con Decreto Supremo N° 187-2021-EF \
publicado el 22 de julio de 2021 en el marco de lo autorizado en los literales \
a), c), d) y e) del numeral 40.1 de la Ley de Presupuesto 2021, ') # Este párrafo tendrá que variar año tras año
    encarg_parrafo7.add_run('se ha realizado una transferencia \
de partidas por el monto de S/. ')
    encarg_parrafo7.add_run(f'{ds_187_at}') #Insertar valor de base de datos    
    encarg_parrafo7.add_run(' a favor de las Unidades Ejecutoras de Educación \
de la Región ') 
    encarg_parrafo7.add_run(region)
    encarg_parrafo7.add_run(' para financiar el costo diferencial de las asignaciones \
temporales a favor los profesores y auxiliares de educación nombrados y contratados.')    
    encarg_parrafo7.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo8 = document.add_paragraph('Actualmente está en gestión en el MINEDU \
la segunda transferencia de recursos por concepto de asignaciones temporales, \
el cual debería aprobarse como máximo el ')
    encarg_parrafo8.add_run('26 de noviembre del 2021') #Esta fecha se actualizará año a año 
    encarg_parrafo8.add_run(' de acuerdo al plazo legal establecido en la Ley de Presupuesto ') 
    encarg_parrafo8.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo8.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo9 = document.add_paragraph('En el siguiente cuadro se muestran el costo y los \
montos programados/transferidos a la Región ')
    encarg_parrafo9.add_run(region)
    encarg_parrafo9.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla_at = document.add_table(tabla_at_resumen.shape[0]+1, tabla_at_resumen.shape[1])
    tabla_at.style = "formato_tabla_minedu"
    for j in range(tabla_at_resumen.shape[-1]):
        tabla_at.cell(0,j).text = tabla_at_resumen.columns[j]
    for i in range(tabla_at_resumen.shape[0]):
        for j in range(tabla_at_resumen.shape[-1]):
            tabla_at.cell(i+1,j).text = str(tabla_at_resumen.values[i,j])

    parrafo_espacio = document.add_paragraph('')        


#----------------------------------------------------------------------------------------------#

    document.add_heading("3.Pago de Beneficios Sociales", level=1)

    region_seleccionada = df_transferencia['region'] == region #Seleccionar region
    tabla_transferencia = df_transferencia[region_seleccionada]
    tabla_transferencia_formato = df_transferencia[region_seleccionada]

    formato_tabla_transferencia = {
        "norma_de_transferencia" : "{}",
        "concepto" : "{}",
        "monto_transferido" : "{:,.0f}",
        }


    tabla_transferencia_formato = tabla_transferencia_formato.transform({k: v.format for k, v in formato_tabla_transferencia.items()})        


    ##Párrafos    
    encarg_parrafo10 = document.add_paragraph('Para la Región ')
    encarg_parrafo10.add_run(region)
    encarg_parrafo10.add_run(', por concepto de Beneficios Sociales se ha calculado\
 un costo total de S/. ')
    encarg_parrafo10.add_run(f'{costo_bs}') #Insertar valor de base de datos    
    encarg_parrafo10.add_run(', aque incluye el pago para la Asignación por \
Tiempo de Servicios (ATS), Compensación por Tiempo de Servicios (CTS) \
y, Subsidio por Luto y Sepelio (SLS) a favor de los profesores y \
auxiliares de educación nombrados y contratados. Dentro de estos, \
de acuerdo con la nueva estrategia para el pago oportuno de Beneficios Sociales \
implementada por el MINEDU, se han aprobado pagos hasta por un costo de S/. ')
    encarg_parrafo10.add_run(f'{lista_bf}') #Insertar valor de base de datos    
    encarg_parrafo10.add_run(' a la fecha')
    
    encarg_parrafo10.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo11 = document.add_paragraph('Para financiar estos conceptos, el ')
    encarg_parrafo11.add_run('2020 ') #Calcular valor de año anterior
    encarg_parrafo11.add_run('el MINEDU gestionó una programación directa de recursos en el PIA ')
    encarg_parrafo11.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo11.add_run(' de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo11.add_run(region)
    encarg_parrafo11.add_run(' por el monto de S/. ')
    encarg_parrafo11.add_run(f'{apm_bs},') #Insertar valor de base de datos    
    encarg_parrafo11.add_run(' lo cual fue comunicado a través del ')
    encarg_parrafo11.add_run('Oficio Múltiple N° 00011-2021-MINEDU/SPE-OPEP-UPP') #Esto cambiará cada año
    encarg_parrafo11.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo12 = document.add_paragraph('Con ')
    encarg_parrafo12.add_run('Decreto Supremo N° 072-2021-EF publicado el 21 de abril de 2021')
    encarg_parrafo12.add_run(' en el marco de lo autorizado en los literales a), d) y e) \
del numeral 40.1 de la Ley de Presupuesto ')
    encarg_parrafo12.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo12.add_run(', se ha realizado una transferencia de partidas por el monto de S/ ')   
    encarg_parrafo12.add_run(f'{ds_72_bs}.') #Insertar valor de base de datos   
    encarg_parrafo12.add_run(', a favor de las Unidades Ejecutoras de Educación de la Región ')   
    encarg_parrafo12.add_run(region)
    encarg_parrafo12.add_run(' para financiar el pago de los beneficios sociales a favor los profesores \
y auxiliares de educación nombrados y contratados que fueron reconocidos hasta el ')   
    encarg_parrafo12.add_run('12 de enero de 2021')
    encarg_parrafo12.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo13 = document.add_paragraph('Asimismo, mediante ')
    encarg_parrafo13.add_run('Decreto Supremo N° 256-2021-EF publicado el 24 de setiembre de 2021, ')  #Esto cambiará cada año
    encarg_parrafo13.add_run(' se realizó la segunda transferencia de recursos por concepto de \
beneficios sociales a favor de docentes y auxiliares nombrados y contratados, cuyos beneficios fueron \
reconocidos hasta el 4 de junio de 2021 ')
    encarg_parrafo13.add_run(', transfiriéndose S/. ')
    encarg_parrafo13.add_run(f'{ds_256_bs}.') #Insertar valor de base de datos   
    encarg_parrafo13.add_run('  a las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo13.add_run(region)
    encarg_parrafo13.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

   
    encarg_parrafo14 = document.add_paragraph('En el siguiente cuadro se muestran el costo y los montos \
programados/transferidos a la Región ')
    encarg_parrafo14.add_run(region)
    encarg_parrafo14.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla_bs = document.add_table(tabla_bs_resumen.shape[0]+1, tabla_bs_resumen.shape[1])
    tabla_bs.style = "formato_tabla_minedu"
      
    for j in range(tabla_bs_resumen.shape[-1]):
        tabla_bs.cell(0,j).text = tabla_bs_resumen.columns[j]
    for i in range(tabla_bs_resumen.shape[0]):
        for j in range(tabla_bs_resumen.shape[-1]):
            tabla_bs.cell(i+1,j).text = str(tabla_bs_resumen.values[i,j])

    parrafo_espacio1 = document.add_paragraph('')    

    encarg_parrafo15 = document.add_paragraph(' De la misma forma, durante el presente año, para la Región ')  
    encarg_parrafo15.add_run(region)
    encarg_parrafo15.add_run(' se ha realizado las siguientes transferencias:')

    encarg_parrafo15.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #Insertar tabla

    tabla_transf = document.add_table(tabla_transferencia_formato.shape[0]+1, tabla_transferencia_formato.shape[1])
    tabla_transf.style = "formato_tabla_minedu"
    ## Header de la tabla
    row = tabla_transf.rows[0].cells
    row[0].text = "NORMA"
    row[1].text = "CONCEPTO"
    row[2].text = "COSTO"

    for i in range(tabla_transferencia_formato.shape[0]):
        for j in range(tabla_transferencia_formato.shape[-1]):
            tabla_transf.cell(i+1,j).text = str(tabla_transferencia_formato.values[i,j])

    parrafo_espacio2 = document.add_paragraph('')        

    encarg_parrafo16 = document.add_paragraph(' Por otro lado, para el año ')
    encarg_parrafo16.add_run('2022 ') #Calcular año posterior
    encarg_parrafo16.add_run('el MINEDU está gestionando la programación parcial de recursos en los \
presupuestos de las Unidades Ejecutoras para atender encargaturas, asignaciones temporales, \
beneficios sociales, entre otros y, el financiamiento restante, se realizará de manera oportuna el ')
    encarg_parrafo16.add_run('2022, ') #Calcular año posterior
    encarg_parrafo16.add_run('preferentemente antes que termine el primer semestre de dicho año fiscal.')    
    encarg_parrafo16.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

#------------------------------------------------------------------------------------------------------------#
#------------------------------------------------------------------------------------------------------------#    


    # Incluimos sección 2 de Proceso de racionalización

    document.add_heading("Sobre el proceso de racionalización", level=1)
    run.underline = True  

    document.add_heading("4. Resultados proceso de racionalización 2020", level=1)
 
    region_seleccionada2 = data_brecha['region'] == region #Seleccionar region    
    tabla_brecha = data_brecha[region_seleccionada2]
    
    excedentes_region = str('{:,.0f}'.format(tabla_brecha["doc_e"].sum()))
    requerimientos_region = str('{:,.0f}'.format(tabla_brecha["doc_req"].sum()))
    requerimiento_neto = str('{:,.0f}'.format(tabla_brecha["req_neto"].sum()))
    excedencia_neta = str('{:,.0f}'.format(tabla_brecha["exc_neto"].sum()))

    region_seleccionada3 = data_brecha_regional['region'] == region #Seleccionar region
    tabla_brecha2 = data_brecha_regional[region_seleccionada3]
    ugel_cant_req = str('{:,.0f}'.format(tabla_brecha2["cant_ugel_req"].sum()))
    ugel_cant_exc = str('{:,.0f}'.format(tabla_brecha2["cant_ugel_exc"].sum()))    
    brecha_region = str('{:,.0f}'.format(tabla_brecha2["brecha_net"].sum()))  

    # Generamos fila total
    total_brecha = tabla_brecha[["region", "doc_req", "doc_e", "doc_e_n_cub_req", "nom_exd_mov1", "doc_e_c", "req_neto", "exc_neto" ]]. \
    groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_brecha = tabla_brecha.append(total_brecha, ignore_index=True)
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_brecha['ugel'] = tabla_brecha['ugel'].fillna("Total")

    tabla_brecha_formato = data_brecha[region_seleccionada2]
    
    formato_tabla_brecha = {
        "ugel": "{}",
        "doc_req" : "{}",
        "doc_e" : "{}",
        "doc_e_n_cub_req" : "{}",
        "nom_exd_mov1" : "{}",
        "doc_e_c" : "{}",
        "req_neto" : "{:.0f}",
        "exc_neto" : "{:.0f}",
        }
    
    tabla_brecha_formato = tabla_brecha
    tabla_brecha_formato = tabla_brecha_formato.transform({k: v.format for k, v in formato_tabla_brecha.items()})
    
    
    #Párrafo 
    racio_parrafo2 = document.add_paragraph('En el proceso de racionalización ',  style="List Bullet")    
    racio_parrafo2.add_run('2020 ') #Calcular valor de año anterior
    racio_parrafo2.add_run(', se identificó en la región ')
    racio_parrafo2.add_run(region)
    racio_parrafo2.add_run(' un total de ')
    negrita1 = racio_parrafo2.add_run(excedentes_region) #Insertar valor de base de datos
    negrita1.bold= True
    negrita2 = racio_parrafo2.add_run(' plazas de docentes de aula excedentes ')
    negrita2.bold= True
    racio_parrafo2.add_run('y ')
    negrita3 = racio_parrafo2.add_run(requerimientos_region) #Insertar valor de base de datos
    negrita3.bold= True    
    negrita4 = racio_parrafo2.add_run(' plazas de requerimiento. ')
    negrita4.bold= True
    racio_parrafo2.add_run('A partir de esos resultados, se procedió a calcular \
el requerimiento y la excedencia por UGEL y el agregado a nivel regional, \
ello se puede observar en las dos últimas columnas del siguiente cuadro: ')  

    # Incluimos tabla brecha
    tabla1_brecha = document.add_table(tabla_brecha_formato.shape[0]+1, tabla_brecha_formato.shape[1])
    tabla1_brecha.style = "formato_tabla_minedu"
    ## Header de la tabla
    row = tabla1_brecha.rows[0].cells
    row[0].text = "UGEL"
    row[1].text = "Total Req"
    row[2].text = "Total Exd"    
    row[3].text = "Exc. Docentes nombrados que pueden cubrir requerimiento"    
    row[4].text = "Exc. Docentes nombrados con dificultad de movimiento"
    row[5].text = "Excedentes vacantes para cubrir requerimientos"
    row[6].text = "Req. Neto final"
    row[7].text = "Exc. Neto final"
    
    ## Contenido de la tabla
    for i in range(tabla_brecha_formato.shape[0]):
        for j in range(tabla_brecha_formato.shape[-1]):
            tabla1_brecha.cell(i+1,j).text = str(tabla_brecha_formato.values[i,j])

    racio_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    parrafo_espacio3 = document.add_paragraph('')        
    
    racio_parrafo3 = document.add_paragraph(' Por lo tanto, a nivel regional \
se contaba con una brecha interna de ',  style="List Bullet")    
    racio_parrafo3.add_run(requerimiento_neto) #Insertar valor de base de datos
    racio_parrafo3.add_run(' plazas en ')
    racio_parrafo3.add_run(ugel_cant_req) #Insertar valor de base de datos
    racio_parrafo3.add_run(' UGEL, y un excedente neto de \
plazas vacantes ascendente a ')
    racio_parrafo3.add_run(excedencia_neta) #Insertar valor de base de datos  
    racio_parrafo3.add_run(' plazas en ')   
    racio_parrafo3.add_run(ugel_cant_exc) #Insertar valor de base de datos
    racio_parrafo3.add_run(' UGEL. Con ello, se obtuvo ')
      
    if region == "AYACUCHO" or region == "HUANCAVELICA" or region=="PUNO" or region=="TUMBES":       
        negrita5 = racio_parrafo3.add_run('una excedencia neta a nivel regional igual a ')
        negrita5.bold= True
        negrita6 = racio_parrafo3.add_run(brecha_region) #Insertar valor de base de datos
        negrita6.bold= True     
        racio_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    else:            
        negrita5 = racio_parrafo3.add_run('un requerimiento neto a nivel regional igual a ')
        negrita5.bold= True
        negrita6 = racio_parrafo3.add_run(brecha_region) #Insertar valor de base de datos
        negrita6.bold= True     
        racio_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    #Creacion proceso de racionalización
    
    region_seleccionada = data_creacion['region'] == region #Seleccionar region
    tabla_creacion = data_creacion[region_seleccionada]
    creacion_region = str('{:,.0f}'.format(tabla_creacion["creacion_total"].sum()))

    # Generamos fila total
    total_creacion = tabla_creacion[["region", "inicial", "primaria", "secundaria", "creacion_total"]]. \
    groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_creacion = tabla_creacion.append(total_creacion, ignore_index=True)
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_creacion['ugel'] = tabla_creacion['ugel'].fillna("Total")

    tabla_creacion_formato = data_creacion[region_seleccionada]
     

    formato_tabla_creacion = {
        "ugel": "{}",
        "inicial": "{:.0f}",
        "primaria": "{:.0f}",
        "secundaria": "{:.0f}",        
        "creacion_total" : "{}",
        }
    
    tabla_creacion_formato = tabla_creacion
    tabla_creacion_formato = tabla_creacion_formato.transform({k: v.format for k, v in formato_tabla_creacion.items()})        

    region_pem = data_creacion_pem['region'] == region #Seleccionar region
    tabla_creacion_pem = data_creacion_pem[region_pem]
    pem_docente = str('{:,.0f}'.format(tabla_creacion_pem["req_doc"].sum()))
    pem_bolsa = str('{:,.0f}'.format(tabla_creacion_pem["req_bolsa"].sum()))
    pem_director = str('{:,.0f}'.format(tabla_creacion_pem["req_director"].sum()))
    pem_subdirector = str('{:,.0f}'.format(tabla_creacion_pem["req_subdir"].sum()))  


    document.add_heading("5. Financiamiento de plazas 2021", level=1)

    creacion_racio = document.add_paragraph().add_run('En el marco del proceso de racionalizacion 2020')

    creacion_racio.italic = True
    creacion_racio.bold = True

    #Párrafo creacion
    racio_parrafo1 = document.add_paragraph('A través del ',  style="List Bullet")
    racio_parrafo1.add_run('DS N° 078-2021-EF') #Esto cambia año tras año
    racio_parrafo1.add_run(' se financiaron ') 
    racio_parrafo1.add_run(creacion_region) #Insertar valor de base de datos 
    racio_parrafo1.add_run(' plazas de docentes de aula en el marco de los resultados del \
proceso de racionalización ') 
    racio_parrafo1.add_run('2020 ') #Calcular valor de año anterior
    racio_parrafo1.add_run('en servicios educativos públicos de la región ')
    racio_parrafo1.add_run(region)
    racio_parrafo1.add_run(' con la siguiente distribución por UGEL:')
    racio_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # Incluimos tabla creaciones
    creacion_titulo = document.add_paragraph('Número de plazas creadas por UGEL y nivel')
    creacion_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    tabla1_creacion = document.add_table(tabla_creacion_formato.shape[0]+1, tabla_creacion_formato.shape[1])
    tabla1_creacion.style = "formato_tabla_minedu"
    ## Header de la tabla
    row = tabla1_creacion.rows[0].cells
    row[0].text = "UGEL"
    row[1].text = "Inicial"
    row[2].text = "Primaria"
    row[3].text = "Secundaria"
    row[4].text = "Número de creaciones"

    ## Contenido de la tabla
    for i in range(tabla_creacion_formato.shape[0]):
        for j in range(tabla_creacion_formato.shape[-1]):
            tabla1_creacion.cell(i+1,j).text = str(tabla_creacion_formato.values[i,j])

#region == "ANCASH" or region == "APURIMAC"  or region == "AREQUIPA" or region == "AYACUCHO" or region == "CALLAO"  or region == "CUSCO" or region == "ICA"

    #Párrafo creacion PEM
    if pem_docente!="0" or pem_bolsa!="0" or pem_director!="0" or pem_subdirector!="0":
        parrafo_espacio_pem = document.add_paragraph('')        

        creacion_pem = document.add_paragraph().add_run('En el marco del Proceso Extraordinario de Matrícula (PEM) 2021')
        creacion_pem.italic = True
        creacion_pem.bold = True
            #Creacion PEM 2021




        # Generamos fila total
        total_creacion_pem = tabla_creacion_pem[["region", "req_doc", "req_bolsa", "req_director", "req_subdir"]]. \
            groupby(by = ["region"], as_index=False).sum()
    
        # Realizamos append del total en la tabla
        tabla_creacion_pem = tabla_creacion_pem.append(total_creacion_pem, ignore_index=True)
    
        #Incluimos palabra "total" y "-" en vez de NaN
        tabla_creacion_pem['ugel'] = tabla_creacion_pem['ugel'].fillna("Total")
        tabla_creacion_pem['modalidad'] = tabla_creacion_pem['modalidad'].fillna("-")

        tabla_creacion_pem_formato = data_creacion_pem[region_pem]
     

        formato_tabla_creacion_pem = {
            "ugel": "{}",
            "modalidad": "{}",
            "req_doc": "{:.0f}",
            "req_bolsa": "{:.0f}",
            "req_director": "{:.0f}",        
            "req_subdir" : "{}",
            }
        
        tabla_creacion_pem_formato = tabla_creacion_pem
        tabla_creacion_pem_formato = tabla_creacion_pem_formato.transform({k: v.format for k, v in formato_tabla_creacion_pem.items()})      

        
        pem_parrafo1 = document.add_paragraph('A través del ',  style="List Bullet")
        pem_parrafo1.add_run('DS N° 065-2021-EF') #Esto cambia año tras año
        pem_parrafo1.add_run(' se financiaron ')
    
        if pem_docente!="0":
            pem_parrafo1.add_run(pem_docente) #Insertar valor de base de datos 
            pem_parrafo1.add_run(' plazas de docentes de aula, ') 
    
        if pem_bolsa!="0":
            pem_parrafo1.add_run(pem_bolsa) #Insertar valor de base de datos 
            pem_parrafo1.add_run(' horas de bolsa, ') 

        if pem_director!="0":
            pem_parrafo1.add_run(pem_director) #Insertar valor de base de datos 
            pem_parrafo1.add_run(' plazas de director, ')         

        if pem_subdirector!="0":
            pem_parrafo1.add_run(pem_subdirector) #Insertar valor de base de datos 
            pem_parrafo1.add_run(' plazas de subdirector, ') 

        pem_parrafo1.add_run('el detalle por UGEL y modalidad se muestra en la siguiente tabla')

        # Incluimos tabla creaciones
        creacion_titulo_pem = document.add_paragraph('Número de plazas creadas por UGEL y modalidad')
        creacion_titulo_pem.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        tabla2_creacion = document.add_table(tabla_creacion_pem_formato.shape[0]+1, tabla_creacion_pem_formato.shape[1])
        tabla2_creacion.style = "formato_tabla_minedu"
        ## Header de la tabla
        row = tabla2_creacion.rows[0].cells
        row[0].text = "UGEL"
        row[1].text = "Modalidad"
        row[2].text = "Plazas docentes de aula"
        row[3].text = "Bolsa de horas"
        row[4].text = "Plazas de director"
        row[5].text = "Plazas de subdirector"

        ## Contenido de la tabla
        for i in range(tabla_creacion_pem_formato.shape[0]):
            for j in range(tabla_creacion_pem_formato.shape[-1]):
                tabla2_creacion.cell(i+1,j).text = str(tabla_creacion_pem_formato.values[i,j])

    document.add_heading("6. Acciones de reordenamiento territorial 2020", level=1)
 
    #Párrafo
    if region == "AREQUIPA" or region == "AYACUCHO" or region=="CUSCO" or region=="HUANCAVELICA" or region=="PASCO" or region=="PUNO" or region=="TACNA":
        region_seleccionada3 = data_bloqueo['region'] == region #Seleccionar region    
        tabla_bloqueo = data_bloqueo[region_seleccionada3]
        bloqueos = str('{:,.0f}'.format(tabla_bloqueo['cant_bloqueos'].sum()))
        racio_parrafo4 = document.add_paragraph('En el marco del proceso de racionalización ',  style="List Bullet") 
        racio_parrafo4.add_run('2020 ') #Calcular valor de año anterior
        racio_parrafo4.add_run(', en la región ')
        racio_parrafo4.add_run(region)    
        racio_parrafo4.add_run(' se inhabilitaron ') 
        racio_parrafo4.add_run(bloqueos) 
        racio_parrafo4.add_run(' plazas de un total de ') 
        racio_parrafo4.add_run(excedencia_neta) #Insertar valor de base de datos
        racio_parrafo4.add_run(' plazas vacantes identificadas como excedencia neta.')
        racio_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    else:    
        racio_parrafo4 = document.add_paragraph('En el marco del proceso de racionalización ',  style="List Bullet") 
        racio_parrafo4.add_run('2020 ') #Calcular valor de año anterior
        racio_parrafo4.add_run(', en la región ')
        racio_parrafo4.add_run(region)    
        racio_parrafo4.add_run(' no se inhabilitaron plazas a pesar de contar con ')
        racio_parrafo4.add_run(excedencia_neta) #Insertar valor de base de datos
        racio_parrafo4.add_run(' plazas vacantes identificadas como excedencia neta.')
        racio_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

#------------------------------------------------------------------------------------------------------------#
#------------------------------------------------------------------------------------------------------------#    
    # Incluimos sección 3 deudas sociales y contratación
    tabla_deuda_social = data_deuda_social
    tabla_deuda_social_formato = data_deuda_social
    
    formato_tabla_deuda_social = {
        "seccion_pliego": "{}",
        "monto" : "{:,.0f}",
        }
    tabla_deuda_social_formato = tabla_deuda_social_formato.transform({k: v.format for k, v in formato_tabla_deuda_social.items()})        

    document.add_heading("Sobre Deudas Sociales y Contratación del DL 276", level=1)
    run.underline = True  

    document.add_heading("7. Deudas Sociales", level=1)

    #Párrafo
    deuda_parrafo1 = document.add_paragraph(' La Deuda Social del Sector Educación \
está constituida por las obligaciones de pago  en materia laboral y previsional, \
respecto del personal Docente y Auxiliares de Educación provenientes \
de la derogada Ley del Profesorado  y su modificatoria; \
del personal Administrativo sujeto al Decreto Legislativo Nº 276, \
Ley de Bases de la Carrera Administrativa; del Personal Administrativo \
en el marco del Texto Único Ordenado (TUO) del Decreto Legislativo Nº 728 \
Ley de Productividad y Competitividad  Laboral  aprobado por \
Decreto Supremo Nº 003-97-TR. ') 
    deuda_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo2 = document.add_paragraph(' La atención de la deuda social \
con el Sector se viene efectuando a través de transferencias de \
partidas del Tesoro Público; es decir, en forma complementaria a \
los recursos presupuestales con que disponen los Pliegos del Gobierno Nacional \
y los Gobiernos Regionales, para la atención del pago de sentencias judiciales \
en calidad de cosa juzgada y en ejecución. Dicho tratamiento, \
se realiza desde el año 2014 en el marco de las Leyes Anuales de Presupuesto, \
habiéndose autorizado diversas habilitaciones presupuestarias \
en el nivel institucional mediante los decretos supremos correspondientes. ')
    deuda_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo3 = document.add_paragraph(' Desde el año 2018 se han venido \
destinando S/ 200 000 000,00 (DOSCIENTOS MILLONES y 00/100 DE SOLES) \
para el sector Educación, que según los criterios de priorización aprobados \
por el Ministerio de Educación, son destinados preferentemente \
al pago de las sentencias judiciales que reconocen \
el derecho de preparación de clases, frente a otros conceptos. ')
    deuda_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo4 = document.add_paragraph(' En ese sentido, mediante el numeral \
1.2 del artículo 1 del Decreto Supremo N° 216-2021-EF, publicado en el \
Diario Oficial “El Peruano” el 27 de agosto de 2021, se autorizó \
la Transferencia de Partidas en el Presupuesto del Sector Público \
para el Año Fiscal 2021, hasta por la suma de S/ 200 000 000,00 \
(DOSCIENTOS MILLONES Y 00/100 SOLES), a favor de diversos Pliegos \
del Gobierno Nacional (dentro de los que se encuentra el Ministerio de Educación \
(MINEDU), y los Gobiernos Regionales, para financiar el pago de \
sentencias judiciales en calidad de cosa juzgada del sector Educación \
y en ejecución al 31 de diciembre de 2020, en el marco del numeral 6 de \
la de la Undécima Disposición Complementaria Final de la Ley N° 31084, \
Ley de presupuesto del año fiscal 2021, con cargo a los recursos de la Reserva \
de Contingencia del Ministerio de Economía y Finanzas. \
El detalle de dicha transferencia de recursos se muestra a continuación:')
    deuda_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # Incluimos tabla deuda social
    deuda_titulo = document.add_paragraph('Transferencia de Partidas a favor de diversos pliegos del Gobierno Nacional \
y Gobiernos Regionales – Sector Educación')
    deuda_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    
    tabla_deuda_social = document.add_table(tabla_deuda_social_formato.shape[0]+1, tabla_deuda_social_formato.shape[1])
    tabla_deuda_social.style = "formato_tabla_minedu"
    ## Header de la tabla
    row = tabla_deuda_social.rows[0].cells
    row[0].text = "Sección / Pliego"
    row[1].text = "Monto"

   ## Contenido de la tabla
    for i in range(tabla_deuda_social_formato.shape[0]):
        for j in range(tabla_deuda_social_formato.shape[-1]):
            tabla_deuda_social.cell(i+1,j).text = str(tabla_deuda_social_formato.values[i,j])    

    parrafo_espacio4 = document.add_paragraph('')    
        
    deuda_parrafo5 = document.add_paragraph('Del mismo modo, es importante \
señalar que las deudas sociales del Sector Educación se atienden siguiendo \
los criterios de priorización establecidos en el presente año fiscal mediante \
el Decreto Supremo N° 003-2021-MINEDU, Decreto Supremo que aprueba \
los criterios de priorización para la atención del pago de sentencias judiciales \
en calidad de cosa juzgada y en ejecución del Sector Educación, \
donde no se establece una distinción entre regímenes laborales \
o entre personal activo o cesante, sino se siguen patrones de antigüedad, \
salud, entre otros')
    deuda_parrafo5.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY


    deuda_parrafo6 = document.add_paragraph('Es igualmente importante señalar, \
que en el caso de existir déficit presupuestal durante el ejercicio fiscal \
en ejecución, para la atención de las planillas continuas de pensionistas, \
derivadas de sentencias en calidad de cosa juzgada, \
su financiamiento deberá tener en cuenta la exoneración prevista \
en el numeral 9.2 del artículo 9. Medidas en materia de modificaciones \
Presupuestarias en el Nivel Funcional Programático, de la Ley Nº 31084, \
Ley de Presupuesto del Sector Público para el Año Fiscal 2021; \
en tanto prescribe que: ')

    formato1 = deuda_parrafo6.add_run(' “A nivel de pliego')
    formato1.italic = True
    formato2 = deuda_parrafo6.add_run(' la Partida de Gasto 2.2.1 “Pensiones” no puede ser habilitadora, ')
    formato2.italic = True
    formato2.bold = True
    formato3 = deuda_parrafo6.add_run('salvo para las habilitaciones que se realicen ')
    formato3.italic = True
    formato3.bold = True
    formato3.underline = True
    formato4 = deuda_parrafo6.add_run('dentro de la misma partida entre unidades \
ejecutoras del mismo pliego presupuestario, y ')
    formato4.italic = True
    formato4.bold = True    
    formato5 = deuda_parrafo6.add_run('para la atención de sentencias judiciales \
en materia pensionaria con calidad de cosa juzgada,')
    formato5.italic = True
    formato5.bold = True
    formato5.underline = True    
    formato6 = deuda_parrafo6.add_run(' en este último caso, previo informe \
favorable de la Dirección General de Presupuesto Público (DGPP), ')
    formato6.italic = True
    formato6.bold = True  
    formato7 = deuda_parrafo6.add_run(' y de corresponder, sobre la base \
de la información registrada en el Aplicativo Informático para el \
 Registro Centralizado de Planillas y de Datos de los Recursos Humanos \
 del Sector Público (AIRHSP) que debe remitir la Dirección General \
 de Gestión Fiscal de los Recursos Humanos a la DGPP.(…)”.')
    formato7.italic = True
    deuda_parrafo6.add_run('(Lo resaltado es nuestro).')
    deuda_parrafo6.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo7 = document.add_paragraph('En el mismo numeral del artículo \
acotado en el párrafo anterior, se establece que las solicitudes de \
informe favorable para su aplicación solo pueden ser presentadas al \
Ministerio de Economía y Finanzas hasta el 15 de octubre de 2021; por lo que, ')
    formato8 = deuda_parrafo7.add_run('las unidades orgánicas pertinentes \
del Pliego Presupuestal involucrado, deben implementar \
las acciones administrativas a que hubiera lugar con sujeción \
a la normatividad presupuestaria y al plazo legal invocados líneas atrás.')
    formato8.bold = True
    deuda_parrafo7.add_run(' Así como, ')
    formato9 = deuda_parrafo7.add_run('de corresponder, \
las acciones pertinentes para el pago de las deudas que se hubieran generado \
por aplicación de la sentencia judicial en calidad de cosa juzgada \
que así lo ordene, en cuyo caso deberá observarse lo dispuesto en \
la Ley Nº 30137 y su Reglamento.')
    formato9.underline = True
    deuda_parrafo7.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_heading("8. Ruta para contratación de personal administrativo", level=1)

    #Párrafo
    pers_adm_parrafo1 = document.add_paragraph('En el marco del inciso \
d) del numeral 8.1 del artículo 8 de la Ley N° 31084, \
Ley de Presupuesto del Sector Público para el año 2021, ')
    formato10 = pers_adm_parrafo1.add_run('no se ha establecido como excepción \
la posibilidad del ingreso de personal administrativo \
bajo el régimen laboral del Decreto Legislativo N° 276.')
    formato10.underline = True
    pers_adm_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    pers_adm_parrafo2 = document.add_paragraph('A su vez, \
la Gerencia de Políticas de Gestión del Servicio Civil de \
la Autoridad Nacional del Servicio Civil (SERVIR) se ha pronunciado \
sobre la imposibilidad de contratar personal administrativo del \
Decreto Legislativo N° 276, emitido mediante Informe Técnico \
N° 331-2021-SERVIR-GPGSC.')
    pers_adm_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    pers_adm_parrafo3 = document.add_paragraph('Dado que el Decreto Legislativo \
N° 276, no solo rige al Sector Educación, sino a todo el Sector Público, \
desde el Ministerio de Educación se propuso incluir la habilitación \
para su contratación en una norma con rango de Ley, la cual autorizaría \
la contratación de personal del Decreto Legislativo N° 276. \
Sin embargo, dicha propuesta no fue incluida en el \
Proyecto de Ley del Presupuesto del Sector Público para el año 2022, \
enviada al Congreso de la República para su aprobación el 31 de agosto de 2021.')
    pers_adm_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_page_break()

    ##############################################
    # Incluimos anexo notas y fechas de actualiz #
    ##############################################
    
    document.add_heading("Notas", level=1)
    nota1 = document.add_paragraph('[1] Este costeo está en función a la información \
registrada en el sistema de control de plazas NEXUS')
    nota1.italic = True
    nota1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    nota2 = document.add_paragraph('[2] Solo se considera los clasificadores \
21.12.11 (JTA), 21.31.15 (CS) y 21.12.21 (Asignación por cargo).')
    nota2.italic = True
    nota2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    nota3 = document.add_paragraph('[3] Mediante el OM, se solicitó habilitar \
la finalidad 0267929 con los recursos destinados a financiar el costo diferencial \
de la asignación por jornada de trabajo adicional y carga social \
(debido al incremento de la RIM), ya que estos estaban programados en \
finalidades que se usaban anteriormente.')
    nota3.italic = True
    nota3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #anexo_parrafo1 = document.add_paragraph("Las fechas de actualización para las \
#secciones del documento se presentan en la tabla siguiente:")
    #tabla_anexo_corte = document.add_table(rows=1, cols=2)
    #tabla_anexo_corte.style = "formato_tabla_minedu"
    #hdr_cells = tabla_anexo_corte.rows[0].cells
    #hdr_cells[0].text = "Sección"
    #hdr_cells[1].text = "Fecha de actualización"
    #for id, name in tabla_fechas_corte:
        #row = tabla_anexo_corte.add_row().cells
        #row[0].text = str(id)
        #row[1].text = str(name)
        
    #######################
    # Guardamos documento #
    #######################
    
    document.save(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
###########################################################
# Creamos tabla con lista de files para enviar por correo #
###########################################################
    
# Generamos lista de AM.
lista_AM = glob.glob(os.path.join(proyecto, f"output/AM_{fecha_actual}/*"))

if platform.system() == "Windows":
    lista_regiones = pd.DataFrame (lista_AM)
    lista_regiones.rename( columns={0:'path'}, inplace=True )
    lista_regiones["path"] = lista_regiones["path"].str.replace('\\', '/')
    lista_regiones[['a', 'b', 'c', 'd']] = lista_regiones["path"].str.split("AM_", expand = True)
    lista_regiones[['date', 'e']] = lista_regiones["c"].str.split("/", expand = True)
    lista_regiones[['region', 'g']] = lista_regiones["d"].str.split("_", expand = True)
    lista_regiones = lista_regiones[["path", "date","region"]]
elif platform.system() == "Darwin":
    lista_regiones = pd.DataFrame (lista_AM)
    lista_regiones.rename( columns={0:'path'}, inplace=True )
    lista_regiones[['a', 'b', 'c']] = lista_regiones["path"].str.split("AM_", expand = True)
    lista_regiones[['date', 'e']] = lista_regiones["b"].str.split("/", expand = True)
    lista_regiones[['region', 'g']] = lista_regiones["c"].str.split("_", expand = True)
    lista_regiones = lista_regiones[["path", "date","region"]]

lista_regiones.to_excel(Path(proyecto, "documentacion", "lista_regiones.xlsx"), index = False)
