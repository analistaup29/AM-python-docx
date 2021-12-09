# -*- coding: utf-8 -*-
"""
Created on Thu Dec  9 14:29:10 2021

@author: analistaup29
"""

# pip install python-docx
# pip install pyjanitor
# pip install nums_from_string
import docx
import pandas as pd
import nums_from_string
import os
import glob
import getpass
from datetime import datetime
from janitor import clean_names
from pathlib import Path


if getpass.getuser() == "analistaup29": # PC Analista UP 29 Minedu
    github = Path("C:/Users/ANALISTAUP29/Documents/GitHub/AM-python-docx")
    proyecto = Path("B:/OneDrive - Ministerio de Educación/unidad_B/2021/4. Herramientas de Seguimiento/13.AM_automatizada")
elif  getpass.getuser() == "bran": # PC Brandon
    github = Path("/Users/bran/Documents/GitHub/AM-python-docx")
    proyecto = Path("/Users/bran/Documents/GitHub/AM-python-docx")


###############################################################################
# Creación de carpeta donde se guardan los outputs #
###############################################################################

# Fecha de hoy
fecha_actual = datetime.today().strftime('%d-%m-%y')

# Creación de carpeta
dir = os.path.join(proyecto, f"output/AM_testeo_grafico/AM_{fecha_actual}")
if not os.path.exists(dir):
    os.mkdir(dir)
    print("Se creó una nueva carpeta")
else:
    print("Ya existe la carpeta")
        
# Path de nueva carpeta
nueva_carpeta = Path(proyecto/ f"output/AM_testeo_grafico/AM_{fecha_actual}")


## A) Base disponibilidad ----------------------------------------------------

# Importamos los nombres de los archivos en la carpeta intervenciones pedagogicas
#lista_archivos_int = glob.glob(os.path.join(proyecto,"input/intervenciones_pedagogicas/*"))
#Mantenemos el corte de disponibilidad más reciente
#fecha_corte_disponibilidad = max(lista_archivos_int, key=os.path.getctime)
# Nos quedamos con el nombre de archivo para la base de disponibilidad
#fecha_corte_disponibilidad = os.path.split(fecha_corte_disponibilidad)
#fecha_corte_disponibilidad = fecha_corte_disponibilidad[1]
# Extraemos la fecha del nombre de archivo
#fecha_corte_disponibilidad = nums_from_string.get_numeric_string_tokens(fecha_corte_disponibilidad)
# Convertimos a formato string
#fecha_corte_disponibilidad = ''.join(fecha_corte_disponibilidad) 
# Convertimos a formato numérico
#fecha_corte_disponibilidad_date = datetime.strptime(fecha_corte_disponibilidad, '%Y%m%d').date()
#mes_disponibilidad = fecha_corte_disponibilidad_date.month
# Damos estilo
#fecha_corte_disponibilidad_date = fecha_corte_disponibilidad_date.strftime("%d %b %Y")

#data_intervenciones = pd.read_excel(proyecto / f"input/intervenciones_pedagogicas/Disponibilidad_Presupuestal_{fecha_corte_disponibilidad}interv.xlsx")
#data_intervenciones = clean_names(data_intervenciones) # Normalizamos nombres


lista_regiones = ["AMAZONAS", "ANCASH", "APURIMAC", "AREQUIPA", "AYACUCHO", "CAJAMARCA", "CUSCO", "HUANCAVELICA", "HUANUCO", "ICA", "JUNIN", "LA LIBERTAD", "LAMBAYEQUE", "LORETO", "MADRE DE DIOS", "MOQUEGUA", "PASCO", "PIURA", "PUNO", "SAN MARTIN", "TACNA", "TUMBES", "UCAYALI", "LIMA PROVINCIAS", "CALLAO"]

for region in lista_regiones:

    # Obtener PIM
    #region_seleccionada = data_intervenciones['region'] == region #Seleccionar region
    #tabla_intervenciones = data_intervenciones[region_seleccionada]   
    #pia_intervenciones_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["pia"].sum()))
    
    # Gráfico
    
    
    ### DOC en Python

    document = docx.Document(proyecto / "input/otros/Formato.docx") # Creación del documento en base al template
    #Título
    title=document.add_heading('AM Temas Presupuestales', 0)
    run = title.add_run()
    run.add_break()
    title.add_run('Región ')
    title.add_run(region)
    run = title.add_run()
    run.add_break()
    title.add_run(datetime.today().strftime('%d-%m-%y'))

    document.add_heading('1. Financiamiento de plazas 2021', level=1)
    pa1=document.add_paragraph('A través del DS 078-2021-EF, se financiaron ', style='List Bullet')
    pa1.add_run(' plazas de docentes de aula en el marco de los resultados del proceso de racionalización 2020 \
    en servicios educativos públicos de la región ') 
    pa1.add_run(region) 
    pa1.add_run(' con la siguiente distribución por UGEL:') 
    
    document.save(nueva_carpeta / f'AM_testeo_{region}_{fecha_actual}.docx')

