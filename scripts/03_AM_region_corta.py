#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Sep 21 15:54:48 2021
"""
import docx
import pandas as pd
import numpy as np
import nums_from_string
import os
import getpass
import glob
import matplotlib.pyplot as plt
from datetime import datetime
#from pyprojroot import here
from janitor import clean_names # pip install pyjanitor
from pathlib import Path
from docx.shared import Pt
from docx.shared import Inches

###############################################################################
# Rutas de los archivos #
###############################################################################

if getpass.getuser() == "analistaup18": # PC Analista UP 18 Minedu
    github = Path("C:/Users/ANALISTAUP18/Documents/GitHub/AM-python-docx")
    proyecto = Path("B:/OneDrive - Ministerio de Educación/unidad_B/2021/4. Herramientas de Seguimiento/13.AM_automatizada")
elif  getpass.getuser() == "bran": # PC Brandon
    github = Path("/Users/bran/Documents/GitHub/AM-python-docx")
    proyecto = Path("/Users/bran/Documents/GitHub/AM-python-docx")

###############################################################################
# Fechas de corte #
###############################################################################

# Importamos los nombres de los archivos en la carpeta input
lista_archivos = os.listdir(Path(proyecto, "input", "intervenciones_pedagogicas"))

# Fecha de hoy
fecha_actual = datetime.today().strftime('%d-%m-%y')

## A) Base disponibilidad
# Importamos los nombres de los archivos en la carpeta intervenciones pedagogicas
lista_archivos_int = glob.glob(os.path.join(proyecto,"input/intervenciones_pedagogicas/*"))
#Mantenemos el corte de disponibilidad más reciente
fecha_corte_disponibilidad = max(lista_archivos_int, key=os.path.getctime)
# Nos quedamos con el nombre de archivo para la base de disponibilidad
fecha_corte_disponibilidad = os.path.split(fecha_corte_disponibilidad)
fecha_corte_disponibilidad = fecha_corte_disponibilidad[1]
# Extraemos la fecha del nombre de archivo
fecha_corte_disponibilidad = nums_from_string.get_numeric_string_tokens(fecha_corte_disponibilidad)
# Convertimos a formato string
fecha_corte_disponibilidad = ''.join(fecha_corte_disponibilidad) 
# Convertimos a formato numérico
fecha_corte_disponibilidad_date = datetime.strptime(fecha_corte_disponibilidad, '%Y%m%d').date()
mes_disponibilidad = fecha_corte_disponibilidad_date.month
# Damos estilo
fecha_corte_disponibilidad_date = fecha_corte_disponibilidad_date.strftime("%d %b %Y")

## B) Siaf de mascarillas
fecha_corte_mascarillas = "03 Oct 2021"

# C) Compromisos de desempeño
fecha_corte_compromisos = "21 Sep 2021"



###############################################################################
# Creación de carpeta donde se guardan los outputs #
###############################################################################

# Creación de carpeta
dir = os.path.join(proyecto, f"output/AM_corta_region/AM_{fecha_actual}")
if not os.path.exists(dir):
    os.mkdir(dir)
    print("Se creó una nueva carpeta")
else:
    print("Ya existe la carpeta")
        
# Path de nueva carpeta
nueva_carpeta = Path(proyecto/ f"output/AM_corta_region/AM_{fecha_actual}")

###############################################################################
# Transformación de Datasets #
###############################################################################

##########################
# Base de disponibilidad 2021 #
##########################

# Base de datos región
## Cargamos nombres de regiones
nombre_regiones = pd.read_excel(proyecto / "input/otros/nombre_regiones.xlsx")

# A) Base de disponibilidad
## Cargamos base de disponibilidad
data_intervenciones = pd.read_excel(proyecto / f"input/intervenciones_pedagogicas/Disponibilidad_Presupuestal_{fecha_corte_disponibilidad}interv.xlsx")
data_intervenciones = clean_names(data_intervenciones) # Normalizamos nombres

# Eliminamos filas de "No hay Intervenciones pedagogicas"
data_intervenciones = data_intervenciones[data_intervenciones['intervencion_pedagogica'] != "No hay Intervenciones Pedagógicas"]
# Eliminamos COAR
data_intervenciones = data_intervenciones[data_intervenciones['intervencion_pedagogica'] != "COAR"]
# Eliminamos  Vacaciones Truncas 
data_intervenciones = data_intervenciones[data_intervenciones['especifica_de_gasto'] != "3.2.8.1.5. VACACIONES TRUNCAS DE C.A.S."]

# Mantenemos variables de interés (PIM, DEVENGADO, COMPROMETIDO CERTIFICADO) y 
# colapsamos a nivel de Region, Intervencion Pedagogica y Cas-No-Cas
data_intervenciones_total = data_intervenciones[["region", "cas_no_cas","intervencion_pedagogica","pia", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}", "costo_enero", "costo_febrero", "costo_marzo", "costo_abril", "costo_mayo", "costo_junio", "costo_julio", "costo_agosto", "costo_septiembre", "costo_octubre", "costo_noviembre", "costo_diciembre", "ds_n°_092_2021_ef_transferencia_convivencia", "ds_n°_169_2021_ef_1°_transferencia_intervenciones", "ds_n°_189_2021_ef_1°_transferencia_acompanatic", "ds_n°_209_2021_ef_2°_transferencia_intervenciones" ,"ds_n°_210_2021_ef_2°_transferencia_acompanatic_previo"
]]. \
   groupby(by = ["region", "intervencion_pedagogica"] , as_index=False).sum()

data_intervenciones = data_intervenciones[["region", "cas_no_cas","intervencion_pedagogica", "pia", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}", "costo_enero", "costo_febrero", "costo_marzo", "costo_abril", "costo_mayo", "costo_junio", "costo_julio", "costo_agosto", "costo_septiembre", "costo_octubre", "costo_noviembre", "costo_diciembre", "ds_n°_092_2021_ef_transferencia_convivencia", "ds_n°_169_2021_ef_1°_transferencia_intervenciones", "ds_n°_189_2021_ef_1°_transferencia_acompanatic", "ds_n°_209_2021_ef_2°_transferencia_intervenciones" ,"ds_n°_210_2021_ef_2°_transferencia_acompanatic_previo"
]]. \
   groupby(by = ["region","cas_no_cas", "intervencion_pedagogica"] , as_index=False).sum()

# Transferencias
data_intervenciones['transferencia'] = data_intervenciones["ds_n°_092_2021_ef_transferencia_convivencia"] + data_intervenciones["ds_n°_169_2021_ef_1°_transferencia_intervenciones"] + data_intervenciones["ds_n°_189_2021_ef_1°_transferencia_acompanatic"] + data_intervenciones["ds_n°_209_2021_ef_2°_transferencia_intervenciones"] + data_intervenciones["ds_n°_210_2021_ef_2°_transferencia_acompanatic_previo"]

# Eliminamos filas con 0 PIM
#condicion_elim = (data_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"] != 0) 
#data_intervenciones = data_intervenciones[condicion_elim]

# Calculamos porcentajes
## Avance PIM
data_intervenciones["avance_pim"] = data_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"]/data_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"]

## Avance costo actual
data_intervenciones["costo_actual"] = data_intervenciones["costo_enero"] + data_intervenciones["costo_febrero"] + data_intervenciones["costo_marzo"] + data_intervenciones["costo_abril"] + data_intervenciones["costo_mayo"] + data_intervenciones["costo_junio"] + data_intervenciones["costo_julio"] + data_intervenciones["costo_agosto"] + data_intervenciones["costo_septiembre"] + data_intervenciones["costo_octubre"]
data_intervenciones["avance_costo"] = data_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"]/data_intervenciones["costo_actual"]

# Mantenemos y ordenamos columnas
data_intervenciones = data_intervenciones[['region', "intervencion_pedagogica", "pia", "transferencia", f"pim_reporte_siaf_{fecha_corte_disponibilidad}", f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}", "avance_pim", "costo_actual", "avance_costo"]]

## Reemplazamos inf por NaN
data_intervenciones.replace([np.inf, -np.inf], np.nan, inplace=True)

## Reemplazamos NaN por 0
data_intervenciones['avance_pim'] = data_intervenciones['avance_pim'].fillna("0").astype(float)
data_intervenciones['avance_costo'] = data_intervenciones['avance_costo'].fillna("0").astype(float)

# Tabla intervenciones CAS
#data_intervenciones_cas = data_intervenciones[data_intervenciones['cas_no_cas'] != "NO CAS"]
# Tabla intervenciones NO CAS
#data_intervenciones_nocas = data_intervenciones[data_intervenciones['cas_no_cas'] != "CAS"]

##########################
# Base de disponibilidad 2020 #
##########################
data_intervenciones_2020 = pd.read_excel(proyecto / "input/am_corta/6. Disponibilidad_Presupuestal_20201231.xlsx", sheet_name="Sheet1")

#######################
# Siaf de mascarillas #
#######################

## Cargamos la base insumo de mascarillas
data_mascarillas = pd.read_excel(proyecto / "input/mascarillas/Incorporación_DU_SIAF_20211108.xlsx", sheet_name='Sheet1')
data_mascarillas = clean_names(data_mascarillas) # Normalizamos nombres

# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
data_mascarillas = data_mascarillas[["region","nom_ue","transferencia","pim","certificado","comprometido_anual","devengado"]]. \
   groupby(by = ["region", "nom_ue"], as_index=False).sum()

data_mascarillas["region"] = data_mascarillas["region"].str.split(". ", n=1).apply(lambda l: "".join(l[1]))

data_mascarillas["UNIDAD EJECUTORA"]=data_mascarillas["nom_ue"]
data_mascarillas["RECURSOS TRANSF. (*)"]=data_mascarillas["transferencia"]
data_mascarillas["PIM"]=data_mascarillas["pim"]
data_mascarillas["CERT. (%)"]=data_mascarillas["certificado"]/data_mascarillas["pim"]
data_mascarillas["COMPRO. (%)"]=data_mascarillas["comprometido_anual"]/data_mascarillas["pim"]
data_mascarillas["DEVENGADO (%)"]=data_mascarillas["devengado"]/data_mascarillas["pim"]

############################
# Compromisos de desempeño #
############################

## Cargamos data de compromisos de desempeño
data_cdd = pd.read_excel(proyecto / "input/compromisos_desempeno/regiones_BD_CDD.xlsx")
data_cdd = clean_names(data_cdd) # Normalizamos nombres
data_cdd["pliego"] = data_cdd["pliego"].str.split(". ", n=1, expand = True)
data_cdd['pliego'] = data_cdd['pliego'].astype('int64') # Convertimos ubigeo a integer

# Corregimos genericas
data_cdd["generica"] = data_cdd["generica"].replace("3. BIENES Y SERVICIOS", "2.3. BIENES Y SERVICIOS")
data_cdd["generica"] = data_cdd["generica"].replace("6. ADQUISICION DE ACTIVOS NO FINANCIEROS", "2.6. ADQUISICION DE ACTIVOS NO FINANCIEROS")	

# Hacemos merge con base de datos región
data_cdd = data_cdd.merge(right = nombre_regiones, how="left", on = "pliego")

# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
data_cdd = data_cdd[["region", "unidad_ejecutora", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region", "programa_presupuestal", "generica"], as_index=False).sum()

######################################################
# Sobre el financiamiento de conceptos remunerativos #
######################################################

# C) Base de Encargaturas
df_consolidado_enc = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'ENC-CONSOLIDADO-VF') 
df_consolidado_enc.fillna(0, inplace =  True)
df_consolidado_enc['COSTO'] = df_consolidado_enc['COSTO-TRAMO I']
df_consolidado_enc['PROGRAMADO POR MINEDU'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']+ df_consolidado_enc['NM ENCARGATURAS']
df_consolidado_enc.rename(columns={'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                   'TRANSFERENCIA DS 217':'TRANSFERENCIA POR DS N° 217-2021-EF',
                                   'UNIDADEJECUTORA':'UNIDAD EJECUTORA'},inplace=True)
df_consolidado_enc['APMeINCREMENTOS'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']

tabla_encargaturas = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIA POR DS N° 217-2021-EF',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]

tabla_encargaturas_resumen = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIA POR DS N° 217-2021-EF',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]



# D) Base de Asignaciones Temporales
df_consolidado_at = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'AT-CONSOLIDADO-VF')   
df_consolidado_at.fillna(0, inplace =  True)      
df_consolidado_at.rename(columns={'REGIÓN':'REGION',
                                  'COSTO-TRAMO I':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                  'TRANSFERENCIA DS 187':'TRANSFERENCIA POR DS N° 187-2021-EF'
    },inplace=True)
df_at=df_consolidado_at[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU','TRANSFERENCIA POR DS N° 187-2021-EF']]

# E) Base de Beneficios Sociales
df_consolidado_bf = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'BS-CONSOLIDADO-VF')           
df_consolidado_bf.fillna(0,inplace = True)
df_consolidado_bf['COSTO BENEFICIARIOS 2020 Y 2021'] = df_consolidado_bf['LISTAS-2021'] + df_consolidado_bf['TRAMO I-BS 2020'] + df_consolidado_bf['TRAMO II-BS 2021']
df_consolidado_bf.rename(columns={'REGIÓN':'REGION',
                                  'COSTO BENEFICIARIOS 2020 Y 2021':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'TRANSFERENCIA DS 072-BS 2020':'TRANSFERENCIA POR DS N° 072-2021-EF',
                                  'TRANSFERENCIA DS 256-BS 2021':'TRANSFERENCIA POR DS N° 256-2021-EF'
    },inplace=True)
df_bs = df_consolidado_bf[['REGION','UNIDAD EJECUTORA','COSTO',
                           'PROGRAMADO POR MINEDU',
                           'TRANSFERENCIA POR DS N° 072-2021-EF',
                           'TRANSFERENCIA POR DS N° 256-2021-EF']]

df_transferencia = pd.read_excel(proyecto / 'input/normas_transferencias/TRANSFERENCIAS 2021.xlsx',sheet_name = 'TRANSFERENCIAS')           
df_transferencia.fillna(0,inplace = True)
df_transferencia = clean_names(df_transferencia)
df_transferencia = df_transferencia[['region', 'norma_de_transferencia', 'concepto', 'monto_transferido']].\
groupby(by = ["region", 'norma_de_transferencia', 'concepto'] , as_index=False).sum()


normas = ["DECRETO DE URGENCIA N 065-2021", "DECRETO SUPREMO N 044-2021-EF", "DECRETO SUPREMO N 078-2021-EF"]


df_transferencia = df_transferencia.loc[df_transferencia['norma_de_transferencia'].isin(normas)]
df_transferencia["concepto"].replace({"CONTRATACIÓN MINDEF": "Contratación de plazas docentes en instituciones educativas de educación básica del Ministerio de Defensa"}, inplace=True)


######################################################
# Sobre el proceso de racionalización #
######################################################

data_creacion = pd.read_excel(proyecto / "input/creaciones/plazas_creacion_racio_2021.xlsx", sheet_name="BD",  skiprows = 2)
data_creacion = clean_names(data_creacion) # Normalizamos nombres
data_creacion = data_creacion[['d_region', 'd_dreugel', 'nivel', 'creacion_total']].\
groupby(by = ["d_region", 'd_dreugel', 'nivel'] , as_index=False).sum()
data_creacion['d_region'] = data_creacion['d_region'].str.split(r'DRE ').str[-1]
data_creacion.loc[data_creacion['d_region']=="DE DIOS", 'd_region'] = "MADRE DE DIOS"


data_creacion = data_creacion.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})
data_creacion.loc[data_creacion['nivel']=="inicial", 'inicial'] = data_creacion['creacion_total']
data_creacion.loc[data_creacion['nivel']=="primaria", 'primaria'] = data_creacion['creacion_total']
data_creacion.loc[data_creacion['nivel']=="secundaria", 'secundaria'] = data_creacion['creacion_total']

data_creacion = data_creacion[['region', 'ugel','inicial', 'primaria', 'secundaria', 'creacion_total']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()

data_creacion.fillna(0, inplace =  True)

## Creación plazas docentes -PEM 2021

data_creacion_pem = pd.read_excel(proyecto / "input/creaciones/plazas_creacion_pem_2021.xlsx")
data_creacion_pem = clean_names(data_creacion_pem) # Normalizamos nombres
data_creacion_pem = data_creacion_pem[['d_region', 'd_dreugel', 'modalidad', 'req_doc', 'req_bolsa', 'req_director', 'req_subdir']].\
groupby(by = ["d_region", 'd_dreugel', 'modalidad'] , as_index=False).sum()
data_creacion_pem['d_region'] = data_creacion_pem['d_region'].str.split(r'DRE ').str[-1]
data_creacion_pem.loc[data_creacion_pem['d_region']=="DE DIOS", 'd_region'] = "MADRE DE DIOS"

data_creacion_pem = data_creacion_pem.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})
data_creacion_pem.fillna(0, inplace =  True)

filtro_ebr = data_creacion_pem['modalidad']=="EBR"
creacion_ebr_pem = data_creacion_pem[filtro_ebr]
filtro_ebe = data_creacion_pem['modalidad']=="EBE"
creacion_ebe_pem = data_creacion_pem[filtro_ebe]

#----------------------------------------------------------------------#
## Brecha de plazas docentes
data_brecha = pd.read_excel(proyecto / "input/brecha/brecha_ugel_2020.xlsx", sheet_name="Data")
# Normalizamos nombres

# Mantenemos variables de interés
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()
data_brecha['doc_e_n_cub_req'] = data_brecha['doc_e_n'] - data_brecha['nom_exd_mov1']
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']]
data_brecha.loc[data_brecha['brecha_net']<=0, 'req_neto'] = -1*data_brecha['brecha_net']
data_brecha.loc[data_brecha['brecha_net']>0, 'exc_neto'] = data_brecha['brecha_net']

#Cantidad de UGEL con requerimiento neto
data_brecha.loc[data_brecha['brecha_net']<0, 'cant_ugel_req'] = 1
data_brecha.loc[data_brecha['brecha_net']>0, 'cant_ugel_exc'] = 1

data_brecha_regional = data_brecha[['region', 'brecha_net', 'cant_ugel_req', 'cant_ugel_exc']].groupby(by = ["region"] , as_index=False).sum()
data_brecha_regional.loc[data_brecha_regional['brecha_net']<=0, 'brecha_net'] = -1*data_brecha_regional['brecha_net']
data_brecha_regional.loc[data_brecha_regional['brecha_net']>0, 'brecha_net'] = data_brecha_regional['brecha_net']
data_brecha_regional.fillna(0, inplace =  True)


data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'req_neto', 'exc_neto', 'brecha_net']]
data_brecha.fillna(0, inplace =  True)


#----------------------------------------------------------------------#
## Bloqueo de plazas
data_bloqueo = pd.read_excel(proyecto / "input/bloqueo/plazas_bloqueo_2020.xlsx")
data_bloqueo = clean_names(data_bloqueo) # Normalizamos nombres
data_bloqueo.fillna(0, inplace =  True)

# Mantenemos variables de interés
data_bloqueo['cant_bloqueos'] = 1
data_bloqueo = data_bloqueo[["descreg", 'cant_bloqueos']].groupby(by = ["descreg"] , as_index=False).sum()
data_bloqueo = data_bloqueo.rename(columns={'descreg':'region'})

#----------------------------------------------------------------------#
## Deuda social
#data_deuda_social = pd.read_excel(proyecto / "input/deudas_sociales/deudas_sociales.xlsx")
#data_deuda_social = clean_names(data_deuda_social) # Normalizamos nombres

###############################################################################
# Gráficos
###############################################################################


###############################################################################
# Nuevos cálculos
###############################################################################

data_inversiones = pd.read_excel(proyecto / "input/am_corta/01. Transferencias inversiones.xlsx", sheet_name="Análisis2", skiprows=4)
data_inversiones = clean_names(data_inversiones) # Normalizamos nombres

#monto_pia_2020 = "monto"
#monto_pim_2020 = "monto"
#monto_devengado_2020 = "monto"
#avance_2020 = "monto"
#monto_transferido_2020 = "monto"

data_kit_lavamanos = pd.read_excel(proyecto / "input/am_corta/02. Kit de higiene y lavamanos.xlsx", sheet_name="Análisis", skiprows=5)
data_kit_lavamanos = clean_names(data_kit_lavamanos) # Normalizamos nombres

#data_cdd_2021 = pd.read_excel(proyecto / "input/am_corta/03. CDD_20211109.xlsx", sheet_name="Sheet1")
#data_cdd_2021 = clean_names(data_cdd_2021) # Normalizamos nombres

data_cdd_2021 = pd.read_excel(proyecto / "input/am_corta/02. CDD_AM Plantilla.xlsx", sheet_name="02. CDD", skiprows=5)
data_cdd_2021 = clean_names(data_cdd_2021) # Normalizamos nombres


# Hacemos merge con base de datos región
#data_cdd_2021 = data_cdd_2021.merge(right = nombre_regiones, how="left", on = "cod_pliego")

# Mantenemos variables de interés (region,  pim, devengado y colapsamos a nivel de Region)
#data_cdd_2021 = data_cdd_2021[["region", "pim", "devengado"]]. \
#    groupby(by = ["region"], as_index=False).sum()

# Asignaciones temporales
data_asignaciones_2021 = pd.read_excel(proyecto / "input/am_corta/08. Asignaciones Temporales.xlsx", sheet_name="Sheet2")


#Deuda social

data_deuda = pd.read_excel(proyecto / "input/am_corta/00b_Deudas sociales.xlsx", sheet_name="Base", skiprows=1)
data_deuda = clean_names(data_deuda) # Normalizamos nombres

# Plazas
data_plazas = pd.read_excel(proyecto / "input/am_corta/10. Plazas financiadas.xlsx", sheet_name="P_Financiadas", skiprows=3)

# Conceptos remunerativos
data_conceptos_remunerativos_2021 = pd.read_excel(proyecto / "input/am_corta/00a_Conceptos remunerativos consolid.xlsx", sheet_name="Base agregada", skiprows=3)
data_conceptos_remunerativos_2021 = clean_names(data_conceptos_remunerativos_2021) # Normalizamos nombres

## Beneficios sociales
data_beneficios_sociales_2021 = pd.read_excel(proyecto / "input/am_corta/09. Beneficios sociales.xlsx", sheet_name="TD_BS", skiprows=2)
data_beneficios_sociales_2021 = clean_names(data_beneficios_sociales_2021)

###############################################################################
# Creación del documento en docx 
###############################################################################

# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS", "ANCASH", "APURIMAC", "AREQUIPA", "AYACUCHO", "CAJAMARCA", "CUSCO", "HUANCAVELICA", "HUANUCO", "ICA", "JUNIN", "LA LIBERTAD", "LAMBAYEQUE", "LORETO", "MADRE DE DIOS", "MOQUEGUA", "PASCO", "PIURA", "PUNO", "SAN MARTIN", "TACNA", "TUMBES", "UCAYALI", "LIMA PROVINCIAS", "CALLAO"]

# For loop para cada región
for region in lista_regiones:

    ###########################################################################
    # Construcción de tablas e indicadores #
    ###########################################################################

    ############################################
    # Tablas e indicadores beneficios sociales #
    ############################################
    
    region_seleccionada = data_beneficios_sociales_2021['region'] == region
    data_beneficios_region = data_beneficios_sociales_2021[region_seleccionada]
    
    costo_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['costo_beneficios']))
    pia_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['pia_beneficios']))
    ds_072_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['transferencia_072']))
    ds_256_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['transferencia_256']))
    
    costo_beneficios_c = data_beneficios_region.iloc[0]['costo_beneficios']
    pia_beneficios_c = data_beneficios_region.iloc[0]['pia_beneficios']
    ds_072_beneficios_c = data_beneficios_region.iloc[0]['transferencia_072']
    ds_256_beneficios_c = data_beneficios_region.iloc[0]['transferencia_256']

    ############################################
    # Tablas e indicadores conceptos remunerativos #
    ############################################
    
    region_seleccionada = data_conceptos_remunerativos_2021['region'] == region

    data_remuneracion_region = data_conceptos_remunerativos_2021[region_seleccionada]
    
    conceptos_remunerativos_2021 =  str('{:,.0f}'.format(data_remuneracion_region.iloc[0]['conceptos_remunerativos']))
    conceptos_remunerativos_2021_c =  data_remuneracion_region.iloc[0]['conceptos_remunerativos']
    
    ############################################
    # Tablas e indicadores asignaciones temporales #
    ############################################
    
    region_seleccionada = data_asignaciones_2021['region'] == region
    
    data_asignacion_region = data_asignaciones_2021[region_seleccionada]
    
    pia_asignacion = str('{:,.0f}'.format(data_asignacion_region.iloc[0]['pia_asignacion']))
    
    pim_asignacion = str('{:,.0f}'.format(data_asignacion_region.iloc[0]['pim_asignacion']))
    
    transferencia_asignacion = str('{:,.0f}'.format(data_asignacion_region.iloc[0]['transferencia_asignacion']))

    transferencia_asignacion_c = data_asignacion_region.iloc[0]['transferencia_asignacion']

    ############################################
    # Tablas e indicadores Plazas #
    ############################################
    
    region_seleccionada = data_plazas['region'] == region #Seleccionar region
    
    data_plazas_region = data_plazas[region_seleccionada]
    
    monto_plazas = str('{:,.0f}'.format(data_plazas_region.iloc[0]['monto_plazas']))
    monto_plazas_c = data_plazas_region.iloc[0]['monto_plazas']
    
    cantidad_plazas = str('{:,.0f}'.format(data_plazas_region.iloc[0]['cantidad_plazas']))

    ############################################
    # Tablas e indicadores Deuda Social #
    ############################################
    
 
    region_seleccionada = data_deuda['region'] == region #Seleccionar region
    
    data_deuda_region = data_deuda[region_seleccionada]
    
    monto_deuda_c = data_deuda_region.iloc[0]['﻿monto_deuda_social']
    monto_deuda = str('{:,.0f}'.format(data_deuda_region.iloc[0]['﻿monto_deuda_social']))

    ############################################
    # Tablas e indicadores cdd_2021 #
    ############################################
    
    region_seleccionada = data_cdd_2021['region'] == region #Seleccionar region
    
    data_transferencia_cdd_2021 = data_cdd_2021[region_seleccionada]
    monto_transferencia_cdd_2021_c = data_transferencia_cdd_2021.iloc[0]['recursos_transferidos_'] # Valor numérico
    monto_transferencia_cdd_2021 = str('{:,.0f}'.format(data_transferencia_cdd_2021.iloc[0]['recursos_transferidos_'])) # Valor con comas (string)
    cdd_ejecucion_2021 = str('{:,.1%}'.format(data_transferencia_cdd_2021["devengado"].sum()/data_transferencia_cdd_2021["pim"].sum()))
    
    ############################################
    # Tablas e indicadores kit/lavamanos #
    ############################################

    region_seleccionada = data_kit_lavamanos['region'] == region #Seleccionar region

    data_region_kit = data_kit_lavamanos[region_seleccionada]
    monto_region_kit = str('{:,.0f}'.format(data_region_kit.iloc[0]['kit_de_higiene_transferencia']))
    monto_region_kit_c = data_region_kit.iloc[0]['kit_de_higiene_transferencia']
    
    data_region_lavamanos = data_kit_lavamanos[region_seleccionada]
    monto_region_lavamanos = str('{:,.0f}'.format(data_region_lavamanos.iloc[0]['lavamanos_transferencia']))
    monto_region_lavamanos_c = data_region_lavamanos.iloc[0]['lavamanos_transferencia']
    
    avance_kit = data_kit_lavamanos[region_seleccionada]
    avance_kit = str('{:,.0f}'.format(avance_kit.iloc[0]['kit_de_higiene_declaracion']))

    avance_lavamanos = data_kit_lavamanos[region_seleccionada]
    avance_lavamanos = str('{:,.0f}'.format(avance_lavamanos.iloc[0]['lavamanos_declaracion']))

    
    ############################################
    # Tablas e indicadores base inversiones #
    ############################################
    
    region_seleccionada = data_inversiones['region'] == region #Seleccionar region
    
    data_region_inversiones = data_inversiones[region_seleccionada]
    monto_region_inversiones_c = data_region_inversiones.iloc[0]['monto_inversiones'] # Valor numérico
    monto_region_inversiones = str('{:,.0f}'.format(data_region_inversiones.iloc[0]['monto_inversiones']))
    
    fila1 = data_inversiones[region_seleccionada]
    fila1 = str(fila1.iloc[0]['texto_1'])

    fila2 = data_inversiones[region_seleccionada]
    fila2 = str(fila2.iloc[0]['texto_2']) 
    
    ############################################
    # Tablas e indicadores base disponibilidad #
    ############################################

    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_intervenciones['region'] == region #Seleccionar region
    tabla_intervenciones = data_intervenciones[region_seleccionada]   
    pia_intervenciones_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["pia"].sum()))
    pim_intervenciones_region_2021 = str('{:,.0f}'.format(tabla_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    ejecucion_intervenciones_region_2021 = str('{:,.0f}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    porcentaje_ejecucion_2021 = str('{:,.1%}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()))
    porcentaje_costomesactual = str('{:,.1%}'.format(tabla_intervenciones[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones["costo_actual"].sum()))
    transferencia_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["transferencia"].sum()))
    transferencia_region_2021_c = tabla_intervenciones["transferencia"].sum()

    
    # TOTAL
    tabla_intervenciones_formato = data_intervenciones[region_seleccionada]
    
    # Generamos porcentaje de avance
    porcentaje_ejecucion_a = tabla_intervenciones_formato[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()
    porcentaje_costomesactual_a = tabla_intervenciones_formato[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato["costo_actual"].sum()
    
    # Generamos fila total
    total_int = tabla_intervenciones_formato.groupby(by = ["region"], as_index=False).sum()

    # Realizamos append del total en la tabla
    tabla_intervenciones_formato = tabla_intervenciones_formato.append(total_int, ignore_index=True)

    # Reemplazamos % de avance pim y avance costo por los valores correctos en fila total
    tabla_intervenciones_formato.iloc[-1, tabla_intervenciones_formato.columns.get_loc('avance_pim')] = porcentaje_ejecucion_a
    tabla_intervenciones_formato.iloc[-1, tabla_intervenciones_formato.columns.get_loc('avance_costo')] = porcentaje_costomesactual_a

    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_intervenciones_formato['intervencion_pedagogica'] = tabla_intervenciones_formato['intervencion_pedagogica'].fillna("Total")
    tabla_intervenciones_formato['avance_pim'] = tabla_intervenciones_formato['avance_pim'].fillna("0").astype(float)
    tabla_intervenciones_formato['avance_costo'] = tabla_intervenciones_formato['avance_costo'].fillna("0").astype(float)
    

    # CAS
    #tabla_intervenciones_formato_cas = data_intervenciones_cas[region_seleccionada]
    
    # Generamos porcentaje de avance CAS
    #porcentaje_ejecucion_cas = tabla_intervenciones_formato_cas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_cas[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()
    #porcentaje_costomesactual_cas = tabla_intervenciones_formato_cas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_cas["costo_actual"].sum()
    
    # Generamos fila total
    #total_int_cas = tabla_intervenciones_formato_cas.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_cas = tabla_intervenciones_formato_cas.append(total_int_cas, ignore_index=True)
    
    # Reemplazamos % de avance pim y avance costo por los valores correctos en fila total
    #tabla_intervenciones_formato_cas.iloc[-1, tabla_intervenciones_formato_cas.columns.get_loc('avance_pim')] = porcentaje_ejecucion_cas
    #tabla_intervenciones_formato_cas.iloc[-1, tabla_intervenciones_formato_cas.columns.get_loc('avance_costo')] = porcentaje_costomesactual_cas
    
    #Incluimos palabra "total" y "-" en vez de NaN
    #tabla_intervenciones_formato_cas['intervencion_pedagogica'] = tabla_intervenciones_formato_cas['intervencion_pedagogica'].fillna("Total")
    #tabla_intervenciones_formato_cas['avance_pim'] = tabla_intervenciones_formato_cas['avance_pim'].fillna("0").astype(float)
    #tabla_intervenciones_formato_cas['avance_costo'] = tabla_intervenciones_formato_cas['avance_costo'].fillna("0").astype(float)
    
    # NO CAS
    #tabla_intervenciones_formato_nocas = data_intervenciones_nocas[region_seleccionada]
    
    # Generamos porcentaje de avance NO CAS
    #porcentaje_ejecucion_nocas = tabla_intervenciones_formato_nocas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_nocas[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"].sum()
    #porcentaje_costomesactual_nocas = tabla_intervenciones_formato_nocas[f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato_nocas["costo_actual"].sum()
    
    # Generamos fila total
    #total_int_nocas = tabla_intervenciones_formato_nocas.groupby(by = ["region"], as_index=False).sum()
    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_nocas = tabla_intervenciones_formato_nocas.append(total_int_nocas, ignore_index=True)
    
    # Reemplazamos % de avance pim y avance costo por los valores correctos en fila total
    #tabla_intervenciones_formato_nocas.iloc[-1, tabla_intervenciones_formato_nocas.columns.get_loc('avance_pim')] = porcentaje_ejecucion_nocas
    #tabla_intervenciones_formato_nocas.iloc[-1, tabla_intervenciones_formato_nocas.columns.get_loc('avance_costo')] = porcentaje_costomesactual_nocas
    
    #Incluimos palabra "total" y "-" en vez de NaN
    #tabla_intervenciones_formato_nocas['intervencion_pedagogica'] = tabla_intervenciones_formato_nocas['intervencion_pedagogica'].fillna("Total")
    #tabla_intervenciones_formato_nocas['avance_pim'] = tabla_intervenciones_formato_nocas['avance_pim'].fillna("0").astype(float)
    #tabla_intervenciones_formato_nocas['avance_costo'] = tabla_intervenciones_formato_nocas['avance_costo'].fillna("0").astype(float)
    
    # Formato para la tabla
    formato_tabla_intervenciones = {
        "intervencion_pedagogica" : "{}",
        f"pim_reporte_siaf_{fecha_corte_disponibilidad}": "{:,.0f}",
        f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}" : "{:,.0f}",
        f"presupuesto_devengado_reporte_siaf_{fecha_corte_disponibilidad}" : "{:,.0f}",
    #    "avance_pim": "{:,.1%}",
    #    "costo_actual": "{:,.0f}",
        "avance_costo": "{:,.1%}",
        }
    #tabla_intervenciones_formato_cas = tabla_intervenciones_formato_cas.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})            
    #tabla_intervenciones_formato_nocas = tabla_intervenciones_formato_nocas.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})            

    tabla_intervenciones_formato = tabla_intervenciones_formato.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})            
    
    ############################################
    # Tablas e indicadores mascarillas #
    ############################################

    # Generamos la tabla "tabla1_mascarilla" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = data_mascarillas['region'] == region
    tabla_mascarillas = data_mascarillas[region_seleccionada]
    # Generamos los indicadores de PIM y ejecución de intervenciones
    transferencia_mascarilla = str('{:,.0f}'.format(tabla_mascarillas["transferencia"].sum()))
    transferencia_mascarilla_c = tabla_mascarillas["transferencia"].sum()
    transferencia_mascarilla_millones = str('{:,.1f}'.format(tabla_mascarillas["transferencia"].sum()/1000000))
    devengado_mascarillas=str('{:.1%}'.format(tabla_mascarillas["devengado"].sum()/tabla_mascarillas["transferencia"].sum()))
    # Generamos la tabla "tabla_mascarillas_formato" - mantiene la región i de la lista de regiones
    tabla_mascarillas_formato = data_mascarillas[region_seleccionada]
    # Formato para la tabla
    formato_tabla_mascarillas = {
        "UNIDAD EJECUTORA": "{}",
        "RECURSOS TRANSF. (*)": "{:,.0f}",
        "PIM" : "{:,.0f}",
        "CERT. (%)" : "{:.1%}",
        "COMPRO. (%)": "{:.1%}",
        "DEVENGADO (%)": "{:.1%}",  
        }
    tabla_mascarillas_formato = tabla_mascarillas_formato.transform({k: v.format for k, v in formato_tabla_mascarillas.items()})  
    
    #################################################
    # Tablas e indicadores compromisos de desempeño #
    #################################################

    # Generamos la tabla "tabla_cdd" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = data_cdd['region'] == region
    tabla_cdd = data_cdd[region_seleccionada]
    
    # Generamos fila total
    total_cdd = tabla_cdd[["region", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_cdd = tabla_cdd.append(total_cdd, ignore_index=True)
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_cdd['programa_presupuestal'] = tabla_cdd['programa_presupuestal'].fillna("Total")
    tabla_cdd['generica'] = tabla_cdd['generica'].fillna("-")

     # Generamos la tabla con formato
    formato_tabla_cdd = {
        "programa_presupuestal": "{}",
        "generica": "{}",
        "monto": "{:,.0f}",
        "ds_085_2021_ef": "{:,.0f}",
        "ds_218_2021_ef": "{:,.0f}",
        "ds_220_2021_ef": "{:,.0f}",
        }
    tabla_cdd_formato = tabla_cdd
    tabla_cdd_formato = tabla_cdd_formato.transform({k: v.format for k, v in formato_tabla_cdd.items()})            
    
    # Generamos CDD transferido 
    cdd_transferido = str('{:,.0f}'.format(tabla_cdd["monto"].sum()))
   
    ########################
    # Tablas Encargaturas #
    ########################

    region_seleccionada = df_consolidado_enc['REGION'] == region
    tabla1 = tabla_encargaturas[region_seleccionada]
    costo_enc = str('{:,.0f}'.format(tabla1["COSTO"].sum())) 
    apmeincre = str('{:,.0f}'.format(tabla1["APMeINCREMENTOS"].sum()))
    prog_gore = str('{:,.0f}'.format(tabla1["PROGRAMADO POR EL PLIEGO REGIONAL"].sum()))
    nm_enca = str('{:,.0f}'.format(tabla1['NM ENCARGATURAS'].sum()))
    ds_217 = str('{:,.0f}'.format(tabla1['TRANSFERENCIA POR DS N° 217-2021-EF'].sum()))
    ds_217_c = tabla1['TRANSFERENCIA POR DS N° 217-2021-EF'].sum()
    
    tabla_encargaturas_resumen = tabla1.groupby(['UNIDAD EJECUTORA'], as_index=False).sum()
    tabla_encargaturas_resumen = tabla_encargaturas_resumen[['UNIDAD EJECUTORA',
                                                             'COSTO', 
                                                             'PROGRAMADO POR MINEDU', 
                                                             'PROGRAMADO POR EL PLIEGO REGIONAL',
                                                             'TRANSFERENCIA POR DS N° 217-2021-EF']] 
    
    # Generamos fila total
    total_enc = tabla_encargaturas_resumen[["UNIDAD EJECUTORA", "COSTO", \
"PROGRAMADO POR MINEDU", "PROGRAMADO POR EL PLIEGO REGIONAL", "TRANSFERENCIA POR DS N° 217-2021-EF"]].sum()
    
    # Realizamos append del total en la tabla
    tabla_encargaturas_resumen = tabla_encargaturas_resumen.append(total_enc, ignore_index=True)
    
    #Incluimos palabra "total"
    tabla_encargaturas_resumen.iloc[-1, tabla_encargaturas_resumen.columns.get_loc('UNIDAD EJECUTORA')] = "Total"

    tabla_encargaturas_resumen = tabla_encargaturas_resumen.round(2)

    ###################################
    #  Tablas Asignaciones Temporales #
    ###################################
    
    region_seleccionada = df_at['REGION'] == region
    tabla2 = df_at[region_seleccionada]
    costo_at = str('{:,.0f}'.format(tabla2["COSTO"].sum()))
    apm_at = str('{:,.0f}'.format(tabla2["PROGRAMADO POR MINEDU"].sum())) 
    ds_187_at = str('{:,.0f}'.format(tabla2["TRANSFERENCIA POR DS N° 187-2021-EF"].sum())) 
    ds_187_at_c = tabla2["TRANSFERENCIA POR DS N° 187-2021-EF"].sum()
    tabla_at_resumen = tabla2.groupby(['UNIDAD EJECUTORA'], as_index=False).sum() 
    
    # Generamos fila total
    total_at = tabla_at_resumen[["UNIDAD EJECUTORA", "COSTO", \
"PROGRAMADO POR MINEDU", "TRANSFERENCIA POR DS N° 187-2021-EF"]].sum()
    
    # Realizamos append del total en la tabla
    tabla_at_resumen = tabla_at_resumen.append(total_at, ignore_index=True)
    
    #Incluimos palabra "total"
    tabla_at_resumen.iloc[-1, tabla_at_resumen.columns.get_loc('UNIDAD EJECUTORA')] = "Total"

    tabla_at_resumen = tabla_at_resumen.round(2)

    ###################################
    #    Tablas Beneficios Sociales   #
    ###################################
    
    region_seleccionada = df_bs['REGION'] == region
    region_seleccionada2 = df_consolidado_bf['REGION'] == region
    tabla3 = df_bs[region_seleccionada]
    tabla3_2 = df_consolidado_bf[region_seleccionada]
    costo_bs = str('{:,.0f}'.format(tabla3["COSTO"].sum()))
    apm_bs = str('{:,.0f}'.format(tabla3["PROGRAMADO POR MINEDU"].sum()))    
    ds_72_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS N° 072-2021-EF"].sum())) 
    ds_72_bs_c = tabla3["TRANSFERENCIA POR DS N° 072-2021-EF"].sum()
    ds_256_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS N° 256-2021-EF"].sum()))
    tabla_bs_resumen = tabla3.groupby(['UNIDAD EJECUTORA'], as_index=False).sum()

    # Generamos fila total
    total_bs = tabla_bs_resumen[["UNIDAD EJECUTORA", "COSTO", \
"PROGRAMADO POR MINEDU", "TRANSFERENCIA POR DS N° 072-2021-EF", "TRANSFERENCIA POR DS N° 256-2021-EF"]].sum()

    # Realizamos append del total en la tabla
    tabla_bs_resumen = tabla_bs_resumen.append(total_bs, ignore_index=True)
    
    #Incluimos palabra "total"
    tabla_bs_resumen.iloc[-1, tabla_bs_resumen.columns.get_loc('UNIDAD EJECUTORA')] = "Total"

    tabla_bs_resumen = tabla_bs_resumen.round(2)
    
    lista_bf = str('{:,.0f}'.format(tabla3_2["LISTAS-2021"].sum()))
    
    
    ########################################
    #    Tablas Financiamiento de Plazas   #
    ########################################

    region_seleccionada = data_creacion['region'] == region #Seleccionar region
    tabla_creacion = data_creacion[region_seleccionada]
    creacion_region = str('{:,.0f}'.format(tabla_creacion["creacion_total"].sum()))
    creacion_region_c = tabla_creacion["creacion_total"].sum()

    # Generamos fila total
    total_creacion = tabla_creacion[["region", "inicial", "primaria", "secundaria", "creacion_total"]]. \
    groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_creacion = tabla_creacion.append(total_creacion, ignore_index=True)
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_creacion['ugel'] = tabla_creacion['ugel'].fillna("Total")

    tabla_creacion_formato = data_creacion[region_seleccionada]
    
    
    ########################################
    #    Tablas Deudas sociales   #
    ########################################

    #tabla_deuda_social = data_deuda_social
    #tabla_deuda_social_formato = data_deuda_social
    
    #formato_tabla_deuda_social = {
    #    "seccion_pliego": "{}",
    #    "monto" : "{:,.0f}",
    #    }
    #tabla_deuda_social_formato = tabla_deuda_social_formato.transform({k: v.format for k, v in formato_tabla_deuda_social.items()})        
    
    ###################################
    #    Suma de variables   #
    ###################################

    # transferencia_region_2021_c
    # ds_217_c
    # ds_187_at_c
    # ds_72_bs_c
    # monto_region_inversiones_c
    # transferencia_mascarilla_c
    # monto_region_kit_c
    # monto_region_lavamanos_c
    # monto_transferencia_cdd_2021_c
    # monto_plazas
    # monto_deuda_c
    
    total_transferido = str('{:,.0f}'.format(np.sum([transferencia_region_2021_c, conceptos_remunerativos_2021_c, monto_deuda_c, monto_region_inversiones_c, transferencia_mascarilla_c, monto_region_kit_c, monto_region_lavamanos_c, monto_transferencia_cdd_2021_c])))
        
    ############################################
    # Tablas de Carátula #
    ############################################
    
    # Conceptos remunerativos: 
        # Pago de encargaturas - ds_217
        # Pago de asignaciones temporales - ds_187_at - transferencia_asignacion_c
        # Pago de beneficios sociales - ds_72_bs
        # Financiamiento de plazas - monto_plazas_c
    
    # Creamos tabla con fechas de corte
    tabla_1 = (
        ("Intervenciones pedagógicas", transferencia_region_2021),
        ("Conceptos remunerativos", conceptos_remunerativos_2021),
        ("Deuda social", monto_deuda),
        ("Inversiones", monto_region_inversiones)
    )
    
    tabla_2 = (
        ("Mascarillas y protectores faciales", transferencia_mascarilla),
        ("Kit de higiene", monto_region_kit),
        ("Estaciones de lavado de manos", monto_region_lavamanos)
    )
    
    tabla_3 = (
        ("Compromisos de desempeño", monto_transferencia_cdd_2021),
    )
    
    tabla_4 = (
        ("TOTAL TRANSFERIDO", total_transferido),
    )
    
    ###########################################################################
    # Inclusión del texto del documento #
    ###########################################################################
    document = docx.Document(proyecto / "input/otros/Formato.docx") # Creación del documento en base al template
    title=document.add_heading('AYUDA MEMORIA', 0) #Título del documento
    run = title.add_run()
    title.add_run(' DE LA REGIÓN ')
    title.add_run(region)
    run = title.add_run()
    
    ###########################################################################
    # Carátula del documento #
    ###########################################################################

    document.add_picture(f"/Users/bran/Documents/GitHub/AM-python-docx/input/maps/{region}.PNG", width=Inches(2.7))
    
    dt_1 = document.add_paragraph("Financiamientos Generales")
    dt_1.style = document.styles['Heading 6']
    
    dt_1_tabla = document.add_table(rows=1, cols=2)
    dt_1_tabla.style = "formato_tabla_minedu"
    hdr_cells = dt_1_tabla.rows[0].cells
    hdr_cells[0].text = "Concepto"
    hdr_cells[1].text = "Transferencia"
    for id, name in tabla_1:
        row = dt_1_tabla.add_row().cells
        row[0].text = str(id)
        row[1].text = str(name)
    
    dt_2 = document.add_paragraph("Financiamiento para la reapertura")
    dt_2.style = document.styles['Heading 6']
    
    dt_2_tabla = document.add_table(rows=1, cols=2)
    dt_2_tabla.style = "formato_tabla_minedu"
    hdr_cells = dt_2_tabla.rows[0].cells
    hdr_cells[0].text = "Concepto"
    hdr_cells[1].text = "Transferencia"
    for id, name in tabla_2:
        row = dt_2_tabla.add_row().cells
        row[0].text = str(id)
        row[1].text = str(name)

    
    dt_3 = document.add_paragraph("Transferencia de Compromisos de desempeño")
    dt_3.style = document.styles['Heading 6']
    
    dt_3_tabla = document.add_table(rows=1, cols=2)
    dt_3_tabla.style = "formato_tabla_minedu"
    hdr_cells = dt_3_tabla.rows[0].cells
    hdr_cells[0].text = "Concepto"
    hdr_cells[1].text = "Transferencia"
    for id, name in tabla_3:
        row = dt_3_tabla.add_row().cells
        row[0].text = str(id)
        row[1].text = str(name)
    
    
    dt_4_tabla = document.add_table(rows=1, cols=2)
    dt_4_tabla.style = "formato_tabla_minedu"
    hdr_cells = dt_4_tabla.rows[0].cells
    for id, name in tabla_4:
        row = dt_4_tabla.add_row().cells
        row[0].text = str(id)
        row[1].text = str(name)
    
    #####################################################
    # Incluimos sección 1 de intervenciones pedagógicas #
    #####################################################
    
    document.add_heading("Intervenciones pedagógicas", level=2) # 1) Intervenciones pedagógicas  
    document.add_heading("Corte: 09/11/2021", level=3)
#    interv_parrafo1 = document.add_paragraph("Al año 2020, la región ")
#    interv_parrafo1.style = document.styles['Heading 5']
#    interv_parrafo1.add_run(region)
#    interv_parrafo1.add_run(" inicio con un PIA de ")
#    interv_parrafo1.add_run(monto_pia_2020)
#    interv_parrafo1.add_run(". Al finalizar el año contaron con un  ")
#    interv_parrafo1.add_run(monto_pim_2020)
#    interv_parrafo1.add_run(" de los cuales se ejecutaron ")
#    interv_parrafo1.add_run(monto_devengado_2020)
#    interv_parrafo1.add_run(" lo cual corresponde a ")
#    interv_parrafo1.add_run(avance_2020)
#    interv_parrafo1.add_run(" del PIM. En el transcurso del 2020, se realizó una transferencia de partidas, por un monto de ")
#    interv_parrafo1.add_run(monto_transferido_2020)
#    interv_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    
    interv_parrafo2 = document.add_paragraph("Al año 2021, la región ")
    interv_parrafo2.style = document.styles['Heading 5']
    interv_parrafo2.add_run(region)
    interv_parrafo2.add_run(" inició con un PIA de S/")
    interv_parrafo2.add_run(pia_intervenciones_region_2021)
    interv_parrafo2.add_run(". En el transcurso del año fiscal, se han realizado cuatro (04) transferencias para complementar el financiamiento de las intervenciones pedagógicas, por un monto total de S/ ")
    interv_parrafo2.add_run(transferencia_region_2021)
    interv_parrafo2.add_run(". Al 9 de noviembre")
    interv_parrafo2.add_run(" cuentan con S/ ")
    interv_parrafo2.add_run(pim_intervenciones_region_2021)
    interv_parrafo2.add_run(" en su PIM. De los cuales se han ejecutado el ")
    interv_parrafo2.add_run(porcentaje_ejecucion_2021)
    interv_parrafo2.add_run(" (Devengado/PIM)")
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #####################################################
    # Incluimos sección 2 de compromisos de desempeño #
    #####################################################
    
    document.add_heading("Compromisos de desempeño", level=2)
    document.add_heading("Corte: 09/11/2021", level=3)
    cdd_parrafo1 = document.add_paragraph("En el marco de la Norma Técnica para la implementación del mecanismo denominado Compromisos de Desempeño 2021, aprobada por Resolución Ministerial N° 042-2021-MINEDU y modificada por la Resolución Ministerial N° 160-2021-MINEDU, se han realizado transferencias de partidas a favor de las Unidades Ejecutoras de Educación del Gobierno Regional de  ")
    cdd_parrafo1.style = document.styles['Heading 5']
    cdd_parrafo1.add_run(region)
    cdd_parrafo1.add_run(" por la suma de S/ ")
    cdd_parrafo1.add_run(monto_transferencia_cdd_2021)
    cdd_parrafo1.add_run(". Con fecha de corte SIAF, la ejecución a nivel regional de los recursos de compromisos de desempeño fue de  ")
    cdd_parrafo1.add_run(cdd_ejecucion_2021)    
    cdd_parrafo1.add_run(" respecto al PIM.")
    cdd_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        
    
    #####################################################
    # Incluimos sección 3 de mascarillas y protectores #
    #####################################################
    
    document.add_heading("Mascarillas y protectores faciales", level=2)
    document.add_heading("Corte: 08/11/2021", level=3)
    mascarilla_parrafo1 = document.add_paragraph("Mediante el DU N° 021-2021, se transfirió S/ ")
    mascarilla_parrafo1.style = document.styles['Heading 5']
    mascarilla_parrafo1.add_run(transferencia_mascarilla)
    mascarilla_parrafo1.add_run(" a las UE del Gobierno Regional de ")
    mascarilla_parrafo1.add_run(region)
    mascarilla_parrafo1.add_run(" para la adquisición y distribución de mascarillas faciales textiles de uso comunitario para estudiantes y personal, así como protectores faciales para el mencionado personal.")
    mascarilla_parrafo1.add_run(" Al corte se ha ejecutado el ")
    mascarilla_parrafo1.add_run(devengado_mascarillas)
    
    #####################################################
    # Incluimos sección 4 de kit de higiene #
    #####################################################
    
    document.add_heading("Kit de higiene y estaciones de lavado de manos", level=2)
    
    kit_parrafo1 = document.add_paragraph("Con respecto al kit de higiene y al lavamanos, se transfirió S/ ")
    kit_parrafo1.style = document.styles['Heading 5']
    kit_parrafo1.add_run(monto_region_kit)
    kit_parrafo1.add_run(" y S/ ")
    kit_parrafo1.add_run(monto_region_lavamanos)
    kit_parrafo1.add_run(" respectivamente. Estos han sido otorgados a las cuentas de los directores de las instituciones educativas en el marco del artículo 42 de la Ley 31084 y el artículo 2 del DU N° 021-2021. De ello, los directores han declarado la ejecución de S/ ")
    kit_parrafo1.add_run(avance_kit)
    kit_parrafo1.add_run(" y S/ ")
    kit_parrafo1.add_run(avance_lavamanos)
    
    #####################################################
    # Incluimos sección 5 de inversiones #
    #####################################################

    document.add_heading("Inversiones", level=2)
    
    inversiones_parrafo1 = document.add_paragraph("")
    inversiones_parrafo1.style = document.styles['Heading 5']
    inversiones_parrafo1.add_run(fila1)
    inversiones_parrafo1.add_run(fila2)
    
    #####################################################
    # Incluimos sección 5 de pago de encargaturas #
    #####################################################
    
    document.add_heading("Pago de Encargaturas", level=2)
    
    encarg_parrafo1 = document.add_paragraph(
    "Para la región ")
    encarg_parrafo1.add_run(f'{region}, por concepto de encargaturas, ') 
    encarg_parrafo1.style = document.styles['Heading 5']
    encarg_parrafo1.add_run(' se ha calculado para el ')
    encarg_parrafo1.add_run(datetime.today().strftime('%Y'))
    encarg_parrafo1.add_run(' un costo de S/ ')
    encarg_parrafo1.add_run(f'{costo_enc}') #Insertar valor de base de datos
    encarg_parrafo1.add_run('. Con Decreto Supremo N° 217-2021-EF \
publicado el 27 de agosto de 2021 en el marco de lo autorizado en el literal b) \
del numeral 40.1 de la Ley de Presupuesto 2021, se ha realizado una transferencia \
de partidas por el monto de S/.')  # Este párrafo tendrá que variar año tras año
    encarg_parrafo1.add_run(f'{ds_217}') #Insertar valor de base de datos    
    encarg_parrafo1.add_run(' a favor de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo1.add_run(region)
    encarg_parrafo1.add_run(' para financiar el costo diferencial.')
    encarg_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    encarg_parrafo1.add_run(" La segunda transferencia de recursos por concepto de encargaturas, el cual debería aprobarse como máximo el 26 de noviembre del 2021.")
    
    #####################################################
    # Incluimos sección 6 de pago de asignaciones temporales #
    #####################################################
    
    document.add_heading("Pago de asignaciones temporales", level=2)
    
    encarg_parrafo5 = document.add_paragraph('Para la Región ')    
    encarg_parrafo5.style = document.styles['Heading 5']
    encarg_parrafo5.add_run(region)
    encarg_parrafo5.add_run(', por concepto de Asignaciones Temporales \
, se ha calculado para el 2021 un costo de \
S/ ')    
    encarg_parrafo5.add_run(f'{costo_at} ') #Insertar valor de base de datos    
    encarg_parrafo5.add_run(". Con un PIA 2021 por el monto de S/ ")
    encarg_parrafo5.add_run(pia_asignacion)
    encarg_parrafo5.add_run('. Con Decreto Supremo N° 187-2021-EF se ha realizado una transferencia \
de partidas por el monto de S/ ')
    encarg_parrafo5.add_run(transferencia_asignacion) #Insertar valor de base de datos    
    encarg_parrafo5.add_run(". La segunda transferencia de recursos por concepto de asignaciones temporales debería aprobarse como máximo el 26 de noviembre del 2021.")

    #####################################################
    # Incluimos sección 7 de pago de beneficios sociales #
    #####################################################
    document.add_heading("Pago de beneficios sociales", level=2)
    
    encarg_parrafo10 = document.add_paragraph('Para la Región ')
    encarg_parrafo10.style = document.styles['Heading 5']
    encarg_parrafo10.add_run(region)
    encarg_parrafo10.add_run(', por concepto de Beneficios Sociales se ha calculado\
 un costo total de S/. ')
    encarg_parrafo10.add_run(f'{costo_bs}') #Insertar valor de base de datos    
    encarg_parrafo10.add_run(' y se han aprobado pagos hasta por un costo de S/. ')
    encarg_parrafo10.add_run(costo_beneficios) #Insertar valor de base de datos    
    encarg_parrafo10.add_run('. Se gestionó una programación directa de recursos en el PIA ')
    encarg_parrafo10.add_run(' por el monto de S/. ')
    encarg_parrafo10.add_run(pia_beneficios) #Insertar valor de base de datos    
    encarg_parrafo10.add_run(' lo cual fue comunicado a través del ')
    encarg_parrafo10.add_run('Oficio Múltiple N° 00011-2021-MINEDU/SPE-OPEP-UPP') #Esto cambiará cada año
    encarg_parrafo10.add_run('. Con Decreto Supremo N° 072-2021-EF')
    encarg_parrafo10.add_run(', se ha transferido ')   
    encarg_parrafo10.add_run(ds_072_beneficios) #Insertar valor de base de datos   
    encarg_parrafo10.add_run('. Mediante ')
    encarg_parrafo10.add_run('Decreto Supremo N° 256-2021-EF se realizó la segunda transferencia por')
    encarg_parrafo10.add_run(' S/ ')
    encarg_parrafo10.add_run(ds_256_beneficios) #Insertar valor de base de datos   
    encarg_parrafo10.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    #####################################################
    # Incluimos sección 8 financiamiento de plazas #
    #####################################################
        
    document.add_heading("Financiamiento de plazas 2021", level=2)
    
    racio_parrafo1 = document.add_paragraph(' A través del DS N° 078-2021-EF se financió ')
    racio_parrafo1.style = document.styles['Heading 5']
    racio_parrafo1.add_run(cantidad_plazas) #Insertar valor de base de datos 
    racio_parrafo1.add_run(' plazas de docentes de aula por un monto de S/ ') 
    racio_parrafo1.add_run(monto_plazas)
    racio_parrafo1.add_run('en el marco de los resultados del proceso de racionalización 2020 en la región de ') #Calcular valor de año anterior
    racio_parrafo1.add_run(region)
    racio_parrafo1.add_run('.') 
    racio_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    #####################################################
    # Incluimos sección 9 deudas sociales #
    #####################################################
    
    document.add_heading("Deudas Sociales", level=2)
    
    deuda_parrafo1 = document.add_paragraph("Mediante el Decreto Supremo N° 216-2021-EF se transfirió S/ ")
    deuda_parrafo1.style = document.styles['Heading 5']
    deuda_parrafo1.add_run(monto_deuda)
    deuda_parrafo1.add_run(" a la región ")
    deuda_parrafo1.add_run(region)
    deuda_parrafo1.add_run(".")
    
    #######################
    # Guardamos documento #
    #######################
    
    document.save(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
###########################################################
# Creamos tabla con lista de files para enviar por correo #
###########################################################
    
# Generamos lista de AM.
#lista_AM = glob.glob(os.path.join(proyecto, f"output/AM_corta_region/AM_{fecha_actual}/*"))

#lista_regiones = pd.DataFrame (lista_AM)
#lista_regiones.rename( columns={0:'path'}, inplace=True )
#lista_regiones[['a', 'b', 'c']] = lista_regiones["path"].str.split("AM_", expand = True)
#lista_regiones[['date', 'e']] = lista_regiones["b"].str.split("/", expand = True)
#lista_regiones[['region', 'g']] = lista_regiones["c"].str.split("_", expand = True)
#lista_regiones = lista_regiones[["path", "date","region"]]
#lista_regiones.to_excel(Path(proyecto, "documentacion", "lista_regiones.xlsx"), index = False)
