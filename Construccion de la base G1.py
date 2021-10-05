# -*- coding: utf-8 -*-
"""
Created on Mon Oct  4 15:42:30 2021

@author: Hugo Fernandez y Carlos Ramirez
"""

# Librerias

import docx
import pandas as pd
from datetime import datetime
from pyprojroot import here # pip install pypro
from janitor import clean_names # pip install pyjanitor
import os
import numpy as np

# Import las bases
os.chdir('D:\Minedu\AM-python-docx') # Directorio del proyecto 
df_tranf = pd.read_excel('input\Bases plazas\TRANSFERENCIAS 2021.xlsx',sheet_name = 'TRANSFERENCIAS')

    # Costeos
df_enc = pd.read_excel('input\Bases plazas\COSTEOS 2021.xlsx',sheet_name = 'ENC') # Encargatura
df_at = pd.read_excel('input\Bases plazas\COSTEOS 2021.xlsx',sheet_name = 'AT')   # Asignaciones temporales
# df_at = pd.read_excel('input\Bases plazas\COSTEOS 2021.xlsx',sheet_name = 'AT')   # Asignaciones temporales


# df_prog = pd.read_excel('input\Bases plazas\PROGRAMACIÓN 2021.xlsx',sheet_name = 'TRANSFERENCIAS')


# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS"]
uwu = 72515691

# For loop para cada región
for region in lista_regiones:
    
    
    #########################################################################
    # Incluimos el código del Documento
    document = docx.Document(here() / "input/Formato.docx") # Creación del documento en base al template
    title=document.add_heading('AM TEMAS PRESUPUESTALES - REGION ', 0) #Título del documento
    title.add_run(region)
    run = title.add_run()
    run.add_break()
    run = title.add_run()
    run.add_break()
    title.add_run(datetime.today().strftime('%d-%m-%y'))
    
    # Incluimos sección 1 de intervenciones pedagógicas
    document.add_heading("Sobre el financiamiento de conceptos remunerativos", level=1)
    run.underline = True
    document.add_heading("1.Pago de Encargaturas", level=1)
    
    ##Párrafos
    encarg_parrafo1 = document.add_paragraph(
    "Para la región ")
    encarg_parrafo1.add_run(region)
    encarg_parrafo1.add_run(', por concepto de encargaturas, se ha calculado para el ')
    encarg_parrafo1.add_run(datetime.today().strftime('%Y'))
    encarg_parrafo1.add_run(', un costo de S/.')
    encarg_parrafo1.add_run(f'{uwu}') #Insertar valor de base de datos
    encarg_parrafo1.add_run(
    ' que incluye la Jornada de Trabajo Adicional de 10 horas \
la carga social vinculada y la asignación por cargo de \
los profesores que asumen cargos de mayor responsabilidad \
mediante encargaturas')
    encarg_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo2 = document.add_paragraph('Para financiar estos conceptos, el')
    #Insertar valor del año pasado
    encarg_parrafo2.add_run('el MINEDU gestionó una programación directa de recursos\
    en el PIA 2021 de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo2.add_run(region)  
    encarg_parrafo2.add_run(' por el monto de S/.')
    encarg_parrafo2.add_run(', XXXXXXX') #Insertar valor de base de datos
    encarg_parrafo2.add_run('  en la finalidad 0267929 Pago de la asignación por jornada \
de trabajo adicional y asignación por cargo de mayor responsabilidad, \
la cuál es usada para financiar las encargaturas. Asimismo, el Pliego Regional \
ya contaba con una programación de ')
    encarg_parrafo2.add_run(', XXXXXXX') #Insertar valor de base de datos
    encarg_parrafo2.add_run(' en la misma finalidad y mediante ')
    encarg_parrafo2.add_run(' Oficio Múltiple N° 00082-2021-MINEDU/SPE-OPEP-UPP, ')
    encarg_parrafo2.add_run('se le solicitó a las Unidades Ejecutoras del Pliego Regional \
#realizar modificaciones presupuestarias por el monto de S/.')
    encarg_parrafo2.add_run(' XXXXXXX') #Insertar valor de base de datos
    encarg_parrafo2.add_run(' para habilitar la finalidad 0267929')
    encarg_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo3 = document.add_paragraph('Con ')
    encarg_parrafo3.add_run('Decreto Supremo 217-2021 \
publicado el 27 de agosto de 2021 en el marco de lo autorizado en el literal b) \
del numeral 40.1 de la Ley de Presupuesto 2021, se ha realizado una transferencia \
de partidas por el monto de S/.')  # Este párrafo tendrá que variar año tras año
    encarg_parrafo3.add_run(' XXXXXXX') #Insertar valor de base de datos    
    encarg_parrafo3.add_run(' a favor de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo3.add_run(region)
    encarg_parrafo3.add_run(' para financiar el costo diferencial.')
    encarg_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo4 = document.add_paragraph('En el siguiente cuadro se muestran el costo \
y los montos programados/transferidos a la Región ')
    encarg_parrafo4.add_run(region)
    encarg_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(here() / f'output/{region}_AM_GG1.docx')