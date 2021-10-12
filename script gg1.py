# -*- coding: utf-8 -*-
"""
Created on Thu Sep 30 10:51:08 2021
"""

# Importar librerías ----------------------------------------------------------
import docx
import pandas as pd
from datetime import datetime
from pyprojroot import here # pip install pypro
from janitor import clean_names # pip install pyjanitor

# Opciones --------------------------------------------------------------------

# Formato de tablas
pd.options.display.float_format = '${:,.0f}'.format

# Transformación de Datasets --------------------------------------------------

## Sobre el financiamiento de conceptos remunerativos
# C) Base de Encargaturas

# D) Base de Asignaciones Temporales

# E) Base de Beneficios Sociales

### Sobre el proceso de racionalización

#----------------------------------------------------------------------#
## Creación plazas docentes
data_creacion = pd.read_excel(here() / "input/Creacion 2021.xlsx", sheet_name="BD",  skiprows = 2)
data_creacion = clean_names(data_creacion) # Normalizamos nombres
data_creacion = data_creacion[['d_region', 'd_dreugel', 'creacion_total']].\
groupby(by = ["d_region", 'd_dreugel'] , as_index=False).sum()
data_creacion['d_region'] = data_creacion['d_region'].str.split(r'DRE ').str[-1]
data_creacion = data_creacion.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})


#----------------------------------------------------------------------#
## Brecha de plazas docentes
data_brecha = pd.read_excel(here() / "input/Brecha UGEL 2020.xlsx", sheet_name="Data")
# Normalizamos nombres


# Mantenemos variables de interés
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()
data_brecha['doc_e_n_cub_req'] = data_brecha['doc_e_n'] - data_brecha['nom_exd_mov1']
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']]
data_brecha.loc[data_brecha['brecha_net']<=0, 'req_neto'] = -1*data_brecha['brecha_net']
data_brecha.loc[data_brecha['brecha_net']>0, 'exc_neto'] = data_brecha['brecha_net']

#Cantidad de UGEL con requerimiento neto
data_brecha.loc[data_brecha['brecha_net']<0, 'cant_ugel_req'] = 1
data_brecha.loc[data_brecha['brecha_net']>0, 'cant_ugel_exc'] = 1

data_brecha_regional = data_brecha[['region', 'brecha_net', 'cant_ugel_req', 'cant_ugel_exc']].groupby(by = ["region"] , as_index=False).sum()
data_brecha_regional.loc[data_brecha_regional['brecha_net']<=0, 'brecha_net'] = -1*data_brecha_regional['brecha_net']
data_brecha_regional.loc[data_brecha_regional['brecha_net']>0, 'brecha_net'] = data_brecha_regional['brecha_net']

data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'req_neto', 'exc_neto', 'brecha_net']]


#----------------------------------------------------------------------#
## Bloqueo de plazas
data_bloqueo = pd.read_excel(here() / "input/Bloqueo 2020.xlsx")
data_bloqueo = clean_names(data_bloqueo) # Normalizamos nombres


# Mantenemos variables de interés
data_bloqueo = pd.read_excel(here() / "input/Bloqueo 2020.xlsx")
data_bloqueo['cant_bloqueos'] = 1
data_bloqueo = data_bloqueo[["descreg", 'cant_bloqueos']].groupby(by = ["descreg"] , as_index=False).sum()
data_bloqueo = data_bloqueo.rename(columns={'descreg':'region'})

#----------------------------------------------------------------------#
## Deuda social
data_deuda_social = pd.read_excel(here() / "input/Deudas sociales.xlsx")
data_deuda_social = clean_names(data_deuda_social) # Normalizamos nombres

# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS", "AREQUIPA", "TACNA"]

# For loop para cada región
for region in lista_regiones:
    
    
    #########################################################################
    # Incluimos el código del Documento
    document = docx.Document(here() / "input/Formato.docx") # Creación del documento en base al template
    title=document.add_heading('AM TEMAS PRESUPUESTALES - REGION ', 0) #Título del documento
    title.add_run(region)
    run = title.add_run()
    run.add_break()
    run = title.add_run()
    run.add_break()
    title.add_run(datetime.today().strftime('%d-%m-%y'))
    
    # Incluimos sección 1 de intervenciones pedagógicas
    document.add_heading("Sobre el financiamiento de conceptos remunerativos", level=1)
    run.underline = True
    document.add_heading("1.Pago de Encargaturas", level=1)
    
    ##Párrafos
    encarg_parrafo1 = document.add_paragraph(
    "Para la región ")
    encarg_parrafo1.add_run(f'{region}, por concepto de Asignaciones Temporales \
por prestar servicios en condiciones especiales') 
    encarg_parrafo1.add_run('se ha calculado para el ')
    encarg_parrafo1.add_run(datetime.today().strftime('%Y'))
    encarg_parrafo1.add_run(', un costo de S/.')
    encarg_parrafo1.add_run(', XXXXXXX') #Insertar valor de base de datos
    encarg_parrafo1.add_run(
    ' que incluye la Jornada de Trabajo Adicional de 10 horas \
la carga social vinculada y la asignación por cargo de \
los profesores que asumen cargos de mayor responsabilidad \
mediante encargaturas')
    nota1 = encarg_parrafo1.add_run('[1]')
    nota1.font.superscript = True

    encarg_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo2 = document.add_paragraph('Para financiar estos conceptos, el')
    #Insertar valor del año pasado
    encarg_parrafo2.add_run('el MINEDU gestionó una programación directa de recursos\
    en el PIA 2021 de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo2.add_run(region)  
    encarg_parrafo2.add_run(' por el monto de S/.')
    encarg_parrafo2.add_run(', XXXXXXX') #Insertar valor de base de datos
    encarg_parrafo2.add_run('  en la finalidad 0267929 Pago de la asignación por jornada \
de trabajo adicional y asignación por cargo de mayor responsabilidad, \
la cuál es usada para financiar las encargaturas. Asimismo, el Pliego Regional \
ya contaba con una programación de ')
    encarg_parrafo2.add_run(', XXXXXXX') #Insertar valor de base de datos
    nota2 = encarg_parrafo2.add_run('[2]')
    nota2.font.superscript = True
    
    encarg_parrafo2.add_run(' en la misma finalidad y mediante ')
    encarg_parrafo2.add_run(' Oficio Múltiple N° 00082-2021-MINEDU/SPE-OPEP-UPP, ')
    encarg_parrafo2.add_run('se le solicitó a las Unidades Ejecutoras del Pliego Regional \
#realizar modificaciones presupuestarias por el monto de S/.')
    encarg_parrafo2.add_run(' XXXXXXX') #Insertar valor de base de datos
    encarg_parrafo2.add_run(' para habilitar la finalidad 0267929')
    nota3 = encarg_parrafo2.add_run('[3]')
    nota3.font.superscript = True
    
    encarg_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo3 = document.add_paragraph('Con ')
    encarg_parrafo3.add_run('Decreto Supremo 217-2021 \
publicado el 27 de agosto de 2021 en el marco de lo autorizado en el literal b) \
del numeral 40.1 de la Ley de Presupuesto 2021, se ha realizado una transferencia \
de partidas por el monto de S/.')  # Este párrafo tendrá que variar año tras año
    encarg_parrafo3.add_run(' XXXXXXX') #Insertar valor de base de datos    
    encarg_parrafo3.add_run(' a favor de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo3.add_run(region)
    encarg_parrafo3.add_run(' para financiar el costo diferencial.')
    encarg_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo4 = document.add_paragraph('En el siguiente cuadro se muestran el costo \
y los montos programados/transferidos a la Región ')
    encarg_parrafo4.add_run(region)
    encarg_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

#----------------------------------------------------------------------------------------------#
    document.add_heading("2.Pago de Asignaciones Temporales", level=1)
    
    ##Párrafos
    encarg_parrafo5 = document.add_paragraph('Para la Región ')    
    encarg_parrafo5.add_run(region)
    encarg_parrafo5.add_run(', por concepto de Asignaciones Temporales \
por prestar servicios en condiciones especiales, se ha calculado para el 2021 un costo de \
S/ ')    
    encarg_parrafo5.add_run(' XXXXXXX') #Insertar valor de base de datos    
    encarg_parrafo5 = document.add_paragraph('que incluye el pago por prestar servicios \
en zonas rurales, de frontera, VRAEM, Instituciones Educativas Unidocentes, Multigrado \
Bilingüe y acreditar dominio de lengua originaria, de los profesores y auxiliares de \
educación nombrados y contratados.')        
    encarg_parrafo5.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    
    encarg_parrafo6 = document.add_paragraph('Para financiar estos conceptos, el ')
    encarg_parrafo6.add_run('2020 ') #Calcular valor de año anterior
    encarg_parrafo6.add_run('el MINEDU gestionó una programación directa de recursos \
en el PIA ')
    encarg_parrafo6.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo6.add_run(' de las Unidades Ejecutoras de Educación de la Región ')    
    encarg_parrafo6.add_run('por el monto de S/.')
    encarg_parrafo6.add_run(' XXXXXXX') #Insertar valor de base de datos    

    encarg_parrafo6.add_run(' en la finalidad 0267928. Pago de las asignaciones \
por tipo y ubicacion de Institucion Educativa la cuál es usada para financiar \
las asignaciones temporales. Asimismo, el Pliego Regional ya contaba con \
una programación de S/.')    
    encarg_parrafo6.add_run(' XXXXXXX') #Insertar valor de base de datos    
    nota4 = encarg_parrafo6.add_run('[4]')
    nota4.font.superscript = True 
    encarg_parrafo6.add_run(' en la misma finalidad.')

    encarg_parrafo6.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
  
    encarg_parrafo7 = document.add_paragraph('Con Decreto Supremo 187-2021 \
publicado el 22 de julio de 2021 en el marco de lo autorizado en los literales \
a), c), d) y e) del numeral 40.1 de la Ley de Presupuesto 2021, ') # Este párrafo tendrá que variar año tras año
    encarg_parrafo7.add_run('se ha realizado una transferencia \
de partidas por el monto de S/. ')
    encarg_parrafo7.add_run(' XXXXXXX') #Insertar valor de base de datos    
    encarg_parrafo7.add_run(' a favor de las Unidades Ejecutoras de Educación \
de la Región ') 
    encarg_parrafo7.add_run(region)
    encarg_parrafo7.add_run(' para financiar el costo diferencial de las asignaciones \
temporales a favor los profesores y auxiliares de educación nombrados y contratados.')    
    encarg_parrafo7.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo8 = document.add_paragraph('Actualmente está en gestión en el MINEDU \
la segunda transferencia de recursos por concepto de asignaciones temporales, \
el cual debería realizarse antes del ')
    encarg_parrafo8.add_run('26 de noviembre del 2021') #Esta fecha se actualizará año a año 
    encarg_parrafo8.add_run(' de acuerdo al plazo legal establecido en la Ley de Presupuesto') 
    encarg_parrafo8.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo8.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    encarg_parrafo9 = document.add_paragraph('En el siguiente cuadro se muestran el costo y los \
montos programados/transferidos a la Región ')
    encarg_parrafo9.add_run(region)
    encarg_parrafo9.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

#----------------------------------------------------------------------------------------------#

    document.add_heading("3.Pago de Beneficios Sociales", level=1)

    ##Párrafos    
    encarg_parrafo10 = document.add_paragraph('Para la Región ')
    encarg_parrafo10.add_run(region)
    encarg_parrafo10.add_run(', de acuerdo con la nueva estrategia para el pago oportuno \
de Beneficios Sociales implementada por el MINEDU, se han aprobado pagos por concepto de \
Asignación por Tiempo de Servicios (ATS), Compensación por Tiempo de Servicios (CTS) y, \
Subsidio por Luto y Sepelio (SLS) hasta por  un costo de S/ ')
    encarg_parrafo10.add_run(' XXXXXXX') #Insertar valor de base de datos    
    encarg_parrafo10.add_run(' a la fecha , a favor de los profesores y auxiliares de educación \
nombrados y contratados.')    
    encarg_parrafo10.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo11 = document.add_paragraph('Para financiar estos conceptos, el ')
    encarg_parrafo11.add_run('2020 ') #Calcular valor de año anterior
    encarg_parrafo11.add_run('el MINEDU gestionó una programación directa de recursos en el PIA ')
    encarg_parrafo11.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo11.add_run(' de las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo11.add_run(region)
    encarg_parrafo11.add_run(' por el monto de S/. ')
    encarg_parrafo11.add_run(' XXXXXXX.') #Insertar valor de base de datos    
    encarg_parrafo11.add_run('  Lo cual fue comunicado a través del ')
    encarg_parrafo11.add_run('Oficio Múltiple N° 00011-2021-MINEDU/SPE-OPEP-UPP') #Esto cambiará cada año
    encarg_parrafo11.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo12 = document.add_paragraph('Con ')
    encarg_parrafo12.add_run('Decreto Supremo 072-2021 publicado el 21 de abril de 2021')
    encarg_parrafo12.add_run('en el marco de lo autorizado en los literales a), d) y e) \
del numeral 40.1 de la Ley de Presupuesto ')
    encarg_parrafo12.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo12.add_run(', se ha realizado una transferencia de partidas por el monto de S/ ')   
    encarg_parrafo12.add_run(' XXXXXXX.') #Insertar valor de base de datos   
    encarg_parrafo12.add_run(', a favor de las Unidades Ejecutoras de Educación de la Región ')   
    encarg_parrafo12.add_run(region)
    encarg_parrafo12.add_run(' para financiar el pago de los beneficios sociales a favor los profesores \
y auxiliares de educación nombrados y contratados que fueron reconocidos hasta el ')   
    encarg_parrafo12.add_run('2020') #Calcular valor de año anterior
    encarg_parrafo12.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo13 = document.add_paragraph('Asimismo, mediante ')
    encarg_parrafo13.add_run('Decreto Supremo 256-2021 publicado el 24 de setiembre de 2021, ')  #Esto cambiará cada año
    encarg_parrafo13.add_run(', se realizó la segunda transferencia de recursos por concepto de \
beneficios sociales a favor de docentes y auxiliares nombrados y contratados, cuyos beneficios fueron \
reconocidos durante el año ')
    encarg_parrafo13.add_run(datetime.today().strftime('%Y')) #Año actual
    encarg_parrafo13.add_run(' transfiriéndose S/. ')
    encarg_parrafo13.add_run(' XXXXXXX.') #Insertar valor de base de datos   
    encarg_parrafo13.add_run('  a las Unidades Ejecutoras de Educación de la Región ')
    encarg_parrafo13.add_run(region)
    encarg_parrafo13.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
   

    encarg_parrafo14 = document.add_paragraph('En el siguiente cuadro se muestran el costo y los montos \
programados/transferidos a la Región ')
    encarg_parrafo14.add_run(region)
    encarg_parrafo14.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo15 = document.add_paragraph(' De la misma forma, durante el presente año, para la Región ')  
    encarg_parrafo15.add_run(region)
    encarg_parrafo15.add_run(' se ha realizado las siguientes transferenciasn')
    
    #Insertar tabla
    
    encarg_parrafo15.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    encarg_parrafo16 = document.add_paragraph(' Por otro lado, para el año ')
    encarg_parrafo16.add_run('2022 ') #Calcular año posterior
    encarg_parrafo16.add_run('el MINEDU está gestionando la programación parcial de recursos en los \
presupuestos de las Unidades Ejecutoras para atender encargaturas, asignaciones temporales, \
beneficios sociales, entre otros y, el financiamiento restante, se realizará de manera oportuna el ')
    encarg_parrafo16.add_run('2022, ') #Calcular año posterior
    encarg_parrafo16.add_run('preferentemente antes que termine el primer semestre de dicho año fiscal.')    
    encarg_parrafo16.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    document.save(here() / f'output/{region}_AM_GG1.docx')

#------------------------------------------------------------------------------------------------------------#
#------------------------------------------------------------------------------------------------------------#    


    # Incluimos sección 2 de Proceso de racionalización
    region_seleccionada = data_creacion['region'] == region #Seleccionar region
    tabla_creacion = data_creacion[region_seleccionada]
    tabla_creacion_formato = data_creacion[region_seleccionada]
    creacion_region = str('{:,.0f}'.format(tabla_creacion["creacion_total"].sum()))
    

    formato_tabla_creacion = {
        "ugel": "{}",
        "creacion_total" : "{}",
        }
    tabla_creacion_formato = tabla_creacion_formato.transform({k: v.format for k, v in formato_tabla_creacion.items()})        

    
    document.add_heading("Sobre el proceso de racionalización", level=1)
    run.underline = True    

    document.add_heading("4. Financiamiento de plazas 2021, en el marco del \
proceso de racionalización 2020", level=1)

    #Párrafo 
    racio_parrafo1 = document.add_paragraph('A través del ',  style="List Bullet")
    racio_parrafo1.add_run('DS 078-2021-EF') #Esto cambia año tras año
    racio_parrafo1.add_run(' se financiaron ') 
    racio_parrafo1.add_run(creacion_region) #Insertar valor de base de datos 
    racio_parrafo1.add_run(' plazas de docentes de aula en el marco de los resultados del \
proceso de racionalización ') 
    racio_parrafo1.add_run('2020 ') #Calcular valor de año anterior
    racio_parrafo1.add_run('en servicios educativos públicos de la región ')
    racio_parrafo1.add_run(region)
    racio_parrafo1.add_run(' con la siguiente distribución por UGEL:')
    racio_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # Incluimos tabla creaciones
    tabla1_creacion = document.add_table(tabla_creacion_formato.shape[0]+1, tabla_creacion_formato.shape[1])
    tabla1_creacion.style = "Colorful List Accent 1"
    ## Header de la tabla
    row = tabla1_creacion.rows[0].cells
    row[0].text = "UGEL"
    row[1].text = "Número de creaciones"
    ## Contenido de la tabla
    for i in range(tabla_creacion_formato.shape[0]):
        for j in range(tabla_creacion_formato.shape[-1]):
            tabla1_creacion.cell(i+1,j).text = str(tabla_creacion_formato.values[i,j])

    document.add_heading("5. Resultados proceso de racionalización 2020", level=1)
 
    region_seleccionada2 = data_brecha['region'] == region #Seleccionar region    
    tabla_brecha = data_brecha[region_seleccionada2]
    tabla_brecha_formato = data_brecha[region_seleccionada2]
    excedentes_region = str('{:,.0f}'.format(tabla_brecha["doc_e"].sum()))
    requerimientos_region = str('{:,.0f}'.format(tabla_brecha["doc_req"].sum()))
    requerimiento_neto = str('{:,.0f}'.format(tabla_brecha["req_neto"].sum()))
    excedencia_neta = str('{:,.0f}'.format(tabla_brecha["exc_neto"].sum()))

    
    region_seleccionada3 = data_brecha_regional['region'] == region #Seleccionar region
    tabla_brecha2 = data_brecha_regional[region_seleccionada3]
    ugel_cant_req = str('{:,.0f}'.format(tabla_brecha2["cant_ugel_req"].sum()))
    ugel_cant_exc = str('{:,.0f}'.format(tabla_brecha2["cant_ugel_exc"].sum()))    
    brecha_region = str('{:,.0f}'.format(tabla_brecha2["brecha_net"].sum()))    

    
    formato_tabla_brecha = {
        "ugel": "{}",
        "doc_req" : "{}",
        "doc_e" : "{}",
        "doc_e_n_cub_req" : "{}",
        "nom_exd_mov1" : "{}",
        "doc_e_c" : "{}",
        "req_neto" : "{}",
        "exc_neto" : "{}",
        }
    tabla_brecha_formato = tabla_brecha_formato.transform({k: v.format for k, v in formato_tabla_brecha.items()})
    
    
    #Párrafo 
    racio_parrafo2 = document.add_paragraph('En el proceso de racionalización ',  style="List Bullet")    
    racio_parrafo2.add_run('2020 ') #Calcular valor de año anterior
    racio_parrafo2.add_run(', se identificó en la región ')
    racio_parrafo2.add_run(region)
    racio_parrafo2.add_run(' un total de ')
    negrita1 = racio_parrafo2.add_run(excedentes_region) #Insertar valor de base de datos
    negrita1.bold= True
    negrita2 = racio_parrafo2.add_run(' plazas de docentes de aula excedentes ')
    negrita2.bold= True
    racio_parrafo2.add_run('y ')
    negrita3 = racio_parrafo2.add_run(requerimientos_region) #Insertar valor de base de datos
    negrita3.bold= True    
    negrita4 = racio_parrafo2.add_run(' plazas de requerimiento. ')
    negrita4.bold= True
    racio_parrafo2.add_run('A partir de esos resultados, se procedió a calcular \
el requerimiento y la excedencia por UGEL y el agregado a nivel regional, \
ello se puede observar en las dos últimas columnas del siguiente cuadro: ')  

    # Incluimos tabla brecha
    tabla1_brecha = document.add_table(tabla_brecha_formato.shape[0]+1, tabla_brecha_formato.shape[1])
    tabla1_brecha.style = "Colorful List Accent 1"
    ## Header de la tabla
    row = tabla1_brecha.rows[0].cells
    row[0].text = "UGEL"
    row[1].text = "Total Req"
    row[2].text = "Total Exd"    
    row[3].text = "Exc. Docentes nombrados que pueden cubrir requerimiento"    
    row[4].text = "Exc. Docentes nombrados con dificultad de movimiento"
    row[5].text = "Req. Neto final"
    row[6].text = "Exc. Neto final"
    
    ## Contenido de la tabla
    for i in range(tabla_brecha_formato.shape[0]):
        for j in range(tabla_brecha_formato.shape[-1]):
            tabla1_brecha.cell(i+1,j).text = str(tabla_brecha_formato.values[i,j])

    racio_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    racio_parrafo3 = document.add_paragraph(' Por lo tanto, a nivel regional \
se contaba con una brecha interna de ',  style="List Bullet")    
    racio_parrafo3.add_run(requerimiento_neto) #Insertar valor de base de datos
    racio_parrafo3.add_run(' plazas en ')
    racio_parrafo3.add_run(ugel_cant_req) #Insertar valor de base de datos
    racio_parrafo3.add_run(' UGEL, y un excedente neto de \
plazas vacantes ascendente a ')
    racio_parrafo3.add_run(excedencia_neta) #Insertar valor de base de datos  
    racio_parrafo3.add_run(' plazas en ')   
    racio_parrafo3.add_run(ugel_cant_exc) #Insertar valor de base de datos
    racio_parrafo3.add_run(' UGEL. Con ello, se obtuvo ')
      
    if region == "AYACUCHO" or region == "HUANCAVELICA" or region=="PUNO" or region=="TUMBES":       
        negrita5 = racio_parrafo3.add_run('una excedencia neta a nivel regional igul a ')
        negrita5.bold= True
        negrita6 = racio_parrafo3.add_run(brecha_region) #Insertar valor de base de datos
        negrita6.bold= True     
        racio_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    else:            
        negrita5 = racio_parrafo3.add_run('un requerimiento neto a nivel regional igul a ')
        negrita5.bold= True
        negrita6 = racio_parrafo3.add_run(brecha_region) #Insertar valor de base de datos
        negrita6.bold= True     
        racio_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_heading("6. Acciones de reordenamiento territorial 2020", level=1)
 
    #Párrafo
    if region == "AREQUIPA" or region == "AYACUCHO" or region=="CUSCO" or region=="HUANCAVELICA" or region=="PASCO" or region=="PUNO" or region=="TACNA":
        region_seleccionada3 = data_bloqueo['region'] == region #Seleccionar region    
        tabla_bloqueo = data_bloqueo[region_seleccionada3]
        bloqueos = str('{:,.0f}'.format(tabla_bloqueo['cant_bloqueos'].sum()))
        racio_parrafo4 = document.add_paragraph('En el marco del proceso de racionalización ',  style="List Bullet") 
        racio_parrafo4.add_run('2020 ') #Calcular valor de año anterior
        racio_parrafo4.add_run(', en la región ')
        racio_parrafo4.add_run(region)    
        racio_parrafo4.add_run(' se inhabilitaron ') 
        racio_parrafo4.add_run(bloqueos) 
        racio_parrafo4.add_run(' plazas de un total de ') 
        racio_parrafo4.add_run(excedencia_neta) #Insertar valor de base de datos
        racio_parrafo4.add_run(' plazas vacantes identificadas como excedencia neta.')
        racio_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    else:    
        racio_parrafo4 = document.add_paragraph('En el marco del proceso de racionalización ',  style="List Bullet") 
        racio_parrafo4.add_run('2020 ') #Calcular valor de año anterior
        racio_parrafo4.add_run(', en la región ')
        racio_parrafo4.add_run(region)    
        racio_parrafo4.add_run(' no se inhabilitaron plazas a pesar de contar con ')
        racio_parrafo4.add_run(excedencia_neta) #Insertar valor de base de datos
        racio_parrafo4.add_run(' plazas vacantes identificadas como excedencia neta.')
        racio_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

#------------------------------------------------------------------------------------------------------------#
#------------------------------------------------------------------------------------------------------------#    
    # Incluimos sección 3 deudas sociales y contratación
    tabla_deuda_social = data_deuda_social
    tabla_deuda_social_formato = data_deuda_social
    
    formato_tabla_deuda_social = {
        "seccion_pliego": "{}",
        "monto" : "{}",
        }
    tabla_deuda_social_formato = tabla_deuda_social_formato.transform({k: v.format for k, v in formato_tabla_deuda_social.items()})        

    document.add_heading("Sobre Deudas Sociales y Contratación del DL 276", level=1)
    run.underline = True  

    document.add_heading("7. Deudas Sociales", level=1)

    #Párrafo
    deuda_parrafo1 = document.add_paragraph(' La Deuda Social del Sector Educación \
está constituida por las obligaciones de pago  en materia laboral y previsional, \
respecto del personal Docente y Auxiliares de Educación provenientes \
de la derogada Ley del Profesorado  y su modificatoria; \
del personal Administrativo sujeto al Decreto Legislativo Nº 276, \
Ley de Bases de la Carrera Administrativa; del Personal Administrativo \
en el marco del Texto Único Ordenado (TUO) del Decreto Legislativo Nº 728 \
Ley de Productividad y Competitividad  Laboral  aprobado por \
Decreto Supremo Nº 003-97-TR. ') 
    deuda_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo2 = document.add_paragraph(' La atención de la deuda social \
con el Sector se viene efectuando a través de transferencias de \
partidas del Tesoro Público; es decir, en forma complementaria a \
los recursos presupuestales con que disponen los Pliegos del Gobierno Nacional \
y los Gobiernos Regionales, para la atención del pago de sentencias judiciales \
en calidad de cosa juzgada y en ejecución. Dicho tratamiento, \
se realiza desde el año 2014 en el marco de las Leyes Anuales de Presupuesto, \
habiéndose autorizado diversas habilitaciones presupuestarias \
en el nivel institucional mediante los decretos supremos correspondientes. ')
    deuda_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo3 = document.add_paragraph(' Desde el año 2018 se han venido \
destinando S/ 200 000 000,00 (DOSCIENTOS MILLONES y 00/100 DE SOLES) \
para el sector Educación, que según los criterios de priorización aprobados \
por el Ministerio de Educación, son destinados preferentemente \
al pago de las sentencias judiciales que reconocen \
el derecho de preparación de clases, frente a otros conceptos. ')
    deuda_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo4 = document.add_paragraph(' En ese sentido, mediante el numeral \
1.2 del artículo 1 del Decreto Supremo N° 216-2021-EF, publicado en el \
Diario Oficial “El Peruano” el 27 de agosto de 2021, se autorizó \
la Transferencia de Partidas en el Presupuesto del Sector Público \
para el Año Fiscal 2021, hasta por la suma de S/ 200 000 000,00 \
(DOSCIENTOS MILLONES Y 00/100 SOLES), a favor de diversos Pliegos \
del Gobierno Nacional (dentro de los que se encuentra el Ministerio de Educación \
(MINEDU), y los Gobiernos Regionales, para financiar el pago de \
sentencias judiciales en calidad de cosa juzgada del sector Educación \
y en ejecución al 31 de diciembre de 2020, en el marco del numeral 6 de \
la de la Undécima Disposición Complementaria Final de la Ley N° 31084, \
Ley de presupuesto del año fiscal 2021, con cargo a los recursos de la Reserva \
de Contingencia del Ministerio de Economía y Finanzas. \
El detalle de dicha transferencia de recursos se muestra a continuación:')
    deuda_parrafo4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # Incluimos tabla deuda social
    tabla_deuda_social = document.add_table(tabla_deuda_social_formato.shape[0]+1, tabla_deuda_social_formato.shape[1])
    tabla_deuda_social.style = "Colorful List Accent 1"
    ## Header de la tabla
    row = tabla_deuda_social.rows[0].cells
    row[0].text = "Sección / Pliego"
    row[1].text = "Monto"

   ## Contenido de la tabla
    for i in range(tabla_deuda_social_formato.shape[0]):
        for j in range(tabla_deuda_social_formato.shape[-1]):
            tabla_deuda_social.cell(i+1,j).text = str(tabla_deuda_social_formato.values[i,j])    
    
    deuda_parrafo5 = document.add_paragraph('Del mismo modo, es importante \
señalar que las deudas sociales del Sector Educación se atienden siguiendo \
los criterios de priorización establecidos en el presente año fiscal mediante \
el Decreto Supremo N° 003-2021-MINEDU, Decreto Supremo que aprueba \
los criterios de priorización para la atención del pago de sentencias judiciales \
en calidad de cosa juzgada y en ejecución del Sector Educación, \
donde no se establece una distinción entre regímenes laborales \
o entre personal activo o cesante, sino se siguen patrones de antigüedad, \
salud, entre otros')
    deuda_parrafo5.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY


    deuda_parrafo6 = document.add_paragraph('Es igualmente importante señalar, \
que en el caso de existir déficit presupuestal durante el ejercicio fiscal \
en ejecución, para la atención de las planillas continuas de pensionistas, \
derivadas de sentencias en calidad de cosa juzgada, \
su financiamiento deberá tener en cuenta la exoneración prevista \
en el numeral 9.2 del artículo 9. Medidas en materia de modificaciones \
Presupuestarias en el Nivel Funcional Programático, de la Ley Nº 31084, \
Ley de Presupuesto del Sector Público para el Año Fiscal 2021; \
en tanto prescribe que: ')

    formato1 = deuda_parrafo6.add_run(' “A nivel de pliego')
    formato1.italic = True
    formato2 = deuda_parrafo6.add_run(' la Partida de Gasto 2.2.1 “Pensiones” no puede ser habilitadora, ')
    formato2.italic = True
    formato2.bold = True
    formato3 = deuda_parrafo6.add_run('salvo para las habilitaciones que se realicen ')
    formato3.italic = True
    formato3.bold = True
    formato3.underline = True
    formato4 = deuda_parrafo6.add_run('dentro de la misma partida entre unidades \
ejecutoras del mismo pliego presupuestario, y ')
    formato4.italic = True
    formato4.bold = True    
    formato5 = deuda_parrafo6.add_run('para la atención de sentencias judiciales \
en materia pensionaria con calidad de cosa juzgada,')
    formato5.italic = True
    formato5.bold = True
    formato5.underline = True    
    formato6 = deuda_parrafo6.add_run(' en este último caso, previo informe \
favorable de la Dirección General de Presupuesto Público (DGPP), ')
    formato6.italic = True
    formato6.bold = True  
    formato7 = deuda_parrafo6.add_run(' y de corresponder, sobre la base \
de la información registrada en el Aplicativo Informático para el \
 Registro Centralizado de Planillas y de Datos de los Recursos Humanos \
 del Sector Público (AIRHSP) que debe remitir la Dirección General \
 de Gestión Fiscal de los Recursos Humanos a la DGPP.(…)”.')
    formato7.italic = True
    deuda_parrafo6.add_run('(Lo resaltado es nuestro).')
    deuda_parrafo6.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    deuda_parrafo7 = document.add_paragraph('En el mismo numeral del artículo \
acotado en el párrafo anterior, se establece que las solicitudes de \
informe favorable para su aplicación solo pueden ser presentadas al \
Ministerio de Economía y Finanzas hasta el 15 de octubre de 2021; por lo que, ')
    formato8 = deuda_parrafo7.add_run('las unidades orgánicas pertinentes \
del Pliego Presupuestal involucrado, deben implementar \
las acciones administrativas a que hubiera lugar con sujeción \
a la normatividad presupuestaria y al plazo legal invocados líneas atrás.')
    formato8.bold = True
    deuda_parrafo7.add_run(' Así como, ')
    formato9 = deuda_parrafo7.add_run('de corresponder, \
las acciones pertinentes para el pago de las deudas que se hubieran generado \
por aplicación de la sentencia judicial en calidad de cosa juzgada \
que así lo ordene, en cuyo caso deberá observarse lo dispuesto en \
la Ley Nº 30137 y su Reglamento.')
    formato9.underline = True
    deuda_parrafo7.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_heading("8. Ruta para contratación de personal administrativo", level=1)

    #Párrafo
    pers_adm_parrafo1 = document.add_paragraph('En el marco del inciso \
d) del numeral 8.1 del artículo 8 de la Ley N° 31084, \
Ley de Presupuesto del Sector Público para el año 2021, ')
    formato10 = pers_adm_parrafo1.add_run('no se ha establecido como excepción \
la posibilidad del ingreso de personal administrativo \
bajo el régimen laboral del Decreto Legislativo N° 276.')
    formato10.underline = True
    pers_adm_parrafo1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    pers_adm_parrafo2 = document.add_paragraph('A su vez, \
la Gerencia de Políticas de Gestión del Servicio Civil de \
la Autoridad Nacional del Servicio Civil (SERVIR) se ha pronunciado \
sobre la imposibilidad de contratar personal administrativo del \
Decreto Legislativo N° 276, emitido mediante Informe Técnico \
N° 331-2021-SERVIR-GPGSC.')
    pers_adm_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    pers_adm_parrafo3 = document.add_paragraph('Dado que el Decreto Legislativo \
N° 276, no solo rige al Sector Educación, sino a todo el Sector Público, \
desde el Ministerio de Educación se propuso incluir la habilitación \
para su contratación en una norma con rango de Ley, la cual autorizaría \
la contratación de personal del Decreto Legislativo N° 276. \
Sin embargo, dicha propuesta no fue incluida en el \
Proyecto de Ley del Presupuesto del Sector Público para el año 2022, \
enviada al Congreso de la República para su aprobación el 31 de agosto de 2021.')
    pers_adm_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_page_break()

#--------------------------------------------------------------------------------#
    document.add_heading("Anexos", level=1)
    nota1 = document.add_paragraph('[1] Este costeo está en función a la información \
registrada en el sistema de control de plazas NEXUS')
    nota1.italic = True
    nota1.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    nota2 = document.add_paragraph('[2] Solo se considera los clasificadores \
21.12.11 (JTA), 21.31.15 (CS) y 21.12.21 (Asignación por cargo).')
    nota2.italic = True
    nota2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    nota3 = document.add_paragraph('[3] Mediante el OM, se solicitó habilitar \
la finalidad 0267929 con los recursos destinados a financiar el costo diferencial \
de la asignación por jornada de trabajo adicional y carga social \
(debido al incremento de la RIM), ya que estos estaban programados en \
finalidades que se usaban anteriormente.')
    nota3.italic = True
    nota3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    nota4 = document.add_paragraph('[4] Se considera el clasificador 21.12.21')
    nota4.italic = True
    nota4.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(here() / f'output/{region}_AM_GG1.docx')
    